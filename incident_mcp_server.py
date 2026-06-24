"""
事件复盘分析MCP服务器 - 生产事件情况表的读取与分析报告生成

工作流位置：
  第1步（MCP节点）：读取生产事件情况表（Excel），输出结构化JSON数据
  第2步（LLM节点）：由大模型根据结构化数据进行复盘分析（外部Dify工作流）
  第3步（MCP节点）：将LLM的分析复盘结果输出为Word文档

依赖安装：
    pip install fastmcp python-docx openpyxl

启动方式（通过main.py统一启动，无需单独运行此文件）：
    python main.py [--transport sse] [--host 0.0.0.0] [--port 18080]

在Dify工作流中使用：
  第1步 - 调用 read_incident_events 工具读取事件情况表
  第2步 - LLM节点分析（在Dify中配置提示词）
  第3步 - 调用 generate_incident_report 工具生成复盘报告Word文档
"""

import json
import os
import sys
import re
import logging
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger('IncidentMCPServer')

# FastMCP导入 - 优先使用共享MCP实例（单端口模式），fallback到独立实例
try:
    from mcp_instance import mcp
    FASTMCP_AVAILABLE = True
    MCP_SHARED_MODE = True  # 共享模式标志
except ImportError:
    try:
        from fastmcp import FastMCP
        FASTMCP_AVAILABLE = True
        MCP_SHARED_MODE = False  # 独立模式
        mcp = FastMCP("IncidentMCPServer")
    except ImportError:
        FASTMCP_AVAILABLE = False
        MCP_SHARED_MODE = False  # 独立模式
        logger.warning("mcp_instance模块和fastmcp均未找到，请确保fastmcp已安装")

# python-docx库（用于生成Word报告）
try:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor, Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx未安装，请运行: pip install python-docx")

# openpyxl库（用于读取Excel事件情况表）
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl未安装，请运行: pip install openpyxl")


# ============================================================
# 事件情况表读取器
# ============================================================

class IncidentEventReader:
    """生产事件情况表读取器 - 从Excel文件读取事件数据并输出结构化JSON"""

    # 期望的Excel列标题映射（支持模糊匹配）
    EXPECTED_HEADER_MAP = {
        "事件编号": "incident_id",
        "事件名称": "incident_name",
        "发生时间": "start_time",
        "结束时间": "end_time",
        "影响时长": "impact_duration_minutes",
        "事件等级": "severity",
        "影响系统": "affected_systems",
        "影响范围": "impact_scope",
        "事件描述": "description",
        "应急处置方式": "emergency_response",
        "处置时长": "response_duration_minutes",
        "根本原因": "root_cause",
        "监控告警情况": "monitoring_alert_status",
        "应急预案": "emergency_plan"
    }

    # 常见同义变体映射
    SYNONYM_MAP = {
        "事件id": "事件编号",
        "事件标题": "事件名称",
        "开始时间": "发生时间",
        "结束时间": "结束时间",
        "影响时长(分钟)": "影响时长",
        "影响时长（分钟）": "影响时长",
        "影响时长min": "影响时长",
        "事件级别": "事件等级",
        "影响系统模块": "影响系统",
        "影响范围描述": "影响范围",
        "事件详细描述": "事件描述",
        "处理方式": "应急处置方式",
        "应急处置措施": "应急处置方式",
        "处置时长(分钟)": "处置时长",
        "处置时长（分钟）": "处置时长",
        "处置时长min": "处置时长",
        "原因分析": "根本原因",
        "监控告警": "监控告警情况",
        "告警情况": "监控告警情况",
        "应急预案情况": "应急预案"
    }

    @staticmethod
    def read_events(file_path: str) -> Dict[str, Any]:
        """
        从Excel文件中读取生产事件情况

        Args:
            file_path: Excel文件路径（支持.xlsx和.xls格式）

        Returns:
            Dict包含events（事件列表）、summary（汇总信息）和metadata（元数据）
        """
        logger.info(f"开始读取生产事件情况表: {file_path}")

        if not file_path or not file_path.strip():
            return {
                "success": False,
                "error": "file_path不能为空",
                "error_type": "ValidationError"
            }

        file_path = file_path.strip()

        if not os.path.exists(file_path):
            return {
                "success": False,
                "error": f"文件不存在: {file_path}",
                "error_type": "FileNotFoundError"
            }

        ext = Path(file_path).suffix.lower()
        if ext not in ['.xlsx', '.xls']:
            return {
                "success": False,
                "error": f"不支持的文件格式: {ext}，仅支持 .xlsx、.xls",
                "error_type": "UnsupportedFormatError"
            }

        if not OPENPYXL_AVAILABLE:
            return {
                "success": False,
                "error": "openpyxl库未安装，请运行: pip install openpyxl",
                "error_type": "ImportError"
            }

        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            ws = wb.active
            if ws is None:
                return {
                    "success": False,
                    "error": "Excel文件中没有工作表",
                    "error_type": "EmptyWorkbookError"
                }

            # 读取所有行
            rows = list(ws.iter_rows(values_only=True))
            wb.close()

            if not rows or len(rows) < 2:
                return {
                    "success": False,
                    "error": "Excel文件为空或只有表头没有数据",
                    "error_type": "EmptyFileError"
                }

            # 解析表头
            header_row = rows[0]
            header_mapping = IncidentEventReader._map_headers(header_row)

            logger.info(f"表头映射: {header_mapping}")

            if not header_mapping:
                return {
                    "success": False,
                    "error": f"无法识别Excel表头，期望的列包括: {', '.join(IncidentEventReader.EXPECTED_HEADER_MAP.keys())}",
                    "error_type": "HeaderParseError"
                }

            # 解析数据行
            events = []
            parse_errors = []

            for row_idx, row in enumerate(rows[1:], start=2):
                if all(cell is None or (isinstance(cell, str) and cell.strip() == '') for cell in row):
                    continue

                event_dict = {}
                valid_row = True
                for col_idx, (field_name, field_key) in header_mapping.items():
                    value = row[col_idx] if col_idx < len(row) else None
                    if value is not None:
                        value = str(value).strip()
                    event_dict[field_key] = value or ""

                # 事件编号和事件名称为必填字段
                if not event_dict.get("incident_id"):
                    parse_errors.append(f"第{row_idx}行缺少事件编号，跳过")
                    valid_row = False

                if valid_row:
                    events.append(event_dict)

            # 构建汇总信息
            total_count = len(events)

            # 按事件等级统计
            severity_stats = {}
            for ev in events:
                sev = ev.get("severity", "未知")
                severity_stats[sev] = severity_stats.get(sev, 0) + 1

            # 按影响系统统计
            system_stats = {}
            for ev in events:
                sys_str = ev.get("affected_systems", "")
                for s in sys_str.split(","):
                    s = s.strip()
                    if s:
                        system_stats[s] = system_stats.get(s, 0) + 1

            # 统计总影响时长
            total_impact_minutes = 0
            for ev in events:
                duration_str = ev.get("impact_duration_minutes", "0")
                try:
                    total_impact_minutes += int(float(duration_str))
                except (ValueError, TypeError):
                    pass

            result = {
                "success": True,
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "total_events": total_count,
                "summary": {
                    "total_events": total_count,
                    "severity_distribution": severity_stats,
                    "affected_systems": system_stats,
                    "total_impact_minutes": total_impact_minutes,
                    "parse_errors_count": len(parse_errors)
                },
                "events": events,
                "parse_errors": parse_errors if parse_errors else None,
                "columns": list(IncidentEventReader.EXPECTED_HEADER_MAP.keys())
            }

            logger.info(
                f"事件读取完成: {total_count}条事件, "
                f"{len(severity_stats)}种等级, {len(system_stats)}个影响系统"
            )
            return result

        except Exception as e:
            logger.exception(f"读取事件情况表失败: {e}")
            return {
                "success": False,
                "error": f"读取事件情况表失败: {str(e)}",
                "error_type": type(e).__name__
            }

    @staticmethod
    def _map_headers(header_row: tuple) -> Dict[int, str]:
        """
        将Excel表头映射到内部字段名

        Args:
            header_row: Excel的第一行（表头）

        Returns:
            Dict: {列索引: (原始表头, 内部字段名)}
        """

        mapping = {}

        # 构建规范化的期望表头：仅保留中文字符和英文
        normalized_expected = {}
        for header_cn, field_key in IncidentEventReader.EXPECTED_HEADER_MAP.items():
            # 提取所有非标点字符用于匹配
            normalized = re.sub(r'[（(）)\)\s\-_∶:，,、]', '', header_cn)
            normalized_expected[normalized] = (header_cn, field_key)

        # 遍历Excel实际表头列
        for col_idx, cell_value in enumerate(header_row):
            if cell_value is None:
                continue

            actual_header = str(cell_value).strip()
            if not actual_header:
                continue

            logger.debug(f"  匹配表头列[{col_idx}]: '{actual_header}'")

            # 第1步：精确匹配
            if actual_header in IncidentEventReader.EXPECTED_HEADER_MAP:
                mapping[col_idx] = (actual_header, IncidentEventReader.EXPECTED_HEADER_MAP[actual_header])
                logger.debug(f"    精确匹配: {actual_header}")
                continue

            # 第2步：同义变体匹配
            if actual_header in IncidentEventReader.SYNONYM_MAP:
                canonical = IncidentEventReader.SYNONYM_MAP[actual_header]
                if canonical in IncidentEventReader.EXPECTED_HEADER_MAP:
                    mapping[col_idx] = (actual_header, IncidentEventReader.EXPECTED_HEADER_MAP[canonical])
                    logger.debug(f"    同义匹配: {actual_header} -> {canonical}")
                    continue

            # 第3步：去除非中文/英文字符后匹配
            actual_normalized = re.sub(r'[（(）)\)\s\-_∶:，,、]', '', actual_header)
            if actual_normalized in normalized_expected:
                canonical, field_key = normalized_expected[actual_normalized]
                mapping[col_idx] = (actual_header, field_key)
                logger.debug(f"    归一化匹配: {actual_header} -> {canonical}")
                continue

            # 第4步：部分匹配（字段至少有3个字符）
            if len(actual_normalized) >= 3:
                for norm_key, (canonical, field_key) in normalized_expected.items():
                    # 检查实际表头是否包含期望表头的核心部分，或反之
                    if len(actual_normalized) >= 3 and (actual_normalized in norm_key or norm_key in actual_normalized):
                        mapping[col_idx] = (actual_header, field_key)
                        logger.debug(f"    部分匹配: {actual_header} -> {canonical}")
                        break

        return mapping


# ============================================================
# 复盘报告生成器
# ============================================================

class IncidentReportGenerator:
    """复盘分析报告生成器 - 将LLM分析结果输出为Word文档"""

    # 安全输出目录（默认当前目录下的 reports 文件夹）
    DEFAULT_OUTPUT_DIR = "reports"

    @staticmethod
    def generate_report(
        analysis_result: str,
        output_path: Optional[str] = None,
        incident_file_name: str = ""
    ) -> Dict[str, Any]:
        """
        根据LLM分析结果生成Word格式的复盘报告

        Args:
            analysis_result: LLM分析结果的JSON字符串
            output_path: 输出Word文件路径（可选，默认在reports目录下自动生成文件名）
            incident_file_name: 原始事件文件名（用于报告标题）

        Returns:
            Dict包含生成结果和文件路径
        """
        start_time = datetime.now()
        call_id = f"report_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始生成复盘报告")

        if not analysis_result or not analysis_result.strip():
            return {
                "success": False,
                "error": "analysis_result不能为空",
                "error_type": "ValidationError",
                "call_id": call_id
            }

        # 解析JSON分析结果
        try:
            analysis_data = json.loads(analysis_result)
        except json.JSONDecodeError as e:
            return {
                "success": False,
                "error": f"analysis_result不是有效的JSON: {str(e)}",
                "error_type": "JSONDecodeError",
                "call_id": call_id
            }

        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx库未安装，请运行: pip install python-docx",
                "error_type": "ImportError",
                "call_id": call_id
            }

        try:
            # 确定输出路径
            if output_path:
                output_path = output_path.strip()
                output_dir = str(Path(output_path).parent)
                os.makedirs(output_dir, exist_ok=True)
            else:
                output_dir = IncidentReportGenerator.DEFAULT_OUTPUT_DIR
                os.makedirs(output_dir, exist_ok=True)
                timestamp = start_time.strftime("%Y%m%d_%H%M%S")
                output_path = os.path.join(output_dir, f"事件复盘报告_{timestamp}.docx")

            # 创建Word文档
            doc = DocxDocument()

            # 设置默认字体
            style = doc.styles['Normal']
            font = style.font
            font.name = '微软雅黑'
            font.size = Pt(11)
            # 设置中文字体
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            # 文档标题
            title_text = "生产事件复盘分析报告"
            if incident_file_name:
                base_name = Path(incident_file_name).stem
                title_text = f"《{base_name}》复盘分析报告"

            para = doc.add_paragraph()
            run = para.add_run(title_text)
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0, 51, 102)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 分隔线
            doc.add_paragraph("─" * 50)

            # 生成日期
            p = doc.add_paragraph()
            run = p.add_run(f"生成日期：{start_time.strftime('%Y-%m-%d %H:%M:%S')}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")

            # ============================================================
            # 第一部分：事件总体总结
            # ============================================================
            IncidentReportGenerator._add_section_heading(doc, "一、事件总体总结")

            summary = analysis_data.get("summary", {})
            if isinstance(summary, dict):
                # 一句话概述
                IncidentReportGenerator._add_field_row(doc, "事件概述", summary.get("overview", ""))
                # 影响时长
                IncidentReportGenerator._add_field_row(doc, "影响时长", summary.get("impact_duration", ""))
                # 事件等级
                IncidentReportGenerator._add_field_row(doc, "事件等级", summary.get("severity", ""))
                # 涉及系统
                IncidentReportGenerator._add_field_row(doc, "涉及系统", summary.get("involved_systems", ""))
                # 影响范围
                IncidentReportGenerator._add_field_row(doc, "影响范围", summary.get("impact_scope", ""))
            else:
                p = doc.add_paragraph(str(summary))

            doc.add_paragraph("")

            # ============================================================
            # 第二部分：深入的原因分析
            # ============================================================
            IncidentReportGenerator._add_section_heading(doc, "二、深入原因分析")

            cause_analysis = analysis_data.get("cause_analysis", {})
            if isinstance(cause_analysis, dict):
                # 直接原因
                IncidentReportGenerator._add_subsection_text(doc, "直接原因", cause_analysis.get("direct_cause", ""))
                # 间接原因
                IncidentReportGenerator._add_subsection_text(doc, "间接原因", cause_analysis.get("indirect_cause", ""))
                # 根因
                IncidentReportGenerator._add_subsection_text(doc, "根本原因", cause_analysis.get("root_cause", ""))
            elif isinstance(cause_analysis, str):
                doc.add_paragraph(cause_analysis)
            else:
                p = doc.add_paragraph(str(cause_analysis))

            doc.add_paragraph("")

            # ============================================================
            # 第三部分：暴露出来的问题
            # ============================================================
            IncidentReportGenerator._add_section_heading(doc, "三、暴露出来的问题")

            issues = analysis_data.get("issues", {})
            if isinstance(issues, dict):
                issue_categories = [
                    ("监控告警方面", "monitoring_alerts"),
                    ("系统架构设计方面", "architecture_design"),
                    ("应急处置和应急预案方面", "emergency_response"),
                    ("其他方面", "others")
                ]
                for label, key in issue_categories:
                    content = issues.get(key, "")
                    if content:
                        IncidentReportGenerator._add_subsection_text(doc, label, content)
            elif isinstance(issues, list):
                for i, issue in enumerate(issues, 1):
                    if isinstance(issue, dict):
                        title = issue.get("category", f"问题{i}")
                        desc = issue.get("description", str(issue))
                        IncidentReportGenerator._add_subsection_text(doc, title, desc)
                    else:
                        doc.add_paragraph(f"{i}. {issue}")
            else:
                doc.add_paragraph(str(issues))

            doc.add_paragraph("")

            # ============================================================
            # 第四部分：改进建议
            # ============================================================
            IncidentReportGenerator._add_section_heading(doc, "四、改进建议")

            suggestions = analysis_data.get("suggestions", [])
            if isinstance(suggestions, list):
                for i, sug in enumerate(suggestions, 1):
                    if isinstance(sug, dict):
                        category = sug.get("category", f"建议{i}")
                        content = sug.get("content", "")
                        priority = sug.get("priority", "")

                        # 建议标题
                        p = doc.add_paragraph()
                        run = p.add_run(f"{i}. {category}")
                        run.bold = True
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 51, 102)

                        # 优先级
                        if priority:
                            priority_colors = {
                                "高": RGBColor(192, 0, 0),
                                "中": RGBColor(196, 128, 0),
                                "低": RGBColor(0, 128, 0)
                            }
                            p2 = doc.add_paragraph()
                            r = p2.add_run(f"   优先级：{priority}")
                            color = priority_colors.get(priority, RGBColor(128, 128, 128))
                            r.font.color.rgb = color
                            r.font.size = Pt(10)

                        # 具体内容
                        if content:
                            doc.add_paragraph(f"   {content}")
                    else:
                        doc.add_paragraph(f"{i}. {sug}")
            elif isinstance(suggestions, dict):
                for key, value in suggestions.items():
                    IncidentReportGenerator._add_subsection_text(doc, key, str(value))
            else:
                doc.add_paragraph(str(suggestions))

            doc.add_paragraph("")

            # ============================================================
            # 第五部分：应急处置方式与处置时长
            # ============================================================
            IncidentReportGenerator._add_section_heading(doc, "五、应急处置方式与处置时长")

            response_info = analysis_data.get("response_info", {})
            if isinstance(response_info, dict):
                IncidentReportGenerator._add_field_row(doc, "应急处置方式",
                    response_info.get("response_method", ""))
                IncidentReportGenerator._add_field_row(doc, "处置时长",
                    response_info.get("response_duration", ""))
                IncidentReportGenerator._add_field_row(doc, "处置效果评估",
                    response_info.get("effectiveness", ""))

                # 处置时间线（如果有）
                timeline = response_info.get("timeline", [])
                if isinstance(timeline, list) and timeline:
                    p = doc.add_paragraph()
                    run = p.add_run("处置时间线：")
                    run.bold = True
                    for step in timeline:
                        if isinstance(step, dict):
                            doc.add_paragraph(
                                f"  {step.get('time', '')} - {step.get('action', '')}"
                            )
                        else:
                            doc.add_paragraph(f"  {step}")
            else:
                doc.add_paragraph(str(response_info))

            doc.add_paragraph("")

            # ============================================================
            # 第六部分：原始事件数据（如果存在）
            # ============================================================
            original_data = analysis_data.get("original_events", [])
            if original_data:
                IncidentReportGenerator._add_section_heading(doc, "六、原始事件数据")
                doc.add_paragraph(f"共 {len(original_data)} 条事件记录")

                for i, ev in enumerate(original_data, 1):
                    if isinstance(ev, dict):
                        p = doc.add_paragraph()
                        r = p.add_run(f"事件{i}：{ev.get('incident_name', ev.get('incident_id', ''))}")
                        r.bold = True
                        r.font.size = Pt(11)

                        for key, label in [
                            ("incident_id", "事件编号"),
                            ("severity", "事件等级"),
                            ("start_time", "发生时间"),
                            ("affected_systems", "影响系统"),
                            ("impact_scope", "影响范围"),
                            ("description", "事件描述"),
                            ("root_cause", "根本原因")
                        ]:
                            val = ev.get(key, "")
                            if val:
                                doc.add_paragraph(f"  {label}：{val}")

            # 保存文档
            doc.save(output_path)
            file_size = os.path.getsize(output_path)

            logger.info(f"[{call_id}] 复盘报告已生成: {output_path} ({file_size} bytes)")

            return {
                "success": True,
                "call_id": call_id,
                "timestamp": start_time.isoformat(),
                "output_path": os.path.abspath(output_path),
                "file_name": Path(output_path).name,
                "file_size": file_size,
                "report_sections": [
                    "事件总体总结",
                    "深入原因分析",
                    "暴露出来的问题",
                    "改进建议",
                    "应急处置方式与处置时长"
                ],
                "incident_file_name": incident_file_name
            }

        except Exception as e:
            logger.exception(f"[{call_id}] 生成复盘报告失败: {e}")
            return {
                "success": False,
                "error": f"生成复盘报告失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }

    @staticmethod
    def _add_section_heading(doc, text: str):
        """添加章节标题"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0, 51, 102)
        # 添加底线
        run.underline = True
        doc.add_paragraph("")

    @staticmethod
    def _add_subsection_text(doc, title: str, content: str):
        """添加子标题和内容"""
        if not content:
            return

        p = doc.add_paragraph()
        r = p.add_run(f"【{title}】")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0, 70, 130)

        doc.add_paragraph(f"  {content}")

    @staticmethod
    def _add_field_row(doc, label: str, value: str):
        """添加带标签的内容行"""
        if not value:
            return

        p = doc.add_paragraph()
        r = p.add_run(f"{label}：")
        r.bold = True
        r.font.size = Pt(11)

        r2 = p.add_run(str(value))
        r2.font.size = Pt(11)


# ============================================================
# MCP工具定义 - 注册到共享MCP实例
# ============================================================

if FASTMCP_AVAILABLE:

    # ============================================================
    # 第1步：读取生产事件情况表
    # ============================================================

    @mcp.tool()
    def read_incident_events(file_path: str) -> str:
        """
        【第1步】读取生产事件情况表Excel文件，输出结构化JSON数据

        读取Excel事件表，输出包含事件列表（含编号、名称、时间、等级、描述等字段）、
        按等级和影响系统的汇总统计。

        Args:
            file_path: 事件情况表Excel文件路径（.xlsx或.xls格式）

        Returns:
            JSON字符串，包含 events（事件列表）、summary（汇总统计）、
            total_events（事件总数）等信息
        """
        start_time = datetime.now()
        call_id = f"read_events_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 读取事件情况表: {file_path}")

        try:
            if not file_path or not file_path.strip():
                return json.dumps({
                    "success": False,
                    "error": "file_path不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id
                }, ensure_ascii=False)

            result = IncidentEventReader.read_events(file_path)
            result["call_id"] = call_id
            result["timestamp"] = datetime.now().isoformat()

            if result.get("success"):
                logger.info(
                    f"[{call_id}] 读取完成: {result.get('total_events', 0)}条事件"
                )
            else:
                logger.warning(
                    f"[{call_id}] 读取失败: {result.get('error', '未知错误')}"
                )

            return json.dumps(result, ensure_ascii=False)

        except Exception as e:
            logger.exception(f"[{call_id}] 读取事件异常")
            return json.dumps({
                "success": False,
                "error": f"读取事件失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)

    # ============================================================
    # 第3步：生成复盘报告Word文档
    # ============================================================

    @mcp.tool()
    def generate_incident_report(
        analysis_result: str,
        output_path: str = "",
        incident_file_name: str = ""
    ) -> str:
        """
        【第3步】将LLM的复盘分析结果输出为Word文档

        接收LLM节点输出的JSON格式分析结果，按结构化模板生成排版美观的
        Word复盘报告，包含事件总结、原因分析、问题暴露、改进建议、应急处置等章节。

        输出文件默认保存在当前目录下的 reports/ 文件夹。

        Args:
            analysis_result: LLM节点输出的复盘分析结果JSON字符串
            output_path: 输出Word文件路径（可选，为空则自动生成到reports/目录）
            incident_file_name: 原始事件文件名（可选，用于报告标题命名）

        Returns:
            JSON字符串，包含 output_path（文件路径）、file_size（文件大小）、
            report_sections（报告包含的章节列表）等
        """
        start_time = datetime.now()
        call_id = f"gen_report_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始生成复盘报告...")

        try:
            if not analysis_result or not analysis_result.strip():
                return json.dumps({
                    "success": False,
                    "error": "analysis_result不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id
                }, ensure_ascii=False)

            # 验证analysis_result是有效JSON
            try:
                json.loads(analysis_result)
            except json.JSONDecodeError as e:
                return json.dumps({
                    "success": False,
                    "error": f"analysis_result不是有效JSON: {str(e)}",
                    "error_type": "JSONDecodeError",
                    "call_id": call_id,
                    "timestamp": datetime.now().isoformat()
                }, ensure_ascii=False)

            # 如果没有传output_path，留空让生成器自动生成
            actual_output = output_path.strip() if output_path else None

            result = IncidentReportGenerator.generate_report(
                analysis_result=analysis_result,
                output_path=actual_output,
                incident_file_name=incident_file_name
            )

            if result.get("success"):
                logger.info(f"[{call_id}] 报告生成完成: {result.get('output_path', '')}")
            else:
                logger.warning(f"[{call_id}] 报告生成失败: {result.get('error', '')}")

            return json.dumps(result, ensure_ascii=False)

        except Exception as e:
            logger.exception(f"[{call_id}] 生成报告异常")
            return json.dumps({
                "success": False,
                "error": f"生成报告失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)

    # ============================================================
    # 辅助工具：获取报告模板
    # ============================================================

    @mcp.tool()
    def get_incident_report_template() -> str:
        """
        获取复盘分析结果的JSON模板

        返回LLM节点应输出的JSON结构模板，帮助Dify工作流中的LLM节点
        生成符合预期的结构化数据。

        Returns:
            JSON字符串，包含报告模板的结构说明和各字段示例
        """
        template = {
            "summary": {
                "overview": "一句话事件概述（如：XX系统XX模块发生XX故障导致XX影响）",
                "impact_duration": "影响时长描述（如：共计XX分钟，影响XX用户）",
                "severity": "事件等级（如：P0重大/P1严重/P2一般）",
                "involved_systems": "涉及系统列表",
                "impact_scope": "影响范围详细描述"
            },
            "cause_analysis": {
                "direct_cause": "直接原因描述",
                "indirect_cause": "间接原因描述（如存在）",
                "root_cause": "根本原因分析"
            },
            "issues": {
                "monitoring_alerts": "监控告警方面暴露的问题",
                "architecture_design": "系统架构设计方面暴露的问题",
                "emergency_response": "应急处置和应急预案方面暴露的问题",
                "others": "其他方面暴露的问题"
            },
            "suggestions": [
                {
                    "category": "建议类别（如：监控告警）",
                    "content": "具体改进建议内容",
                    "priority": "优先级（高/中/低）"
                }
            ],
            "response_info": {
                "response_method": "应急处置方式详细描述",
                "response_duration": "处置时长",
                "effectiveness": "处置效果评估",
                "timeline": [
                    {"time": "时间点", "action": "处置动作"}
                ]
            },
            "original_events": []
        }

        return json.dumps({
            "success": True,
            "template": template,
            "usage": (
                "在Dify工作流中，LLM节点应严格按此模板输出JSON。"
                "字段名不可修改，缺失的字段可留空字符串或空列表。"
                "suggestions数组中可以包含多条建议。"
                "timeline数组可包含多个处置步骤。"
                "original_events数组用于回传原始事件数据（可选）。"
            ),
            "call_id": f"template_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            "timestamp": datetime.now().isoformat()
        }, ensure_ascii=False)


# ============================================================
# 独立启动入口（仅用于测试）
# ============================================================

if __name__ == "__main__":
    """独立运行此文件（仅作测试用，生产环境请使用 main.py）"""
    if not FASTMCP_AVAILABLE:
        print("错误: fastmcp未安装，请运行: pip install fastmcp")
        sys.exit(1)

    if not DOCX_AVAILABLE:
        print("警告: python-docx未安装，报告生成功能不可用")
        print("请运行: pip install python-docx")

    if not OPENPYXL_AVAILABLE:
        print("警告: openpyxl未安装，事件读取功能不可用")
        print("请运行: pip install openpyxl")

    print("=" * 60)
    print("事件复盘分析MCP服务器（独立模式，仅用于测试）")
    print("=" * 60)
    print("可用工具:")
    print("  1. read_incident_events      - 第1步：读取事件情况表")
    print("  2. generate_incident_report  - 第3步：生成复盘报告")
    print("  3. get_incident_report_template - 获取报告模板")
    print("\n注意：生产环境请使用 python main.py 统一启动")

    # 默认以SSE模式启动
    port = 8003 if len(sys.argv) < 2 else int(sys.argv[1])
    print(f"\n启动服务 (端口: {port})...")
    print("-" * 60)
    mcp.run(transport='sse', host='0.0.0.0', port=port)
