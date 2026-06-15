"""
文档修订MCP节点 - 第6步
使用 FastMCP 框架实现
功能：接收检查建议JSON和原文档，输出修订后的新文档（使用修订模式标记）
支持格式：Word文档(.docx)、WPS文档(.wps和.wpsx)
"""

import json
import os
import sys
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass
from datetime import datetime
from enum import Enum

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger('RevisionMCPNode')

# FastMCP导入
try:
    from fastmcp import FastMCP
    FASTMCP_AVAILABLE = True
except ImportError:
    FASTMCP_AVAILABLE = False
    logger.warning("fastmcp未安装，请运行: pip install fastmcp")

# 文档处理库
try:
    from docx import Document as DocxDocument
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx未安装，请运行: pip install python-docx")


class Severity(Enum):
    """建议严重程度"""
    CRITICAL = "Critical"
    HIGH = "High"
    MEDIUM = "Medium"
    LOW = "Low"


class IssueType(Enum):
    """问题类型"""
    FORMAT = "format"
    CONTENT = "content"
    LANGUAGE = "language"
    LOGIC = "logic"


@dataclass
class CheckSuggestion:
    """检查建议数据模型"""
    id: str
    rule_id: str
    rule_name: str
    type: str
    severity: str
    section: str
    original_text: str
    suggestion: str
    reason: str

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'CheckSuggestion':
        """从字典创建CheckSuggestion"""
        return cls(
            id=data.get('id', ''),
            rule_id=data.get('rule_id', ''),
            rule_name=data.get('rule_name', ''),
            type=data.get('type', ''),
            severity=data.get('severity', 'Medium'),
            section=data.get('section', ''),
            original_text=data.get('original_text', ''),
            suggestion=data.get('suggestion', ''),
            reason=data.get('reason', '')
        )


class DocumentParser:
    """文档解析器 - 支持docx/wps/wpsx"""

    @staticmethod
    def parse_document(file_path: str) -> Dict[str, Any]:
        """解析文档，返回结构和文本内容"""
        logger.info(f"开始解析文档: {file_path}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = Path(file_path).suffix.lower()

        if ext == '.docx':
            return DocumentParser._parse_docx(file_path)
        elif ext in ['.wps', '.wpsx']:
            return DocumentParser._parse_wps(file_path)
        else:
            raise ValueError(f"不支持的文档格式: {ext}")

    @staticmethod
    def _parse_docx(file_path: str) -> Dict[str, Any]:
        """解析Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装，请运行: pip install python-docx")

        logger.info("解析Word文档 (.docx)")

        doc = DocxDocument(file_path)
        paragraphs = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({
                    'index': i,
                    'text': text,
                    'style': para.style.name if para.style else 'Normal',
                    'element': para._element
                })

        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                'index': i,
                'data': table_data
            })

        logger.info(f"解析完成: {len(paragraphs)}个段落, {len(tables)}个表格")

        return {
            'type': 'docx',
            'paragraphs': paragraphs,
            'tables': tables,
            'document': doc,
            'full_text': '\n'.join([p['text'] for p in paragraphs])
        }

    @staticmethod
    def _parse_wps(file_path: str) -> Dict[str, Any]:
        """解析WPS文档 (.wps 和 .wpsx) - 使用OOXML格式，与DOCX兼容"""
        import zipfile

        logger.info("解析WPS文档 (.wps/.wpsx)")

        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装，请运行: pip install python-docx")

        try:
            # WPS/WPSX现在使用OOXML格式，可以直接用python-docx打开
            doc = DocxDocument(file_path)

            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs.append({
                        'index': i,
                        'text': text,
                        'style': para.style.name if para.style else 'Normal',
                        'element': para._element
                    })

            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    'index': i,
                    'data': table_data
                })

            logger.info(f"解析完成: {len(paragraphs)}个段落, {len(tables)}个表格")

            return {
                'type': 'wps',
                'paragraphs': paragraphs,
                'tables': tables,
                'document': doc,
                'full_text': '\n'.join([p['text'] for p in paragraphs]),
                'zipfile': file_path
            }
        except Exception as e:
            logger.error(f"WPS解析失败: {e}")
            raise


class RevisionMCPNode:
    """
    MCP节点 - 文档修订核心类
    提供统一的API接口，用于文档解析和修订
    """
    
    def __init__(self):
        self.name = "DocumentRevisionNode"
        self.version = "1.1.0"
        self.description = "文档检查建议修订节点"
    
    def revise_document(
        self,
        file_path: str,
        suggestions_json: str,
        output_path: str = "",
        use_track_changes: str = "true",
        preserve_original: bool = True
    ) -> str:
        """
        根据检查建议修订文档
        
        Args:
            file_path: 原文档路径
            suggestions_json: 检查建议JSON字符串
            output_path: 输出路径，可选
            use_track_changes: 是否使用修订模式（"true"/"false"）
            
        Returns:
            JSON字符串，包含修订结果
        """
        import json
        from datetime import datetime
        
        start_time = datetime.now()
        call_id = f"revise_{start_time.strftime('%Y%m%d_%H%M%S')}"
        
        logger.info(f"[{call_id}] 开始修订文档")
        
        try:
            # 1. 验证参数
            if not file_path or not file_path.strip():
                return {
                    "success": False,
                    "error": "file_path不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id
                }
            
            if not suggestions_json or not suggestions_json.strip():
                return {
                    "success": False,
                    "error": "suggestions_json不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id
                }
            
            file_path = file_path.strip()
            suggestions_json = suggestions_json.strip()
            use_track_changes_flag = use_track_changes.strip().lower() == "true"
            
            # 2. 验证文件存在
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"文件不存在: {file_path}",
                    "error_type": "FileNotFoundError",
                    "call_id": call_id
                }
            
            # 3. 解析建议JSON（适配第4步输出格式）
            try:
                suggestions_data = json.loads(suggestions_json)
                
                # 适配第4步输出格式：{"suggestions": [...], "statistics": {...}}
                if isinstance(suggestions_data, dict):
                    if "suggestions" in suggestions_data:
                        suggestions_list = suggestions_data["suggestions"]
                    else:
                        raise ValueError("JSON对象必须包含'suggestions'字段")
                elif isinstance(suggestions_data, list):
                    suggestions_list = suggestions_data
                else:
                    raise ValueError("格式不正确")
                
                # 过滤掉用户拒绝的建议
                suggestions_list = [
                    s for s in suggestions_list 
                    if s.get("accepted", True) and s.get("status") != "rejected"
                ]
                
            except Exception as e:
                return {
                    "success": False,
                    "error": f"建议数据格式错误: {str(e)}",
                    "error_type": "DataFormatError",
                    "call_id": call_id
                }
            
            # 4. 转换为CheckSuggestion对象
            suggestions = [CheckSuggestion.from_dict(s) for s in suggestions_list]
            logger.info(f"[{call_id}] 解析成功，共{len(suggestions)}条建议")
            
            # 5. 解析文档
            try:
                doc_data = DocumentParser.parse_document(file_path)
            except Exception as e:
                return {
                    "success": False,
                    "error": f"文档解析失败: {str(e)}",
                    "error_type": "ParseError",
                    "call_id": call_id
                }
            
            # 6. 应用修订
            try:
                if use_track_changes_flag and doc_data['type'] == 'docx':
                    reviser = TrackChangesReviser(doc_data)
                    result = reviser.apply_suggestions(suggestions)
                else:
                    # 对于WPS或非修订模式，使用简单替换
                    result = self._apply_simple_revisions(doc_data, suggestions)
            except Exception as e:
                return {
                    "success": False,
                    "error": f"应用修订失败: {str(e)}",
                    "error_type": "RevisionError",
                    "call_id": call_id
                }
            
            # 7. 生成输出路径
            if not output_path:
                from pathlib import Path
                file_path_obj = Path(file_path)
                output_path = str(file_path_obj.parent / f"{file_path_obj.stem}_revised{file_path_obj.suffix}")
            
            # 8. 保存文档
            try:
                if doc_data['type'] == 'docx':
                    result['document'].save(output_path)
                elif doc_data['type'] == 'wps':
                    doc_data['document'].save(output_path)
            except Exception as e:
                return {
                    "success": False,
                    "error": f"保存文档失败: {str(e)}",
                    "error_type": "SaveError",
                    "call_id": call_id
                }
            
            # 9. 构建成功结果
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            
            successful_revisions = [r for r in result['revisions'] if r['status'] == 'success']
            
            # 统计
            stats = {"by_severity": {}, "by_type": {}}
            for r in successful_revisions:
                sev = r.get('severity', 'Medium')
                typ = r.get('type', 'content')
                stats["by_severity"][sev] = stats["by_severity"].get(sev, 0) + 1
                stats["by_type"][typ] = stats["by_type"].get(typ, 0) + 1
            
            return {
                "success": True,
                "output_path": output_path,
                "total_suggestions": len(suggestions),
                "applied_revisions": result['revision_count'],
                "revisions_detail": successful_revisions,
                "skipped_revisions": len([r for r in result['revisions'] if r['status'] != 'success']),
                "statistics": stats,
                "call_id": call_id,
                "timestamp": end_time.isoformat(),
                "duration_seconds": duration
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"未知错误: {str(e)}",
                "error_type": "UnexpectedError",
                "call_id": call_id
            }
    
    def _apply_simple_revisions(self, doc_data: Dict, suggestions: List[CheckSuggestion]) -> Dict:
        """简单修订（非修订模式）"""
        doc = doc_data['document']
        revision_count = 0
        revisions = []
        
        for suggestion in suggestions:
            original_text = suggestion.original_text.strip()
            new_text = suggestion.suggestion.strip()
            
            if not original_text or not new_text:
                continue
            
            found = False
            for para in doc.paragraphs:
                if original_text in para.text:
                    found = True
                    para.text = para.text.replace(original_text, new_text)
                    
                    revisions.append({
                        'id': suggestion.id,
                        'rule_id': suggestion.rule_id,
                        'rule_name': suggestion.rule_name,
                        'type': suggestion.type,
                        'severity': suggestion.severity,
                        'original': original_text[:50],
                        'new': new_text[:50],
                        'status': 'success'
                    })
                    
                    revision_count += 1
                    break
            
            if not found:
                revisions.append({
                    'id': suggestion.id,
                    'rule_id': suggestion.rule_id,
                    'rule_name': suggestion.rule_name,
                    'type': suggestion.type,
                    'severity': suggestion.severity,
                    'original': original_text[:50],
                    'new': new_text[:50],
                    'status': 'not_found'
                })
        
        return {
            'document': doc,
            'revision_count': revision_count,
            'revisions': revisions,
            'total_suggestions': len(suggestions)
        }
    
    def get_tool_definition(self) -> Dict[str, Any]:
        """
        返回MCP工具定义
        
        Returns:
            工具定义字典
        """
        return {
            "name": self.name,
            "description": self.description,
            "version": self.version,
            "tools": [
                {
                    "name": "revise_document",
                    "description": "根据检查建议修订文档（使用修订模式标记修改）",
                    "input_schema": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "原文档路径（支持.docx, .wps, .wpsx）"
                            },
                            "suggestions_json": {
                                "type": "string",
                                "description": "检查建议JSON字符串"
                            },
                            "output_path": {
                                "type": "string",
                                "description": "输出路径，可选"
                            },
                            "use_track_changes": {
                                "type": "string",
                                "description": "是否使用修订模式（true/false），默认true"
                            }
                        },
                        "required": ["file_path", "suggestions_json"]
                    }
                },
                {
                    "name": "parse_document",
                    "description": "解析文档并返回结构化内容",
                    "input_schema": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "文档路径"
                            }
                        },
                        "required": ["file_path"]
                    }
                }
            ]
        }
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        解析文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            字典，包含文档结构
        """
        from datetime import datetime
        
        try:
            doc_data = DocumentParser.parse_document(file_path)
            
            return {
                "success": True,
                "file_path": file_path,
                "type": doc_data['type'],
                "paragraph_count": len(doc_data['paragraphs']),
                "table_count": len(doc_data['tables']),
                "paragraphs": [p['text'] for p in doc_data['paragraphs']],
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"解析失败: {str(e)}"
            }


class TrackChangesReviser:
    """修订模式文档修订器 - 使用Word修订追踪功能"""

    def __init__(self, doc_data: Dict[str, Any]):
        self.doc_data = doc_data
        self.revisions: List[Dict[str, Any]] = []
        self.revision_id = 0

    def apply_suggestions(self, suggestions: List[CheckSuggestion]) -> Dict[str, Any]:
        """应用修订建议到文档，使用修订模式标记"""
        logger.info(f"开始应用修订建议，共{len(suggestions)}条（使用修订模式）")

        if self.doc_data['type'] == 'docx':
            return self._apply_to_docx_track_changes(suggestions)
        elif self.doc_data['type'] == 'wps':
            return self._apply_to_wps(suggestions)
        else:
            raise ValueError(f"不支持的文档类型: {self.doc_data['type']}")

    def _get_next_revision_id(self) -> str:
        """获取下一个修订ID"""
        self.revision_id += 1
        return str(self.revision_id)

    def _create_deleted_text(self, text: str, revision_id: str) -> OxmlElement:
        """创建删除标记 w:del"""
        # 创建 w:del 元素
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), revision_id)
        del_elem.set(qn('w:author'), 'DocumentReviewer')
        del_elem.set(qn('w:date'), datetime.now().isoformat())

        # 创建 w:r 元素
        run_elem = OxmlElement('w:r')

        # 创建 w:rPr 元素（删除线样式）
        rpr_elem = OxmlElement('w:rPr')
        strike_elem = OxmlElement('w:strike')
        rpr_elem.append(strike_elem)
        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), 'FF0000')  # 红色
        rpr_elem.append(color_elem)
        run_elem.append(rpr_elem)

        # 创建 w:t 元素
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        run_elem.append(text_elem)

        del_elem.append(run_elem)
        return del_elem

    def _create_inserted_text(self, text: str, revision_id: str) -> OxmlElement:
        """创建插入标记 w:ins"""
        # 创建 w:ins 元素
        ins_elem = OxmlElement('w:ins')
        ins_elem.set(qn('w:id'), revision_id)
        ins_elem.set(qn('w:author'), 'DocumentReviewer')
        ins_elem.set(qn('w:date'), datetime.now().isoformat())

        # 创建 w:r 元素
        run_elem = OxmlElement('w:r')

        # 创建 w:rPr 元素（下划线样式）
        rpr_elem = OxmlElement('w:rPr')
        underline_elem = OxmlElement('w:u')
        underline_elem.set(qn('w:val'), 'single')
        rpr_elem.append(underline_elem)
        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), '0000FF')  # 蓝色
        rpr_elem.append(color_elem)
        run_elem.append(rpr_elem)

        # 创建 w:t 元素
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        run_elem.append(text_elem)

        ins_elem.append(run_elem)
        return ins_elem

    def _apply_to_docx_track_changes(self, suggestions: List[CheckSuggestion]) -> Dict[str, Any]:
        """使用Word修订追踪模式修订文档"""
        doc = self.doc_data['document']
        revision_count = 0

        for suggestion in suggestions:
            original_text = suggestion.original_text.strip()
            new_text = suggestion.suggestion.strip()

            if not original_text or not new_text:
                logger.warning(f"跳过空文本建议: {suggestion.id}")
                continue

            found = False
            for para in doc.paragraphs:
                if original_text in para.text:
                    found = True
                    revision_id = self._get_next_revision_id()

                    self.revisions.append({
                        'id': suggestion.id,
                        'rule_id': suggestion.rule_id,
                        'rule_name': suggestion.rule_name,
                        'type': suggestion.type,
                        'severity': suggestion.severity,
                        'original': original_text[:50],
                        'new': new_text[:50],
                        'status': 'success'
                    })

                    # 获取段落XML元素
                    p_elem = para._element

                    if original_text == para.text:
                        # 整段替换
                        # 清空段落内容
                        for child in list(p_elem):
                            if child.tag.endswith('r'):  # 只删除run元素
                                p_elem.remove(child)

                        # 添加删除标记
                        del_elem = self._create_deleted_text(original_text, revision_id)
                        p_elem.append(del_elem)

                        # 添加插入标记
                        ins_elem = self._create_inserted_text(new_text, revision_id)
                        p_elem.append(ins_elem)
                    else:
                        # 部分文本替换 - 在run级别操作
                        for run in para.runs:
                            if original_text in run.text:
                                run_text = run.text
                                parts = run_text.split(original_text, 1)

                                # 清空当前run
                                run._r.clear()

                                # 添加前半部分（如果有）
                                if parts[0]:
                                    t_elem = OxmlElement('w:t')
                                    t_elem.text = parts[0]
                                    run._r.append(t_elem)

                                # 添加删除标记
                                del_elem = self._create_deleted_text(original_text, revision_id)
                                run._r.append(del_elem)

                                # 添加插入标记
                                ins_elem = self._create_inserted_text(new_text, revision_id)
                                run._r.append(ins_elem)

                                # 添加后半部分（如果有）
                                if len(parts) > 1 and parts[1]:
                                    t_elem = OxmlElement('w:t')
                                    t_elem.text = parts[1]
                                    run._r.append(t_elem)

                                break  # 只处理第一个匹配的run

                    revision_count += 1
                    logger.info(f"  [REVISED] [{suggestion.severity}] {suggestion.rule_name}")
                    break

            if not found:
                self.revisions.append({
                    'id': suggestion.id,
                    'rule_id': suggestion.rule_id,
                    'rule_name': suggestion.rule_name,
                    'type': suggestion.type,
                    'severity': suggestion.severity,
                    'original': original_text[:50],
                    'new': new_text[:50],
                    'status': 'not_found'
                })
                logger.warning(f"  [NOT FOUND] 未找到原文: {original_text[:30]}...")

        return {
            'document': doc,
            'revision_count': revision_count,
            'revisions': self.revisions,
            'total_suggestions': len(suggestions)
        }

    def _apply_to_wps(self, suggestions: List[CheckSuggestion]) -> Dict[str, Any]:
        """修订WPS文档 - 使用与DOCX相同的修订模式"""
        # WPS现在使用OOXML格式，可以直接使用docx的修订逻辑
        doc = self.doc_data['document']
        revision_count = 0

        for suggestion in suggestions:
            original_text = suggestion.original_text.strip()
            new_text = suggestion.suggestion.strip()

            if not original_text or not new_text:
                logger.warning(f"跳过空文本建议: {suggestion.id}")
                continue

            found = False
            for para in doc.paragraphs:
                if original_text in para.text:
                    found = True
                    revision_id = self._get_next_revision_id()

                    self.revisions.append({
                        'id': suggestion.id,
                        'rule_id': suggestion.rule_id,
                        'rule_name': suggestion.rule_name,
                        'type': suggestion.type,
                        'severity': suggestion.severity,
                        'original': original_text[:50],
                        'new': new_text[:50],
                        'status': 'success'
                    })

                    # 获取段落XML元素
                    p_elem = para._element

                    if original_text == para.text:
                        # 整段替换
                        # 清空段落内容
                        for child in list(p_elem):
                            if child.tag.endswith('r'):  # 只删除run元素
                                p_elem.remove(child)

                        # 添加删除标记
                        del_elem = self._create_deleted_text(original_text, revision_id)
                        p_elem.append(del_elem)

                        # 添加插入标记
                        ins_elem = self._create_inserted_text(new_text, revision_id)
                        p_elem.append(ins_elem)
                    else:
                        # 部分文本替换 - 在run级别操作
                        for run in para.runs:
                            if original_text in run.text:
                                run_text = run.text
                                parts = run_text.split(original_text, 1)

                                # 清空当前run
                                run._r.clear()

                                # 添加前半部分（如果有）
                                if parts[0]:
                                    t_elem = OxmlElement('w:t')
                                    t_elem.text = parts[0]
                                    run._r.append(t_elem)

                                # 添加删除标记
                                del_elem = self._create_deleted_text(original_text, revision_id)
                                run._r.append(del_elem)

                                # 添加插入标记
                                ins_elem = self._create_inserted_text(new_text, revision_id)
                                run._r.append(ins_elem)

                                # 添加后半部分（如果有）
                                if len(parts) > 1 and parts[1]:
                                    t_elem = OxmlElement('w:t')
                                    t_elem.text = parts[1]
                                    run._r.append(t_elem)

                                break  # 只处理第一个匹配的run

                    revision_count += 1
                    logger.info(f"  [REVISED] [{suggestion.severity}] {suggestion.rule_name}")
                    break

            if not found:
                self.revisions.append({
                    'id': suggestion.id,
                    'rule_id': suggestion.rule_id,
                    'rule_name': suggestion.rule_name,
                    'type': suggestion.type,
                    'severity': suggestion.severity,
                    'original': original_text[:50],
                    'new': new_text[:50],
                    'status': 'not_found'
                })
                logger.warning(f"  [NOT FOUND] 未找到原文: {original_text[:30]}...")

        return {
            'document': doc,
            'revision_count': revision_count,
            'revisions': self.revisions,
            'total_suggestions': len(suggestions)
        }


# ============================================================
# FastMCP 服务器
# ============================================================

if FASTMCP_AVAILABLE:
    mcp = FastMCP("DocumentRevisionNode")

    @mcp.tool()
    def revise_document(
        file_path: str,
        suggestions_json: str,
        output_path: str = "",
        use_track_changes: str = "true"
    ) -> str:
        """
    根据检查建议修订文档，支持修订模式标记修改
    
    按流程依次完成参数验证、文件存在性检查、建议JSON解析（兼容第4步输出格式及纯数组格式）、
    文档解析、修订应用（docx格式支持修订模式，其他格式使用传统替换方式）及文档保存。
    自动过滤用户拒绝的建议，并在输出路径未指定时默认在原文件名后添加_revised后缀。
    
    Args:
        file_path (str): 原文档路径，支持 .docx、.wps、.wpsx 格式
        suggestions_json (str): 检查建议JSON字符串，支持包含 'suggestions' 或 'data' 字段的对象格式及纯数组格式
        output_path (str): 输出文件路径，可选，默认在原文件名后加 _revised 后缀
        use_track_changes (str): 是否使用修订模式，'true' 启用，其他值为禁用，默认 'true'
    
    Returns:
        str: JSON字符串，包含以下字段：
            - success (bool): 操作是否成功
            - output_path (str): 修订后文件的保存路径（成功时）
            - applied_revisions (int): 成功应用的修订数量（成功时）
            - total_suggestions (int): 总建议数量（成功时）
            - revisions_detail (list): 各条修订的详细结果（成功时）
            - statistics (dict): 按严重程度和类型统计的修订数据（成功时）
            - error (str): 错误描述（失败时）
            - error_type (str): 错误类型，如 ValidationError、FileNotFoundError、JSONParseError、ParseError、RevisionError、SaveError（失败时）
            - call_id (str): 调用唯一标识
            - timestamp (str): 时间戳
    
    Raises:
        无显式抛出异常，所有错误均通过返回JSON中的 error_type 字段标识
        """
        start_time = datetime.now()
        call_id = f"revise_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始修订文档")
        logger.info(f"[{call_id}] 输入 - file_path: {file_path}")
        logger.info(f"[{call_id}] 输入 - suggestions_json长度: {len(suggestions_json)}")
        logger.info(f"[{call_id}] 输入 - use_track_changes: {use_track_changes}")

        try:
            # 1. 验证参数
            if not file_path or not file_path.strip():
                error_result = {
                    "success": False,
                    "error": "file_path不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            if not suggestions_json or not suggestions_json.strip():
                error_result = {
                    "success": False,
                    "error": "suggestions_json不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            file_path = file_path.strip()
            suggestions_json = suggestions_json.strip()
            use_track_changes_flag = use_track_changes.strip().lower() == "true"

            # 2. 验证文件存在
            if not os.path.exists(file_path):
                error_result = {
                    "success": False,
                    "error": f"文件不存在: {file_path}",
                    "error_type": "FileNotFoundError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            # 3. 解析建议JSON（适配第4步输出格式）
            try:
                suggestions_data = json.loads(suggestions_json)
            
                # 适配第4步输出格式：{"suggestions": [...], "statistics": {...}}
                if isinstance(suggestions_data, dict):
                    if "suggestions" in suggestions_data:
                        suggestions_list = suggestions_data["suggestions"]
                        logger.info(f"[{call_id}] 检测到第4步输出格式，提取suggestions数组")
                    elif "data" in suggestions_data:
                        suggestions_list = suggestions_data["data"]
                    else:
                        raise ValueError("JSON对象必须包含'suggestions'或'data'字段")
                elif isinstance(suggestions_data, list):
                    suggestions_list = suggestions_data
                else:
                    raise ValueError("suggestions_json必须是数组格式或包含suggestions的对象")
            
                # 过滤掉用户拒绝的建议（如果存在accepted字段）
                suggestions_list = [
                    s for s in suggestions_list 
                    if s.get("accepted", True) or s.get("status") != "rejected"
                ]
            
                logger.info(f"[{call_id}] 有效建议数量: {len(suggestions_list)}")
            
            except json.JSONDecodeError as e:
                error_result = {
                    "success": False,
                    "error": f"JSON解析错误: {str(e)}",
                    "error_type": "JSONParseError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)
            except Exception as e:
                error_result = {
                    "success": False,
                    "error": f"建议数据格式错误: {str(e)}",
                    "error_type": "DataFormatError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            # 转换为CheckSuggestion对象
            try:
                suggestions = [CheckSuggestion.from_dict(s) for s in suggestions_list]
                logger.info(f"[{call_id}] 解析成功，共{len(suggestions)}条建议")
            except Exception as e:
                error_result = {
                    "success": False,
                    "error": f"建议数据字段缺失: {str(e)}",
                    "error_type": "FieldMissingError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            # 4. 解析文档
            try:
                doc_data = DocumentParser.parse_document(file_path)
            except Exception as e:
                error_result = {
                    "success": False,
                    "error": f"文档解析失败: {str(e)}",
                    "error_type": "ParseError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            # 5. 应用修订（使用修订模式）
            try:
                if use_track_changes_flag and doc_data['type'] == 'docx':
                    reviser = TrackChangesReviser(doc_data)
                    result = reviser.apply_suggestions(suggestions)
                else:
                    # 不使用修订模式时，使用简单替换
                    result = self._apply_simple_revisions(doc_data, suggestions)
            except Exception as e:
                error_result = {
                    "success": False,
                    "error": f"应用修订失败: {str(e)}",
                    "error_type": "RevisionError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            # 6. 生成输出路径
            if not output_path or not output_path.strip():
                file_path_obj = Path(file_path)
                output_path = str(file_path_obj.parent / f"{file_path_obj.stem}_revised{file_path_obj.suffix}")
            else:
                output_path = output_path.strip()

            # 7. 保存文档
            try:
                if doc_data['type'] == 'docx':
                    result['document'].save(output_path)
                    logger.info(f"[{call_id}] DOCX文档已保存: {output_path}")
                elif doc_data['type'] == 'wps':
                    # WPS保存逻辑
                    output_path = _save_wps(doc_data, output_path, suggestions)
                    logger.info(f"[{call_id}] WPS文档已保存: {output_path}")
            except Exception as e:
                error_result = {
                    "success": False,
                    "error": f"保存文档失败: {str(e)}",
                    "error_type": "SaveError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            # 8. 构建成功结果
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()

            successful_revisions = [r for r in result['revisions'] if r['status'] == 'success']

            success_result = {
                "success": True,
                "output_path": output_path,
                "use_track_changes": use_track_changes_flag,
                "total_suggestions": len(suggestions),
                "applied_revisions": result['revision_count'],
                "revisions_detail": successful_revisions,
                "skipped_revisions": len([r for r in result['revisions'] if r['status'] != 'success']),
                "statistics": {
                    "by_severity": _count_by(successful_revisions, 'severity'),
                    "by_type": _count_by(successful_revisions, 'type')
                },
                "call_id": call_id,
                "timestamp": end_time.isoformat(),
                "duration_seconds": duration
            }

            logger.info(f"[{call_id}] 修订完成! 耗时: {duration:.2f}秒")
            logger.info(f"[{call_id}] 输出: {success_result['output_path']}")
            logger.info(f"[{call_id}] 应用: {success_result['applied_revisions']}/{success_result['total_suggestions']}")

            return json.dumps(success_result, ensure_ascii=False, indent=2)

        except Exception as e:
            error_result = {
                "success": False,
                "error": f"未知错误: {str(e)}",
                "error_type": "UnexpectedError",
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }
            logger.exception(f"[{call_id}] 未捕获异常")
            return json.dumps(error_result, ensure_ascii=False)

    @mcp.tool()
    def parse_document(file_path: str) -> str:
        """
        解析文档并返回结构化内容

        Args:
            file_path: 文档路径（支持.docx, .wps, .wpsx）

        Returns:
            JSON字符串，包含文档结构和文本内容
        """
        start_time = datetime.now()
        call_id = f"parse_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始解析文档: {file_path}")

        try:
            if not file_path or not file_path.strip():
                error_result = {
                    "success": False,
                    "error": "file_path不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id
                }
                return json.dumps(error_result, ensure_ascii=False)

            file_path = file_path.strip()

            if not os.path.exists(file_path):
                error_result = {
                    "success": False,
                    "error": f"文件不存在: {file_path}",
                    "error_type": "FileNotFoundError",
                    "call_id": call_id
                }
                return json.dumps(error_result, ensure_ascii=False)

            doc_data = DocumentParser.parse_document(file_path)

            result = {
                "success": True,
                "file_path": file_path,
                "type": doc_data['type'],
                "paragraph_count": len(doc_data['paragraphs']),
                "table_count": len(doc_data['tables']),
                "paragraphs": [p['text'] for p in doc_data['paragraphs']],
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }

            logger.info(f"[{call_id}] 解析成功: {result['paragraph_count']}个段落")
            return json.dumps(result, ensure_ascii=False, indent=2)

        except Exception as e:
            error_result = {
                "success": False,
                "error": f"解析失败: {str(e)}",
                "error_type": "ParseError",
                "call_id": call_id
            }
            logger.exception(f"[{call_id}] 解析异常")
            return json.dumps(error_result, ensure_ascii=False)

    @mcp.tool()
    def validate_suggestions(suggestions_json: str) -> str:
        """
        验证检查建议JSON格式

        Args:
            suggestions_json: 检查建议JSON字符串

        Returns:
            JSON字符串，包含验证结果
        """
        start_time = datetime.now()
        call_id = f"validate_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始验证建议JSON")

        try:
            if not suggestions_json or not suggestions_json.strip():
                return json.dumps({
                    "success": False,
                    "error": "suggestions_json不能为空",
                    "error_type": "ValidationError"
                }, ensure_ascii=False)

            suggestions_data = json.loads(suggestions_json)

            if not isinstance(suggestions_data, list):
                return json.dumps({
                    "success": False,
                    "error": "suggestions_json必须是数组格式",
                    "error_type": "ValidationError"
                }, ensure_ascii=False)

            # 验证每条建议的必需字段
            required_fields = ['id', 'rule_id', 'rule_name', 'type', 'severity', 'original_text', 'suggestion']
            valid_suggestions = []
            invalid_suggestions = []

            for i, s in enumerate(suggestions_data):
                missing_fields = [f for f in required_fields if f not in s]
                if missing_fields:
                    invalid_suggestions.append({
                        "index": i,
                        "id": s.get('id', 'unknown'),
                        "missing_fields": missing_fields
                    })
                else:
                    valid_suggestions.append(CheckSuggestion.from_dict(s))

            result = {
                "success": True,
                "total": len(suggestions_data),
                "valid_count": len(valid_suggestions),
                "invalid_count": len(invalid_suggestions),
                "invalid_details": invalid_suggestions,
                "call_id": call_id
            }

            if invalid_suggestions:
                result["success"] = False
                result["error"] = f"有{len(invalid_suggestions)}条建议格式不正确"

            logger.info(f"[{call_id}] 验证完成: {result['valid_count']}条有效, {result['invalid_count']}条无效")
            return json.dumps(result, ensure_ascii=False, indent=2)

        except json.JSONDecodeError as e:
            error_result = {
                "success": False,
                "error": f"JSON解析错误: {str(e)}",
                "error_type": "JSONParseError"
            }
            logger.error(f"[{call_id}] JSON解析失败")
            return json.dumps(error_result, ensure_ascii=False)

        except Exception as e:
            error_result = {
                "success": False,
                "error": f"验证失败: {str(e)}",
                "error_type": "UnexpectedError"
            }
            logger.exception(f"[{call_id}] 验证异常")
            return json.dumps(error_result, ensure_ascii=False)

    @mcp.tool()
    def get_tools_info() -> str:
        """
        获取工具定义信息

        Returns:
            JSON字符串，包含所有工具的定义
        """
        logger.info("[tools_info] 获取工具定义")

        tools_info = {
            "name": "DocumentRevisionNode",
            "description": "文档检查建议修订节点 - 根据LLM生成的检查建议对文档进行修订（使用修订模式标记）",
            "version": "1.1.0",
            "supported_formats": [".docx", ".wps", ".wpsx"],
            "tools": [
                {
                    "name": "revise_document",
                    "description": "根据检查建议修订文档（使用修订模式标记修改）",
                    "parameters": {
                        "file_path": "string - 原文档路径",
                        "suggestions_json": "string - 检查建议JSON",
                        "output_path": "string - 输出路径(可选)",
                        "use_track_changes": "string - 是否使用修订模式(true/false，默认true)"
                    },
                    "returns": "JSON - 修订结果"
                },
                {
                    "name": "parse_document",
                    "description": "解析文档并返回结构化内容",
                    "parameters": {
                        "file_path": "string - 文档路径"
                    },
                    "returns": "JSON - 文档结构"
                },
                {
                    "name": "validate_suggestions",
                    "description": "验证检查建议JSON格式",
                    "parameters": {
                        "suggestions_json": "string - 检查建议JSON"
                    },
                    "returns": "JSON - 验证结果"
                }
            ]
        }

        return json.dumps(tools_info, ensure_ascii=False, indent=2)


def _count_by(items: List[Dict], key: str) -> Dict[str, int]:
    """统计列表中每个值出现的次数"""
    counts = {}
    for item in items:
        val = item.get(key, 'unknown')
        counts[val] = counts.get(val, 0) + 1
    return counts


def _save_wps(doc_data: Dict, output_path: str, suggestions: List[CheckSuggestion]) -> str:
    """保存WPS文档 - 创建完整的WPS兼容结构"""
    import zipfile

    doc_type = doc_data.get('type', 'wps')

    if doc_type == 'wps':
        return _save_wps_format(doc_data, output_path)
    else:
        return _save_wpsx_format(doc_data, output_path)


def _save_wps_format(doc_data: Dict, output_path: str) -> str:
    """保存.wps格式文档 - 使用docx库生成OOXML格式，WPS可兼容打开"""
    from docx import Document
    from docx.shared import Pt, Inches

    # 创建新文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 添加段落
    paragraphs = doc_data.get('paragraphs', [])
    for para_data in paragraphs:
        text = para_data.get('text', '')
        if text.strip():
            para = doc.add_paragraph(text)

    # 保存为docx格式，然后重命名为.wps
    temp_docx = output_path.replace('.wps', '_temp.docx')
    doc.save(temp_docx)

    # 将docx内容复制到.wps文件（WPS可以识别OOXML格式）
    import zipfile
    with zipfile.ZipFile(temp_docx, 'r') as zf_src:
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf_dst:
            for item in zf_src.namelist():
                zf_dst.writestr(item, zf_src.read(item))

    # 删除临时文件
    try:
        os.remove(temp_docx)
    except:
        pass

    return output_path


def _save_wpsx_format(doc_data: Dict, output_path: str) -> str:
    """保存.wpsx格式文档"""
    import zipfile

    paragraphs = doc_data.get('paragraphs', [])

    def escape_xml(text):
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    # 生成document.xml
    document_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>'''

    for para in paragraphs:
        text = escape_xml(para['text'])
        document_xml += f'\n    <w:p><w:r><w:t>{text}</w:t></w:r></w:p>'

    document_xml += '''
    <w:sectPr>
      <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="1440" w:right="1800" w:bottom="1440" w:left="1800"/>
    </w:sectPr>
  </w:body>
</w:document>'''

    # 样式定义
    styles_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:docDefaults>
    <w:rPrDefault>
      <w:rPr>
        <w:rFonts w:ascii="宋体" w:eastAsia="宋体" w:hAnsi="宋体"/>
        <w:sz w:val="24"/>
        <w:szCs w:val="24"/>
      </w:rPr>
    </w:rPrDefault>
  </w:docDefaults>
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:rPr>
      <w:rFonts w:ascii="宋体" w:eastAsia="宋体"/>
      <w:sz w:val="24"/>
    </w:rPr>
  </w:style>
</w:styles>'''

    # 内容类型定义
    content_types_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>'''

    # 根关系定义
    rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''

    # word目录下的关系定义
    word_rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>'''

    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('word/document.xml', document_xml)
        zf.writestr('word/styles.xml', styles_xml)
        zf.writestr('word/_rels/document.xml.rels', word_rels_xml)
        zf.writestr('[Content_Types].xml', content_types_xml)
        zf.writestr('_rels/.rels', rels_xml)

    return output_path


def _generate_wps_content(doc_data: Dict, suggestions: List[CheckSuggestion]) -> str:
    """生成WPS内容XML（兼容旧版本）"""
    paragraphs = doc_data.get('paragraphs', [])

    def escape_xml(text):
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    xml_parts = ['<?xml version="1.0" encoding="UTF-8"?>']
    xml_parts.append('<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">')
    xml_parts.append('<w:body>')

    for para in paragraphs:
        text = escape_xml(para['text'])
        xml_parts.append(f'<w:p><w:r><w:t>{text}</w:t></w:r></w:p>')

    xml_parts.append('</w:body></w:document>')
    return '\n'.join(xml_parts)


# ============================================================
# 数据转换工具函数（用于第4步到第6步的数据流转）
# ============================================================

def transform_step4_to_step6(step4_output_json: str, accepted_only: bool = True) -> str:
    """
    将第4步输出转换为第6步输入格式
    
    支持的处理：
    1. 提取suggestions数组
    2. 过滤用户拒绝的建议（如果accepted_only=True）
    3. 移除第6步不需要的字段
    4. 返回标准JSON字符串
    
    Args:
        step4_output_json: 第4步输出JSON字符串
        accepted_only: 是否只保留用户接受的建议（默认True）
        
    Returns:
        JSON字符串（数组格式），可直接传递给第6步
    """
    try:
        data = json.loads(step4_output_json)
        
        # 提取建议数组
        if isinstance(data, dict):
            suggestions = data.get("suggestions", [])
        elif isinstance(data, list):
            suggestions = data
        else:
            raise ValueError("输入格式不正确")
        
        # 过滤拒绝的建议
        if accepted_only:
            suggestions = [
                s for s in suggestions 
                if s.get("accepted", True) and s.get("status") != "rejected"
            ]
        
        # 只保留第6步需要的字段
        filtered = []
        for s in suggestions:
            filtered.append({
                "id": s.get("id", ""),
                "rule_id": s.get("rule_id", ""),
                "rule_name": s.get("rule_name", ""),
                "type": s.get("type", ""),
                "severity": s.get("severity", "Medium"),
                "section": s.get("section", ""),
                "original_text": s.get("original_text", ""),
                "suggestion": s.get("suggestion", ""),
                "reason": s.get("reason", "")
            })
        
        return json.dumps(filtered, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"数据转换失败: {e}")
        raise


def prepare_suggestions_for_revision(
    suggestions_json: str,
    accepted_ids: str = ""
) -> str:
    """
    准备建议数据用于文档修订
    
    此函数用于第5步（用户确认）后，准备数据给第6步
    
    Args:
        suggestions_json: 建议JSON字符串（可以是第4步输出或直接数组）
        accepted_ids: 用户接受的建议ID列表，逗号分隔（如"sugg_001,sugg_002"）
                     如果为空，则接受所有建议
                     
    Returns:
        JSON字符串（数组格式）
    """
    try:
        data = json.loads(suggestions_json)
        
        # 提取建议数组
        if isinstance(data, dict):
            suggestions = data.get("suggestions", [])
        elif isinstance(data, list):
            suggestions = data
        else:
            raise ValueError("输入必须是JSON数组或包含suggestions的对象")
        
        # 如果指定了接受的ID，则过滤
        if accepted_ids and accepted_ids.strip():
            accepted_set = set(id.strip() for id in accepted_ids.split(","))
            suggestions = [s for s in suggestions if s.get("id", "") in accepted_set]
        
        # 转换为标准格式
        result = []
        for s in suggestions:
            result.append({
                "id": s.get("id", ""),
                "rule_id": s.get("rule_id", ""),
                "rule_name": s.get("rule_name", ""),
                "type": s.get("type", ""),
                "severity": s.get("severity", "Medium"),
                "section": s.get("section", ""),
                "original_text": s.get("original_text", ""),
                "suggestion": s.get("suggestion", ""),
                "reason": s.get("reason", "")
            })
        
        logger.info(f"准备了{len(result)}条建议用于修订")
        return json.dumps(result, ensure_ascii=False)
        
    except Exception as e:
        logger.error(f"准备建议数据失败: {e}")
        raise


# ============================================================
# 主入口
# ============================================================

def main():
    """主入口函数 - 支持stdio和SSE两种传输方式"""
    import argparse

    parser = argparse.ArgumentParser(description='文档修订MCP节点服务')
    parser.add_argument(
        '--transport',
        choices=['stdio', 'sse'],
        default='stdio',
        help='传输方式: stdio (默认) 或 sse'
    )
    parser.add_argument(
        '--host',
        default='0.0.0.0',
        help='SSE模式下的主机地址 (默认: 0.0.0.0)'
    )
    parser.add_argument(
        '--port',
        type=int,
        default=8000,
        help='SSE模式下的端口号 (默认: 8000)'
    )

    args = parser.parse_args()

    print("=" * 60)
    print("文档修订MCP节点服务（修订模式版）")
    print("=" * 60)

    if not FASTMCP_AVAILABLE:
        print("\n错误: fastmcp未安装")
        print("请运行: pip install fastmcp")
        sys.exit(1)

    print(f"\n传输方式: {args.transport.upper()}")

    if args.transport == 'sse':
        print(f"服务地址: http://{args.host}:{args.port}")
        print(f"SSE端点: http://{args.host}:{args.port}/sse")
        print(f"消息端点: http://{args.host}:{args.port}/messages")

    print("\n可用工具:")
    print("  1. revise_document   - 修订文档（使用修订模式）")
    print("  2. parse_document     - 解析文档")
    print("  3. validate_suggestions - 验证建议")
    print("  4. get_tools_info     - 工具信息")
    print("\n启动服务...")
    print("-" * 60)

    # 启动FastMCP服务
    if args.transport == 'sse':
        mcp.run(transport='sse', host=args.host, port=args.port)
    else:
        mcp.run()


if __name__ == "__main__":
    main()
