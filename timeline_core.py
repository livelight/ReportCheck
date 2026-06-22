"""
事件分析报告时间轴图生成 MCP 服务 - 核心功能模块

此模块包含所有业务逻辑的实现函数，可以被测试脚本直接调用，
也可以被 MCP 服务器包装为工具。
"""

import json
import logging
import hashlib
import shutil
from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),  # 控制台输出
        logging.FileHandler('mcp_server.log', encoding='utf-8')  # 文件日志
    ]
)
logger = logging.getLogger("TimelineMCP")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import matplotlib
    matplotlib.use('Agg')  # 非交互式后端
    import matplotlib.pyplot as plt
    import matplotlib.dates as mdates
    from matplotlib.font_manager import FontProperties
    import numpy as np
except ImportError as e:
    logger.error(f"导入依赖失败: {e}")
    raise ImportError("请安装所有依赖: pip install -r requirements.txt")

# 支持的文档格式
SUPPORTED_FORMATS = ['.docx', '.wpsx']
# WPS 二进制格式需要额外转换（通过 LibreOffice）
WPS_LEGACY_FORMAT = '.wps'

# ==================== 全局状态管理 ====================

# 工作目录配置
WORK_DIR = Path("./workspace").resolve()
WORK_DIR.mkdir(exist_ok=True)

# 会话状态存储（简单内存存储，生产环境建议使用数据库）
session_states: Dict[str, Dict] = {}


# ==================== 文档格式支持 ====================

def open_document(file_path: str):
    """
    打开文档，支持 .docx 和 .wpsx 格式
    
    .wpsx 本质上是与 .docx 相同的 ZIP+XML 格式，
    可以直接用 python-docx 读取。
    
    Args:
        file_path: 文档路径
    
    Returns:
        python-docx Document 对象
    
    Raises:
        ValueError: 不支持的格式
        Exception: 文档打开失败
    """
    suffix = Path(file_path).suffix.lower()
    if suffix in ['.docx', '.wpsx']:
        return Document(file_path)
    elif suffix == WPS_LEGACY_FORMAT:
        raise ValueError(
            f".wps 是 WPS 旧版二进制格式，不支持直接读取。\n"
            f"请使用 WPS Office 打开后另存为 .docx 或 .wpsx 格式，\n"
            f"或使用 LibreOffice 命令行转换: "
            f"libreoffice --headless --convert-to docx \"{file_path}\""
        )
    else:
        raise ValueError(f"不支持的文件格式: {suffix}")


def is_supported_format(file_path: str) -> bool:
    """检查文件格式是否被支持"""
    suffix = Path(file_path).suffix.lower()
    return suffix in SUPPORTED_FORMATS


# ==================== 安全工具函数 ====================

def safe_path(base_dir: Path, user_path: str) -> Path:
    """
    安全路径校验，防止目录穿越攻击
    
    Args:
        base_dir: 允许的根目录
        user_path: 用户提供的相对路径
    
    Returns:
        解析后的绝对路径
    
    Raises:
        ValueError: 如果路径不安全
    """
    resolved = (base_dir / user_path).resolve()
    if not str(resolved).startswith(str(base_dir.resolve())):
        raise ValueError(f"不安全的路径访问: {user_path}")
    return resolved


def generate_error_response(error_msg: str, error_code: str = "UNKNOWN_ERROR") -> str:
    """
    生成标准化的错误响应
    
    Args:
        error_msg: 错误消息
        error_code: 错误代码
    
    Returns:
        JSON 格式的错误响应字符串
    """
    return json.dumps({
        "success": False,
        "error": {
            "code": error_code,
            "message": error_msg
        },
        "timestamp": datetime.now().isoformat()
    }, ensure_ascii=False, indent=2)


def generate_success_response(data: Any, message: str = "操作成功") -> str:
    """
    生成标准化的成功响应
    
    Args:
        data: 返回数据
        message: 成功消息
    
    Returns:
        JSON 格式的成功响应字符串
    """
    return json.dumps({
        "success": True,
        "message": message,
        "data": data,
        "timestamp": datetime.now().isoformat()
    }, ensure_ascii=False, indent=2)


# ==================== 辅助函数 ====================

def extract_event_from_text(text: str, index: int) -> Optional[Dict]:
    """
    从文本中提取单个事件（简化版，可根据需要增强）
    
    Args:
        text: 段落文本
        index: 段落索引
    
    Returns:
        事件字典或 None
    """
    import re
    
    # 简单的时间模式匹配（可扩展）
    time_patterns = [
        r'\d{4}年\d{1,2}月\d{1,2}日\s*\d{1,2}[:::]\d{2}',
        r'\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}',
        r'\d{4}/\d{2}/\d{2}\s+\d{2}:\d{2}',
        r'\d{1,2}月\d{1,2}日\s*\d{1,2}[:::]\d{2}',
    ]
    
    for pattern in time_patterns:
        match = re.search(pattern, text)
        if match:
            # 从描述中剔除已匹配的时间前缀，避免标签中重复显示
            desc = text[:150]
            time_str = match.group()
            # 去掉描述开头的时间字符串（含其后的空格/分隔符）
            desc_without_time = re.sub(r'^\s*' + re.escape(time_str) + r'\s*[，,、:\-]\s*', '', desc)
            desc_without_time = re.sub(r'^\s*' + re.escape(time_str) + r'\s*', '', desc_without_time)
            if not desc_without_time.strip():
                desc_without_time = desc  # 如果去掉时间后为空，保留原文
            
            return {
                "index": index,
                "time_str": time_str,
                "description": desc_without_time[:150],
                "raw_text": text
            }
    
    return None


def parse_and_sort_events(events: List[Dict]) -> List[Dict]:
    """
    解析事件时间并排序
    
    Args:
        events: 原始事件列表
    
    Returns:
        解析并排序后的事件列表
    """
    from datetime import datetime
    import re
    
    parsed = []
    
    for event in events:
        time_str = event.get("time_str", "")
        
        # 尝试多种时间格式解析
        formats = [
            "%Y年%m月%d日 %H:%M",
            "%Y-%m-%d %H:%M",
            "%Y/%m/%d %H:%M",
            "%m月%d日 %H:%M",
            "%Y年%m月%d日",
            "%Y-%m-%d",
        ]
        
        parsed_time = None
        for fmt in formats:
            try:
                # 清理时间字符串
                clean_time = re.sub(r'\s+', ' ', time_str.strip())
                parsed_time = datetime.strptime(clean_time, fmt)
                break
            except ValueError:
                continue
        
        if parsed_time:
            parsed.append({
                "time": parsed_time,
                "description": event.get("description", ""),
                "original_time_str": time_str
            })
    
    # 按时间排序
    parsed.sort(key=lambda x: x["time"])
    
    return parsed


def create_timeline_chart(
    events: List[Dict],
    output_path: str,
    title: str = "事件时间轴",
    figsize: tuple = (16, 9),
    font_path: Optional[str] = None,
    important_indices: Optional[List[int]] = None
) -> Path:
    """
    创建专业风格时间轴图表（含自适应标签布局，避免时间相近的事件标签重叠）
    
    标签自动换行策略：
    - 描述文本超出 20 个中文字符宽度时自动换行，而非直接截断
    - 每行最多显示 20 个字，超长描述分多行显示
    - 最多显示 4 行（约 80 个字），超出部分用省略号
    
    Args:
        events: 解析后的事件列表
        output_path: 输出文件路径
        title: 图表标题
        figsize: 图片尺寸
        font_path: 字体路径
        important_indices: 重要中间事件的索引列表（从0开始），会用特殊星标标记
    
    Returns:
        输出文件的 Path 对象
    """
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyBboxPatch
    
    # 设置中文字体
    if font_path:
        font_prop = FontProperties(fname=font_path)
    else:
        font_prop = FontProperties()
        font_prop.set_family(['SimHei', 'Microsoft YaHei', 'Arial Unicode MS'])
    
    # 重要事件集合
    if important_indices is None:
        important_indices = []
    important_set = set(important_indices)
    
    # ===== 自适应标签布局算法 =====
    times = [e["time"] for e in events]
    descriptions = [e["description"] for e in events]
    n = len(events)
    
    # 计算时间跨度，决定分组粒度
    total_span = (max(times) - min(times)).total_seconds()
    cluster_threshold = max(900, min(3600, total_span * 0.08))
    
    # 分组
    groups = []
    if n > 0:
        current_group_start = 0
        for i in range(1, n):
            gap = (times[i] - times[i - 1]).total_seconds()
            if gap > cluster_threshold:
                groups.append((current_group_start, i))
                current_group_start = i
        groups.append((current_group_start, n))
    
    # 为每个事件分配 y 层级
    BASE_Y_LEVELS = [-0.45, 0.45, -0.75, 0.75, -1.05, 1.05, -1.35, 1.35]
    y_positions = [0.0] * n
    label_offsets = []
    
    for gs, ge in groups:
        group_size = ge - gs
        if group_size == 1:
            y_positions[gs] = 0.4
            label_offsets.append((0.55, 'bottom'))
        else:
            for j in range(gs, ge):
                local_idx = j - gs
                level = BASE_Y_LEVELS[local_idx % len(BASE_Y_LEVELS)]
                y_positions[j] = level
                if level > 0:
                    label_offsets.append((level + 0.15, 'bottom'))
                else:
                    label_offsets.append((level - 0.15, 'top'))
    
    # 动态调整图片高度
    max_y_level = max(abs(y) for y in y_positions) if y_positions else 0.4
    required_height = max(figsize[1], 4 + max_y_level * 4)
    fig_height = max(figsize[1], required_height)
    
    fig, ax = plt.subplots(figsize=(figsize[0], fig_height))
    
    # ===== 专业配色方案 =====
    # 使用蓝-青-紫渐变的专业色系
    timeline_colors = [
        '#2563EB',  # 主蓝
        '#0D9488',  # 青绿
        '#7C3AED',  # 紫色
        '#0891B2',  # 青色
        '#4F46E5',  # 靛蓝
        '#059669',  # 翠绿
        '#8B5CF6',  # 淡紫
        '#0284C7',  # 天蓝
        '#0EA5E9',  # 亮蓝
        '#6366F1',  # 蓝紫
    ]
    # 如果事件多于颜色数，循环使用
    event_colors = [timeline_colors[i % len(timeline_colors)] for i in range(n)]
    
    # ===== 绘制背景 =====
    ax.set_facecolor('#FAFBFC')
    fig.patch.set_facecolor('white')
    
    # ===== 绘制时间主轴 =====
    # 主轴线：粗且有渐变感
    time_range = [min(times), max(times)]
    # 画一条粗的主轴线
    ax.plot(time_range, [0, 0], color='#CBD5E1', linewidth=4, zorder=1, solid_capstyle='round')
    # 叠加细的深色线增强层次
    ax.plot(time_range, [0, 0], color='#94A3B8', linewidth=1.5, zorder=2, solid_capstyle='round')
    
    # ===== 绘制顶部和底部点缀线（平衡视觉） =====
    for y_base in [1.8, -1.8]:
        ax.axhline(y=y_base, xmin=0.05, xmax=0.95, color='#F1F5F9', linewidth=0.5, zorder=0)
    
    # ===== 自动换行标签函数 =====
    def wrap_text(text: str, max_chars_per_line: int = 20, max_lines: int = 4) -> str:
        if len(text) <= max_chars_per_line:
            return text
        lines = []
        for i in range(0, len(text), max_chars_per_line):
            if len(lines) >= max_lines:
                break
            lines.append(text[i:i + max_chars_per_line])
        total_chars = len(lines) * max_chars_per_line
        if total_chars < len(text):
            last = lines[-1]
            if len(last) > 3:
                lines[-1] = last[:-1] + '…'
            else:
                lines[-1] = last + '…'
        return '\n'.join(lines)
    
    # ===== 绘制事件点和标签 =====
    for i, (time, y_pos, desc) in enumerate(zip(times, y_positions, descriptions)):
        color = event_colors[i]
        is_top = y_pos > 0
        is_first = (i == 0)
        is_last = (i == n - 1)
        is_important = i in important_set
        
        # 普通事件标记尺寸
        outer_size = 14
        inner_size = 8
        
        # --- 连接虚线：从主轴到事件点 ---
        conn_style = '--'
        conn_width = 1.2
        conn_alpha = 0.35
        if is_first or is_last:
            conn_style = '-'
            conn_width = 1.8
            conn_alpha = 0.5
        elif is_important:
            conn_style = '-'
            conn_width = 1.5
            conn_alpha = 0.45
        
        ax.plot(
            [time, time], [0, y_pos],
            color=color, linewidth=conn_width, linestyle=conn_style,
            alpha=conn_alpha, zorder=3, solid_capstyle='round'
        )
        
        # --- 事件点标记 ---
        if is_first:
            # 开始事件：绿色大菱形
            marker_color = '#059669'
            ax.scatter(
                [time], [y_pos],
                s=100, marker='D', color='white',
                zorder=6, edgecolors=marker_color, linewidths=2.5
            )
            ax.scatter(
                [time], [y_pos],
                s=40, marker='D', color=marker_color,
                zorder=7, edgecolors='none', alpha=0.5
            )
        elif is_last:
            # 结束事件：红色大方形
            marker_color = '#DC2626'
            ax.scatter(
                [time], [y_pos],
                s=100, marker='s', color='white',
                zorder=6, edgecolors=marker_color, linewidths=2.5
            )
            ax.scatter(
                [time], [y_pos],
                s=40, marker='s', color=marker_color,
                zorder=7, edgecolors='none', alpha=0.5
            )
        elif is_important:
            # 重要中间事件：金色大星标 + 光晕
            marker_color = '#F59E0B'  # 琥珀金
            # 脉冲光晕（两层）
            for radius_factor, alpha_val in [(1.8, 0.08), (1.3, 0.15)]:
                ax.scatter(
                    [time], [y_pos],
                    s=(inner_size * radius_factor) ** 2,
                    color=marker_color, alpha=alpha_val,
                    zorder=4, edgecolors='none'
                )
            # 外圈白底
            ax.scatter(
                [time], [y_pos],
                s=(inner_size * 1.3) ** 2, color='white',
                zorder=5, edgecolors=marker_color, linewidths=2.5
            )
            # 内星形
            ax.scatter(
                [time], [y_pos],
                s=(inner_size * 0.9) ** 2, marker='*',
                color=marker_color,
                zorder=6, edgecolors='none'
            )
            # 改用金色作为事件颜色（标签、连线统一）
            color = marker_color
        else:
            # 普通事件：圆形
            marker_color = color
            ax.scatter(
                [time], [y_pos],
                s=outer_size ** 2, color=color,
                alpha=0.2, zorder=4, edgecolors='none'
            )
            ax.scatter(
                [time], [y_pos],
                s=inner_size ** 2, color='white',
                zorder=5, edgecolors=color, linewidths=2.5
            )
            ax.scatter(
                [time], [y_pos],
                s=(inner_size * 0.4) ** 2, color=color,
                zorder=6, edgecolors='none', alpha=0.8
            )
        
        # --- 事件编号标识（所有事件都保留） ---
        # 开始/结束使用白色编号，重要事件使用深色，普通事件使用各自颜色
        if is_first or is_last:
            num_color = 'white'
        elif is_important:
            num_color = '#92400E'  # 深金色，保证在星标上可读
        else:
            num_color = color
        
        ax.annotate(
            str(i + 1),
            xy=(time, y_pos),
            ha='center', va='center',
            fontsize=5.5, fontweight='bold',
            color=num_color, zorder=8
        )
        
        # --- 开始/结束特殊标注文字 ---
        if is_first:
            ax.annotate(
                '开始', xy=(time, y_pos), xytext=(time, y_pos + (0.25 if is_top else -0.25)),
                ha='center', va='bottom' if is_top else 'top',
                fontsize=7, fontweight='bold', color='#059669',
                fontproperties=font_prop, zorder=8
            )
        elif is_last:
            ax.annotate(
                '结束', xy=(time, y_pos), xytext=(time, y_pos + (0.25 if is_top else -0.25)),
                ha='center', va='bottom' if is_top else 'top',
                fontsize=7, fontweight='bold', color='#DC2626',
                fontproperties=font_prop, zorder=8
            )
        
        # --- 标签卡片 ---
        time_str = time.strftime('%m-%d %H:%M')
        wrapped_desc = wrap_text(desc, max_chars_per_line=20, max_lines=4)
        label = f"{time_str}\n{wrapped_desc}"
        text_offset, va = label_offsets[i]
        
        # 如果是多行描述，微调标签偏移量
        desc_lines = wrapped_desc.count('\n') + 1
        if desc_lines > 1:
            if is_top:
                text_offset += 0.05 * (desc_lines - 1)
            else:
                text_offset -= 0.05 * (desc_lines - 1)
        
        # 标签边框颜色与事件颜色一致
        bbox_color = color
        
        ax.annotate(
            label,
            xy=(time, y_pos),
            xytext=(time, text_offset),
            ha='center',
            va=va,
            fontsize=8,
            fontproperties=font_prop,
            bbox=dict(
                boxstyle='round,pad=0.4',
                facecolor='white',
                alpha=0.95,
                edgecolor=bbox_color,
                linewidth=0.8
            ),
            arrowprops=dict(
                arrowstyle='->',
                connectionstyle='arc3,rad=0',
                alpha=0.3,
                color=bbox_color,
                linewidth=0.8
            )
        )
    
    # ===== x 轴日期格式化（单行，避免和底部标签重叠） =====
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%m/%d %H:%M'))
    ax.xaxis.set_major_locator(mdates.AutoDateLocator())
    ax.tick_params(axis='x', which='both', length=0, labelsize=8, colors='#64748B')
    plt.setp(ax.get_xticklabels(), ha='center', fontproperties=font_prop)
    
    # ===== 标题区域（增加顶部间距，避免重叠） =====
    title_main = title
    subtitle = f"共 {n} 个事件  |  时间范围: {min(times).strftime('%m-%d %H:%M')} ~ {max(times).strftime('%m-%d %H:%M')}"
    
    # 使用 fig.suptitle 将主标题放在图形顶部，不占据 axes 空间
    fig.suptitle(
        title_main,
        fontsize=18, fontweight='bold',
        fontproperties=font_prop,
        color='#1E293B',
        x=0.02, ha='left',
        y=0.98, va='top'
    )
    # 副标题放在主标题下方
    fig.text(
        0.02, 0.93, subtitle,
        fontsize=9, color='#94A3B8',
        fontproperties=font_prop,
        va='top', ha='left'
    )
    
    # 隐藏 axes 自身的标题
    ax.set_title('')
    
    # ===== 隐藏坐标轴边框 =====
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_color('#E2E8F0')
    ax.spines['bottom'].set_linewidth(0.5)
    
    # ===== 移除 y 轴刻度 =====
    ax.set_yticks([])
    
    # ===== 设置 y 轴范围 =====
    y_margin = max_y_level + 0.8
    ax.set_ylim(-y_margin, y_margin)
    
    # ===== x 轴范围留边 =====
    x_padding = (max(times) - min(times)) * 0.03
    if n > 1:
        ax.set_xlim(min(times) - x_padding, max(times) + x_padding)
    
    # ===== 图例 =====
    legend_elements = []
    legend_elements.append(
        plt.Line2D([0], [0], marker='D', color='w',
                   markerfacecolor='#059669', markersize=6, label='开始')
    )
    legend_elements.append(
        plt.Line2D([0], [0], marker='s', color='w',
                   markerfacecolor='#DC2626', markersize=6, label='结束')
    )
    if important_set:
        legend_elements.append(
            plt.Line2D([0], [0], marker='*', color='w',
                       markerfacecolor='#F59E0B', markersize=8, label='重要事件')
        )
    if n > 2:
        legend_elements.append(
            plt.Line2D([0], [0], marker='o', color='w',
                       markerfacecolor='#2563EB', markersize=5, label='普通事件')
        )
    legend_elements.append(
        plt.Line2D([0], [0], marker='s', color='w',
                   markerfacecolor='#DC2626', markersize=6, label='结束')
    )
    if n > 2:
        legend_elements.append(
            plt.Line2D([0], [0], marker='o', color='w',
                       markerfacecolor='#2563EB', markersize=5, label='中间事件')
        )
    if legend_elements:
        ax.legend(
            handles=legend_elements,
            loc='upper right',
            fontsize=7,
            framealpha=0.8,
            edgecolor='#E2E8F0',
            prop=font_prop
        )
    
    # ===== 调整布局 =====
    plt.tight_layout(rect=[0, 0, 1, 0.88])  # 顶部留出 12% 空间给标题区域
    
    # ===== 保存图片 =====
    plt.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    
    return Path(output_path)


def find_insertion_point(doc: Document, keyword: str, position: str) -> int:
    """
    查找插入位置
    
    Args:
        doc: 文档对象
        keyword: 关键词
        position: 位置策略
    
    Returns:
        插入位置的段落索引
    """
    for idx, paragraph in enumerate(doc.paragraphs):
        if keyword in paragraph.text:
            if position == "before":
                return idx
            elif position == "after":
                return min(idx + 1, len(doc.paragraphs))
    
    # 如果未找到关键词，返回文档末尾
    return len(doc.paragraphs)


# ==================== 核心业务函数 ====================

def analyze_document_structure_impl(doc_path: str) -> str:
    """
    分析 Word 文档的结构，识别章节、段落和关键内容
    
    Args:
        doc_path: Word 文档的相对路径（相对于 workspace 目录）
    
    Returns:
        JSON 字符串，包含文档结构信息
    """
    try:
        logger.info(f"[analyze_document_structure] 开始分析文档: {doc_path}")
        
        # 安全路径校验
        full_path = safe_path(WORK_DIR, doc_path)
        
        if not full_path.exists():
            logger.error(f"[analyze_document_structure] 文件不存在: {full_path}")
            return generate_error_response(f"文件不存在: {doc_path}", "FILE_NOT_FOUND")
        
        if not is_supported_format(str(full_path)):
            suffix = full_path.suffix.lower()
            if suffix == WPS_LEGACY_FORMAT:
                return generate_error_response(
                    ".wps 是 WPS 旧版二进制格式，不支持直接读取。请使用 WPS Office 打开后另存为 .docx 或 .wpsx 格式。",
                    "WPS_LEGACY_FORMAT"
                )
            logger.error(f"[analyze_document_structure] 不支持的文件格式: {full_path.suffix}")
            return generate_error_response(f"仅支持 .docx 和 .wpsx 格式的文档", "UNSUPPORTED_FORMAT")
        
        # 读取文档
        doc = open_document(str(full_path))
        
        structure = {
            "file_name": full_path.name,
            "total_paragraphs": len(doc.paragraphs),
            "sections": [],
            "headings": []
        }
        
        current_section = None
        
        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            style = paragraph.style.name if paragraph.style else "Normal"
            
            # 识别标题（Heading 样式）
            if 'Heading' in style or '标题' in style:
                heading_info = {
                    "index": idx,
                    "text": text,
                    "style": style,
                    "level": int(style.replace('Heading ', '').replace('标题 ', '')) if any(c.isdigit() for c in style) else 1
                }
                structure["headings"].append(heading_info)
                
                # 开始新章节
                current_section = {
                    "title": text,
                    "start_index": idx,
                    "paragraphs": []
                }
                structure["sections"].append(current_section)
            elif current_section is not None:
                current_section["paragraphs"].append({
                    "index": idx,
                    "text": text[:200],  # 限制长度避免过大
                    "style": style
                })
        
        logger.info(f"[analyze_document_structure] 分析完成，发现 {len(structure['headings'])} 个标题")
        return generate_success_response(
            structure, 
            f"文档分析完成，共 {len(structure['headings'])} 个标题，{len(structure['sections'])} 个章节"
        )
    
    except ValueError as e:
        logger.error(f"[analyze_document_structure] 路径安全校验失败: {str(e)}")
        return generate_error_response(str(e), "SECURITY_ERROR")
    except Exception as e:
        logger.error(f"[analyze_document_structure] 处理异常: {str(e)}", exc_info=True)
        return generate_error_response(f"文档分析失败: {str(e)}", "ANALYSIS_ERROR")


def extract_timeline_events_impl(doc_path: str, section_keywords: str = "处置,过程,时间线,经过") -> str:
    """
    从文档中提取时间线相关的事件信息
    
    Args:
        doc_path: Word 文档的相对路径
        section_keywords: 用于定位相关章节的关键词，逗号分隔
    
    Returns:
        JSON 字符串，包含提取的事件列表
    """
    try:
        logger.info(f"[extract_timeline_events] 开始提取事件，关键词: {section_keywords}")
        
        full_path = safe_path(WORK_DIR, doc_path)
        
        if not full_path.exists():
            return generate_error_response(f"文件不存在: {doc_path}", "FILE_NOT_FOUND")
        
        doc = open_document(str(full_path))
        keywords = [kw.strip() for kw in section_keywords.split(",") if kw.strip()]
        
        events = []
        current_relevant_section = False
        
        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            style = paragraph.style.name if paragraph.style else "Normal"
            
            # 检查是否进入相关章节
            if any(kw in text for kw in keywords) and ('Heading' in style or '标题' in style):
                current_relevant_section = True
                logger.info(f"[extract_timeline_events] 找到相关章节: {text}")
                continue
            
            # 如果在相关章节中，尝试提取事件
            if current_relevant_section:
                # 如果遇到新标题，结束当前章节
                if 'Heading' in style or '标题' in style:
                    current_relevant_section = False
                    continue
                
                # 简单的事件提取逻辑（基于时间表达式）
                event = extract_event_from_text(text, idx)
                if event:
                    events.append(event)
        
        # 生成会话 ID 用于后续操作
        session_id = hashlib.md5(f"{doc_path}{datetime.now().isoformat()}".encode()).hexdigest()[:12]
        session_states[session_id] = {
            "doc_path": doc_path,
            "events": events,
            "created_at": datetime.now().isoformat()
        }
        
        logger.info(f"[extract_timeline_events] 提取完成，共 {len(events)} 个事件，会话ID: {session_id}")
        return generate_success_response(
            {
                "session_id": session_id,
                "event_count": len(events),
                "events": events
            },
            f"成功提取 {len(events)} 个事件"
        )
    
    except Exception as e:
        logger.error(f"[extract_timeline_events] 处理异常: {str(e)}", exc_info=True)
        return generate_error_response(f"事件提取失败: {str(e)}", "EXTRACTION_ERROR")


def review_extracted_events_impl(session_id: str) -> str:
    """
    查看已提取的事件列表，供人工审核
    
    Args:
        session_id: 会话 ID（从 extract_timeline_events 返回）
    
    Returns:
        JSON 字符串，包含事件详情
    """
    try:
        logger.info(f"[review_extracted_events] 查看会话: {session_id}")
        
        if session_id not in session_states:
            return generate_error_response(f"会话不存在: {session_id}", "SESSION_NOT_FOUND")
        
        session = session_states[session_id]
        
        logger.info(f"[review_extracted_events] 返回 {len(session['events'])} 个事件")
        return generate_success_response(
            {
                "session_id": session_id,
                "doc_path": session["doc_path"],
                "event_count": len(session["events"]),
                "events": session["events"]
            },
            "事件列表获取成功"
        )
    
    except Exception as e:
        logger.error(f"[review_extracted_events] 处理异常: {str(e)}", exc_info=True)
        return generate_error_response(f"查看事件失败: {str(e)}", "REVIEW_ERROR")


def modify_event_impl(session_id: str, event_index: int, new_time: str = "", new_description: str = "") -> str:
    """
    修改指定事件的详细信息
    
    Args:
        session_id: 会话 ID
        event_index: 事件在列表中的索引（从0开始）
        new_time: 新的时间字符串（可选）
        new_description: 新的描述（可选）
    
    Returns:
        JSON 字符串，包含修改后的事件信息
    """
    try:
        logger.info(f"[modify_event] 修改会话 {session_id} 的事件 {event_index}")
        
        if session_id not in session_states:
            return generate_error_response(f"会话不存在: {session_id}", "SESSION_NOT_FOUND")
        
        session = session_states[session_id]
        
        if event_index < 0 or event_index >= len(session["events"]):
            return generate_error_response(
                f"事件索引超出范围: {event_index} (总共 {len(session['events'])} 个事件)",
                "INDEX_OUT_OF_RANGE"
            )
        
        event = session["events"][event_index]
        
        # 更新字段
        if new_time:
            old_time = event.get("time_str", "")
            event["time_str"] = new_time
            logger.info(f"[modify_event] 时间修改: '{old_time}' -> '{new_time}'")
        
        if new_description:
            old_desc = event.get("description", "")
            event["description"] = new_description
            logger.info(f"[modify_event] 描述已更新")
        
        logger.info(f"[modify_event] 事件 {event_index} 修改完成")
        return generate_success_response(
            {
                "session_id": session_id,
                "modified_event_index": event_index,
                "updated_event": event
            },
            f"事件 {event_index} 修改成功"
        )
    
    except Exception as e:
        logger.error(f"[modify_event] 处理异常: {str(e)}", exc_info=True)
        return generate_error_response(f"修改事件失败: {str(e)}", "MODIFY_ERROR")


def delete_event_impl(session_id: str, event_index: int) -> str:
    """
    删除指定的事件
    
    Args:
        session_id: 会话 ID
        event_index: 事件在列表中的索引（从0开始）
    
    Returns:
        JSON 字符串，包含删除确认信息
    """
    try:
        logger.info(f"[delete_event] 删除会话 {session_id} 的事件 {event_index}")
        
        if session_id not in session_states:
            return generate_error_response(f"会话不存在: {session_id}", "SESSION_NOT_FOUND")
        
        session = session_states[session_id]
        
        if event_index < 0 or event_index >= len(session["events"]):
            return generate_error_response(
                f"事件索引超出范围: {event_index}",
                "INDEX_OUT_OF_RANGE"
            )
        
        deleted_event = session["events"].pop(event_index)
        
        logger.info(f"[delete_event] 事件 {event_index} 已删除")
        return generate_success_response(
            {
                "session_id": session_id,
                "deleted_event_index": event_index,
                "remaining_event_count": len(session["events"])
            },
            f"事件 {event_index} 删除成功，剩余 {len(session['events'])} 个事件"
        )
    
    except Exception as e:
        logger.error(f"[delete_event] 处理异常: {str(e)}", exc_info=True)
        return generate_error_response(f"删除事件失败: {str(e)}", "DELETE_ERROR")


def add_manual_event_impl(session_id: str, time_str: str, description: str) -> str:
    """
    手动添加一个新事件到事件列表
    
    Args:
        session_id: 会话 ID
        time_str: 事件时间（格式如: 2024-01-15 14:30）
        description: 事件描述
    
    Returns:
        JSON 字符串，包含添加的事件信息
    """
    try:
        logger.info(f"[add_manual_event] 向会话 {session_id} 添加事件")
        
        if session_id not in session_states:
            return generate_error_response(f"会话不存在: {session_id}", "SESSION_NOT_FOUND")
        
        session = session_states[session_id]
        
        new_event = {
            "index": -1,  # 手动添加的事件没有原始索引
            "time_str": time_str,
            "description": description,
            "is_manual": True
        }
        
        session["events"].append(new_event)
        
        logger.info(f"[add_manual_event] 事件添加成功，当前共 {len(session['events'])} 个事件")
        return generate_success_response(
            {
                "session_id": session_id,
                "added_event": new_event,
                "total_event_count": len(session["events"])
            },
            f"事件添加成功，当前共 {len(session['events'])} 个事件"
        )
    
    except Exception as e:
        logger.error(f"[add_manual_event] 处理异常: {str(e)}", exc_info=True)
        return generate_error_response(f"添加事件失败: {str(e)}", "ADD_EVENT_ERROR")


def generate_timeline_image_impl(
    session_id: str,
    output_filename: str = "timeline.png",
    title: str = "事件处置时间轴",
    width: int = 16,
    height: int = 9,
    font_path: str = "",
    important_indices: str = ""
) -> str:
    """
    根据提取的事件生成时间轴图
    
    Args:
        session_id: 会话 ID
        output_filename: 输出图片文件名
        title: 图表标题
        width: 图片宽度（英寸）
        height: 图片高度（英寸）
        font_path: 中文字体路径（可选，用于正确显示中文）
        important_indices: 重要中间事件的索引列表，逗号分隔（如 "3,5,7"），在图表中用星标突出显示
    
    Returns:
        JSON 字符串，包含生成结果和图片路径
    """
    try:
        logger.info(f"[generate_timeline_image] 为会话 {session_id} 生成时间轴图")
        
        if session_id not in session_states:
            return generate_error_response(f"会话不存在: {session_id}", "SESSION_NOT_FOUND")
        
        session = session_states[session_id]
        events = session["events"]
        
        if not events:
            return generate_error_response("没有可绘制的事件，请先提取或添加事件", "NO_EVENTS")
        
        # 解析时间并排序
        parsed_events = parse_and_sort_events(events)
        
        if not parsed_events:
            return generate_error_response("无法解析任何有效的时间信息", "PARSE_TIME_ERROR")
        
        # 解析重要事件索引
        imp_indices = []
        if important_indices:
            for part in important_indices.split(","):
                part = part.strip()
                if part:
                    try:
                        idx = int(part)
                        if 0 <= idx < len(parsed_events):
                            imp_indices.append(idx)
                    except ValueError:
                        pass
        
        # 生成图片
        output_path = WORK_DIR / output_filename
        image_path = create_timeline_chart(
            parsed_events,
            str(output_path),
            title=title,
            figsize=(width, height),
            font_path=font_path if font_path else None,
            important_indices=imp_indices
        )
        
        logger.info(f"[generate_timeline_image] 时间轴图生成成功: {image_path}")
        return generate_success_response(
            {
                "session_id": session_id,
                "output_file": output_filename,
                "absolute_path": str(image_path),
                "event_count": len(parsed_events),
                "title": title
            },
            f"时间轴图生成成功: {output_filename}"
        )
    
    except Exception as e:
        logger.error(f"[generate_timeline_image] 处理异常: {str(e)}", exc_info=True)
        return generate_error_response(f"生成时间轴图失败: {str(e)}", "GENERATION_ERROR")


def insert_timeline_to_docx_impl(
    doc_path: str,
    session_id: str,
    image_filename: str = "timeline.png",
    target_section_keyword: str = "时间轴",
    position: str = "after"
) -> str:
    """
    将生成的时间轴图插入到 Word 文档中
    
    Args:
        doc_path: Word 文档相对路径
        session_id: 会话 ID
        image_filename: 时间轴图片文件名
        target_section_keyword: 目标章节关键词
        position: 插入位置（before/after/end）
    
    Returns:
        JSON 字符串，包含插入结果
    """
    try:
        logger.info(f"[insert_timeline_to_docx] 将时间轴插入到文档: {doc_path}")
        
        full_path = safe_path(WORK_DIR, doc_path)
        
        if not full_path.exists():
            return generate_error_response(f"文件不存在: {doc_path}", "FILE_NOT_FOUND")
        
        image_path = WORK_DIR / image_filename
        if not image_path.exists():
            return generate_error_response(f"图片文件不存在: {image_filename}", "IMAGE_NOT_FOUND")
        
        # 加载文档（支持 .docx 和 .wpsx）
        doc = open_document(str(full_path))
        
        # 查找插入位置
        insert_index = find_insertion_point(doc, target_section_keyword, position)
        
        # 插入图片
        if insert_index < len(doc.paragraphs):
            ref_paragraph = doc.paragraphs[insert_index]
            new_paragraph = ref_paragraph.insert_paragraph_before()
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(6))
        else:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(6))
        
        # 保存文档（创建备份）
        backup_path = full_path.with_suffix('.backup.docx')
        shutil.copy2(full_path, backup_path)
        logger.info(f"[insert_timeline_to_docx] 已创建备份: {backup_path}")
        
        doc.save(str(full_path))
        
        logger.info(f"[insert_timeline_to_docx] 时间轴插入成功")
        return generate_success_response(
            {
                "doc_path": doc_path,
                "session_id": session_id,
                "image_inserted": image_filename,
                "backup_created": str(backup_path.name)
            },
            "时间轴图已成功插入文档"
        )
    
    except Exception as e:
        logger.error(f"[insert_timeline_to_docx] 处理异常: {str(e)}", exc_info=True)
        return generate_error_response(f"插入时间轴失败: {str(e)}", "INSERT_ERROR")


def clear_session_impl(session_id: str) -> str:
    """
    清除会话数据，释放内存
    
    Args:
        session_id: 会话 ID
    
    Returns:
        JSON 字符串，包含清除结果
    """
    try:
        logger.info(f"[clear_session] 清除会话: {session_id}")
        
        if session_id not in session_states:
            return generate_error_response(f"会话不存在: {session_id}", "SESSION_NOT_FOUND")
        
        del session_states[session_id]
        
        logger.info(f"[clear_session] 会话 {session_id} 已清除")
        return generate_success_response(
            {"session_id": session_id},
            "会话数据已清除"
        )
    
    except Exception as e:
        logger.error(f"[clear_session] 处理异常: {str(e)}", exc_info=True)
        return generate_error_response(f"清除会话失败: {str(e)}", "CLEAR_ERROR")


def list_active_sessions_impl() -> str:
    """
    列出所有活跃的会话
    
    Returns:
        JSON 字符串，包含会话列表
    """
    try:
        logger.info("[list_active_sessions] 列出活跃会话")
        
        sessions_info = []
        for sid, session in session_states.items():
            sessions_info.append({
                "session_id": sid,
                "doc_path": session["doc_path"],
                "event_count": len(session["events"]),
                "created_at": session["created_at"]
            })
        
        logger.info(f"[list_active_sessions] 共 {len(sessions_info)} 个活跃会话")
        return generate_success_response(
            {
                "total_sessions": len(sessions_info),
                "sessions": sessions_info
            },
            f"共有 {len(sessions_info)} 个活跃会话"
        )
    
    except Exception as e:
        logger.error(f"[list_active_sessions] 处理异常: {str(e)}", exc_info=True)
        return generate_error_response(f"列出会话失败: {str(e)}", "LIST_ERROR")
