"""
端到端演示：生成银行IT事件分析报告 → 读取 → 检查 → 修订
"""
import json
import os
import sys
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# 第1步：生成一份模拟的"银行IT系统生产事件分析报告"（包含典型问题）
# ============================================================

def create_bank_incident_report():
    """创建一份有典型错误的银行IT生产事件分析报告"""
    doc = DocxDocument()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    title = doc.add_heading('', level=0)
    run = title.add_run('生产事件分析报告')
    run.font.size = Pt(22)

    # === 一、事件概述 ===
    doc.add_heading('一、事件概述', level=2)

    doc.add_paragraph(
        '2024年3月15日14:23分，核心交易系统发生服务中断事件，'
        '持续时间约45分钟，影响范围涉及个人网银、手机银行、ATM等渠道。'
        '经紧急处理后于15:08分恢复服务。'
    )

    # === 二、事件影响 ===
    doc.add_heading('二、事件影响', level=2)

    doc.add_paragraph(
        '本次事件影响了约3.2万名用户无法正常使用银行服务，'
        '其中个人网银用户约1.8万人，手机银行用户约1.2万人，'
        'ATM用户约2千人。事件期间共收到客户投诉电话1,256通。'
    )

    # === 三、处置过程（包含错误：项目组名称不规范） ===
    doc.add_heading('三、处置过程', level=2)

    doc.add_paragraph(
        '14:25 - 监控系统发出告警，项目组值班人员发现交易响应时间异常升高。\n'
        '14:28 - 项目组启动应急响应流程，通知相关技术人员。\n'
        '14:35 - 初步判断为数据库连接池耗尽，项目组决定重启数据库连接池。\n'
        '14:42 - 重启后问题未解决，项目组进一步排查发现是核心交易模块存在死锁。\n'
        '14:50 - 项目组对死锁进程进行强制终止，并调整连接池参数。\n'
        '15:08 - 系统恢复正常，项目组持续监控30分钟确认稳定。'
    )

    # === 四、根因分析（包含错别字"奔溃"、分析深度不足） ===
    doc.add_heading('四、根因分析', level=2)

    doc.add_paragraph(
        '经分析，发现是数据库连接池耗尽导致系统无法响应请求。'
        '在事件发生前，系统并发量突然增加，导致连接池资源被迅速耗尽。'
        '数据库连接池最大连接数设置为50，在高并发场景下出现了系统奔溃的情况。'
    )

    # === 五、改进措施（包含不具体、编号问题） ===
    doc.add_heading('五、改进措施', level=2)

    doc.add_paragraph(
        '1、调大数据库连接池；\n'
        '2、增加监控告警；\n'
        '3、完善应急响应流程；\n'
        '4、定期进行压力测试。'
    )

    # === 六、事件总结 ===
    doc.add_heading('六、事件总结', level=2)

    doc.add_paragraph(
        '本次事件暴露了系统在容量规划、监控告警、应急响应等方面存在不足。'
        '项目组将认真总结经验教训，落实改进措施，避免类似事件再次发生。'
    )

    # 保存
    path = 'test_bank_incident_report.docx'
    doc.save(path)
    return path


# ============================================================
# 第2步：读取文档内容
# ============================================================

def read_document(file_path):
    """使用 DocumentReader 读取文档"""
    from unified_mcp_server import DocumentReader
    reader = DocumentReader()
    result = reader.read_document(file_path)
    return result


# ============================================================
# 第3步：模拟LLM检查结果（硬性规则 + 软性规则 → 合并为 suggestions）
# ============================================================

def create_check_suggestions():
    """模拟第2步（硬性规则）+ 第3步（软性规则）+ 第4步（合并排序）的输出"""
    suggestions = [
        # === 硬性规则检查结果 ===
        {
            "id": "hard_001",
            "rule_id": "NAMING_TEAM_001",
            "rule_name": "项目组名称不规范",
            "type": "format",
            "severity": "Critical",
            "section": "三、处置过程",
            "original_text": "项目组值班人员发现交易响应时间异常升高",
            "suggestion": "核心交易系统项目组值班人员发现交易响应时间异常升高",
            "reason": '不能直接写“项目组”，必须写成“部门+系统+项目组”形式，全文保持一致。处置过程中出现5处“项目组”均不规范。'
        },
        {
            "id": "hard_002",
            "rule_id": "NAMING_SYSTEM_001",
            "rule_name": "系统名称首次出现未定义简称",
            "type": "format",
            "severity": "Medium",
            "section": "一、事件概述",
            "original_text": "核心交易系统发生服务中断事件",
            "suggestion": '核心交易系统（以下简称“核心系统”）发生服务中断事件',
            "reason": '系统名称首次出现时应使用“中文全称（以下简称“XX”）”格式'
        },
        {
            "id": "hard_003",
            "rule_id": "PUNCT_COMMA_001",
            "rule_name": "标点符号不符合规范",
            "type": "format",
            "severity": "Medium",
            "section": "一、事件概述",
            "original_text": "持续时间约45分钟,影响范围涉及个人网银",
            "suggestion": "持续时间约45分钟，影响范围涉及个人网银",
            "reason": "使用了英文逗号，应使用中文全角逗号"
        },
        {
            "id": "hard_004",
            "rule_id": "PUNCT_BRACKET_001",
            "rule_name": "括号使用了英文半角",
            "type": "format",
            "severity": "Low",
            "section": "四、根因分析",
            "original_text": "最大连接数设置为50",
            "suggestion": "最大连接数设置为50（当前值）",
            "reason": "补充说明应使用中文全角括号（）"
        },
        # === 软性规则检查结果 ===
        {
            "id": "soft_001",
            "rule_id": "LANGUAGE_TYPO_001",
            "rule_name": "存在错别字",
            "type": "language",
            "severity": "High",
            "section": "四、根因分析",
            "original_text": "出现了系统奔溃的情况",
            "suggestion": "出现了系统崩溃的情况",
            "reason": "'奔溃'为错别字，正确写法应为'崩溃'"
        },
        {
            "id": "soft_002",
            "rule_id": "CONTENT_ROOT_CAUSE_001",
            "rule_name": "根因分析深度不足",
            "type": "content",
            "severity": "Critical",
            "section": "四、根因分析",
            "original_text": "经分析，发现是数据库连接池耗尽导致系统无法响应请求。在事件发生前，系统并发量突然增加，导致连接池资源被迅速耗尽。",
            "suggestion": "经5Why分析，根本原因如下：\n【Why 1】为什么系统无法响应？→ 数据库连接池耗尽，所有线程阻塞等待连接\n【Why 2】为什么连接池耗尽？→ 并发请求量在14:20-14:23间突增300%，且核心交易模块存在慢SQL未优化\n【Why 3】为什么缺乏限流保护？→ 网关层未配置并发限流策略，上游流量直接透传至数据库\n【Why 4】为什么未配置限流？→ 容量规划仅在常规场景下评估，未考虑极端并发场景\n【Why 5】为什么容量规划不完善？→ 缺乏系统性的非功能需求评审机制\n根本原因：① 容量规划机制不完善 ② 限流保护策略缺失 ③ 慢SQL优化滞后",
            "reason": "当前分析停留在表面现象（连接池耗尽），未进行深度根因分析（5Why），缺乏对系统架构层面缺陷的挖掘"
        },
        {
            "id": "soft_003",
            "rule_id": "CONTENT_MEASURE_TARGET_001",
            "rule_name": "改进措施缺乏具体细节",
            "type": "content",
            "severity": "High",
            "section": "五、改进措施",
            "original_text": "1、调大数据库连接池；",
            "suggestion": "1、调大数据库连接池并优化连接管理\n   - 当前配置：最大连接数50，最大等待时间30s\n   - 调整后配置：最大连接数200，最大等待时间10s，空闲超时5min\n   - 调整依据：基于峰值并发量（TPS峰值约2000）的2倍余量计算\n   - 责任人：DBA团队（张三）\n   - 完成时间：2024年3月22日前",
            "reason": "措施过于笼统，缺乏具体数值、技术参数、调整依据、责任人和完成时间"
        },
        {
            "id": "soft_004",
            "rule_id": "CONTENT_MEASURE_MONITOR_001",
            "rule_name": "监控告警措施不够具体",
            "type": "content",
            "severity": "Medium",
            "section": "五、改进措施",
            "original_text": "2、增加监控告警；",
            "suggestion": "2、完善多层级监控告警体系\n   - 应用层：新增连接池使用率告警（阈值>70%告警，>90%严重告警）\n   - 数据库层：新增慢SQL监控（执行时间>1s记录，>5s告警）\n   - 网关层：新增并发限流告警（触发限流时即时通知）\n   - 告警渠道：企业微信+短信双通道\n   - 责任人：运维团队（李四）\n   - 完成时间：2024年3月25日前",
            "reason": "未说明监控的具体指标、阈值和告警渠道"
        },
        {
            "id": "soft_005",
            "rule_id": "LOGIC_ORDER_001",
            "rule_name": "改进措施排序缺乏优先级",
            "type": "logic",
            "severity": "Low",
            "section": "五、改进措施",
            "original_text": "1、调大数据库连接池；\n2、增加监控告警；\n3、完善应急响应流程；\n4、定期进行压力测试。",
            "suggestion": "按紧急度和影响范围重新排序：\nP0（本周内）：调大连接池 + 增加限流保护（直接防止再次发生）\nP1（两周内）：增加监控告警 + 优化慢SQL（提升发现和响应能力）\nP2（一个月内）：完善应急流程 + 定期压测（长效机制）",
            "reason": "当前措施没有区分紧急度，所有措施并列，不利于执行"
        }
    ]
    return suggestions


# ============================================================
# 第4步：执行文档修订
# ============================================================

def revise_document(file_path, suggestions):
    """使用 DocumentReviser 执行修订"""
    from unified_mcp_server import DocumentReviser

    reviser = DocumentReviser()
    suggestions_json = json.dumps(suggestions, ensure_ascii=False, indent=2)

    result = reviser.revise_document(
        file_path=file_path,
        suggestions_json=suggestions_json,
        revision_mode="track_changes"  # 使用 Track Changes 模式
    )
    return result


# ============================================================
# 第5步：对比展示
# ============================================================

def show_comparison(original_path, revised_path):
    """对比原始文档和修订后文档的内容"""
    print("\n" + "=" * 80)
    print("  原始文档 vs 修订后文档 对比")
    print("=" * 80)

    orig_doc = DocxDocument(original_path)
    rev_doc = DocxDocument(revised_path)

    print(f"\n{'─' * 40} 原始文档内容 {'─' * 40}")
    for i, para in enumerate(orig_doc.paragraphs):
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            print(f"  [{style}] {para.text[:100]}")

    print(f"\n{'─' * 40} 修订后文档内容 {'─' * 40}")
    for i, para in enumerate(rev_doc.paragraphs):
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            text = para.text[:150]
            # 标注修订标记
            print(f"  [{style}] {text}")


def show_revision_details(result):
    """展示修订详情"""
    print("\n" + "=" * 80)
    print("  修订执行结果")
    print("=" * 80)

    print(f"\n  状态: {'✅ 成功' if result.get('success') else '❌ 失败'}")
    print(f"  输出文件: {result.get('output_path')}")
    print(f"  修订模式: {result.get('revision_mode', 'N/A')}")
    print(f"  总建议数: {result.get('total_suggestions')}")
    print(f"  已应用:   {result.get('applied_revisions')}")
    print(f"  已跳过:   {result.get('skipped_revisions')}")

    print(f"\n{'─' * 40} 逐条修订明细 {'─' * 40}")
    for rev in result.get('revisions_detail', []):
        status_icon = "✅" if rev['status'] == 'applied' else "⏭️"
        print(f"\n  {status_icon} [{rev['severity']:8}] {rev['rule_name']} ({rev['type']})")
        print(f"     原文: {rev['original'][:80]}...")
        print(f"     改为: {rev['new'][:80]}...")
        print(f"     原因: {rev.get('reason', 'N/A')[:80]}")

    # 统计
    if result.get('statistics'):
        stats = result['statistics']
        print(f"\n{'─' * 40} 修订统计 {'─' * 40}")
        print(f"\n  按严重程度:")
        for sev, count in stats.get('by_severity', {}).items():
            bar = '█' * count
            print(f"    {sev:10}: {count} {bar}")
        print(f"\n  按类型:")
        for typ, count in stats.get('by_type', {}).items():
            bar = '█' * count
            print(f"    {typ:10}: {count} {bar}")


# ============================================================
# 主流程
# ============================================================

def main():
    print("=" * 80)
    print("  银行IT生产事件分析报告 — 端到端审校演示")
    print("  流程：生成报告 → 读取 → LLM检查(模拟) → 修订")
    print("=" * 80)

    # Step 1: 生成报告
    print("\n[Step 1] 生成模拟的银行IT事件分析报告...")
    docx_path = create_bank_incident_report()
    print(f"  ✅ 已生成: {docx_path}")
    print(f"     大小: {os.path.getsize(docx_path)} bytes")

    # Step 2: 读取文档
    print("\n[Step 2] 读取文档内容...")
    try:
        doc_data = read_document(docx_path)
        print(f"  ✅ 读取成功")
        print(f"     段落数: {doc_data.get('paragraph_count', 'N/A')}")
        print(f"     内容预览:")
        for para in doc_data.get('paragraphs', [])[:8]:
            text = para.get('text', '')[:80]
            if text.strip():
                print(f"       {text}...")
    except Exception as e:
        print(f"  ⚠️ 读取异常（将跳过读取步骤继续）: {e}")

    # Step 3: 创建检查建议
    print("\n[Step 3] 模拟 LLM 检查结果（硬性规则 + 软性规则）...")
    suggestions = create_check_suggestions()
    print(f"  ✅ 共 {len(suggestions)} 条建议")
    for s in suggestions:
        print(f"     [{s['severity']:8}] [{s['type']:8}] {s['rule_name']}")

    # Step 4: 执行修订
    print("\n[Step 4] 执行文档修订（Track Changes 模式）...")
    try:
        result = revise_document(docx_path, suggestions)
        show_revision_details(result)
        revised_path = result.get('output_path')
    except Exception as e:
        print(f"  ❌ 修订失败: {e}")
        import traceback
        traceback.print_exc()
        return

    # Step 5: 对比展示
    if revised_path and os.path.exists(revised_path):
        show_comparison(docx_path, revised_path)
        print(f"\n{'=' * 80}")
        print(f"  完成！输出文件：")
        print(f"    原始报告: {os.path.abspath(docx_path)}")
        print(f"    修订报告: {os.path.abspath(revised_path)}")
        print(f"    建议数据: 共 {len(suggestions)} 条（4条硬性 + 5条软性）")
        print(f"{'=' * 80}")
    else:
        print(f"\n  ❌ 修订后文件未生成")


if __name__ == "__main__":
    main()
