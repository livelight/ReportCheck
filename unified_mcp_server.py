"""
统一MCP服务器 - 同时启动读入文档和修改文档的所有MCP服务

功能：
1. 第1步 - 文档读取：读取文档内容，输出JSON
2. 第1.5步 - 系统信息读取：从Excel系统信息表读取系统基础信息
3. 第6步 - 文档修订：根据检查建议修订文档

支持格式：Word文档(.docx)、WPS文档(.wps和.wpsx)
系统信息格式：Excel表格(.xlsx)
启动方式：stdio模式 或 SSE模式

使用方法:
    python unified_mcp_server.py [--transport stdio|sse] [--host 0.0.0.0] [--port 8000]

端点(SSE模式):
    GET /sse - SSE连接端点
    POST /messages - 发送消息
    GET /health - 健康检查
    GET /tools - 获取工具定义
"""

import json
import os
import re
import sys
import logging
import zipfile
import argparse
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, asdict
from datetime import datetime
from enum import Enum

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger('UnifiedMCPServer')

# FastMCP导入 - 使用全局唯一的MCP实例
try:
    from mcp_instance import mcp
    FASTMCP_AVAILABLE = True
except ImportError:
    FASTMCP_AVAILABLE = False
    logger.warning("mcp_instance模块未找到，请确保mcp_instance.py存在")

# 文档处理库
try:
    from docx import Document as DocxDocument
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx未安装，请运行: pip install python-docx")

# Excel处理库（用于读取系统信息表）
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl未安装，请运行: pip install openpyxl")


# ============================================================
# 数据模型
# ============================================================

@dataclass
class DocumentMetadata:
    """文档元数据模型"""
    title: str = ""
    author: str = ""
    created_date: str = ""
    modified_date: str = ""
    file_type: str = ""
    file_size: int = 0
    paragraph_count: int = 0
    table_count: int = 0
    page_count: Optional[int] = None
    company: Optional[str] = None


@dataclass
class CheckSuggestion:
    """检查建议数据模型"""
    id: str = ""
    rule_id: str = ""
    rule_name: str = ""
    type: str = ""
    severity: str = "Medium"
    section: str = ""
    original_text: str = ""
    suggestion: str = ""
    reason: str = ""

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'CheckSuggestion':
        """从字典创建CheckSuggestion"""
        return cls(
            id=data.get('id', ''),
            rule_id=data.get('rule_id', ''),
            rule_name=data.get('rule_name', ''),
            type=data.get('type', ''),
            severity=data.get('severity', 'Medium'),
            section=data.get('section', ''),
            original_text=data.get('original_text', ''),
            suggestion=data.get('suggestion', ''),
            reason=data.get('reason', '')
        )


class Severity(Enum):
    """建议严重程度"""
    CRITICAL = "Critical"
    HIGH = "High"
    MEDIUM = "Medium"
    LOW = "Low"


# ============================================================
# 文档读取器
# ============================================================

class DocumentReader:
    """文档读取器 - 支持docx/wps/wpsx"""

    @staticmethod
    def read_document(file_path: str) -> Dict[str, Any]:
        """
        读取文档，返回文本内容和元数据
        
        Args:
            file_path: 文档路径
            
        Returns:
            Dict包含document_content和document_metadata
        """
        logger.info(f"开始读取文档: {file_path}")

        if not file_path or not file_path.strip():
            raise ValueError("文件路径不能为空")

        file_path = file_path.strip()

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = Path(file_path).suffix.lower()

        if ext == '.docx':
            return DocumentReader._read_docx(file_path)
        elif ext in ['.wps', '.wpsx']:
            return DocumentReader._read_wps(file_path)
        else:
            raise ValueError(f"不支持的文档格式: {ext}，仅支持 .docx, .wps, .wpsx")

    @staticmethod
    def _read_docx(file_path: str) -> Dict[str, Any]:
        """读取Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装，请运行: pip install python-docx")

        logger.info("读取Word文档 (.docx)")

        doc = DocxDocument(file_path)
        
        # 提取段落文本
        paragraphs = []
        full_text_lines = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({
                    'index': i,
                    'text': text,
                    'style': para.style.name if para.style else 'Normal'
                })
                full_text_lines.append(text)

        # 提取表格内容
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.rows[0].cells) if table.rows else 0,
                'data': table_data
            })

        # 提取元数据
        metadata = DocumentReader._extract_docx_metadata(doc, file_path)

        # 构建完整文本
        full_text = '\n'.join(full_text_lines)
        
        # 添加表格文本表示
        if tables:
            full_text += '\n\n【表格内容】\n'
            for table in tables:
                full_text += f'\n表{table["index"] + 1}:\n'
                for row_data in table['data']:
                    full_text += '| ' + ' | '.join(row_data) + ' |\n'

        logger.info(f"读取完成: {len(paragraphs)}个段落, {len(tables)}个表格")

        return {
            'success': True,
            'document_content': full_text,
            'document_metadata': {
                'title': metadata.title,
                'author': metadata.author,
                'created_date': metadata.created_date,
                'modified_date': metadata.modified_date,
                'file_type': metadata.file_type,
                'file_size': metadata.file_size,
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'page_count': metadata.page_count,
                'company': metadata.company
            },
            'structured_content': {
                'paragraphs': paragraphs,
                'tables': tables
            },
            'file_path': file_path
        }

    @staticmethod
    def _read_wps(file_path: str) -> Dict[str, Any]:
        """读取WPS文档（使用OOXML格式解析）"""
        logger.info("读取WPS文档 (.wps/.wpsx)")

        # WPS文档实际上是OOXML格式，可以像docx一样解析
        try:
            return DocumentReader._read_docx(file_path)
        except Exception as e:
            logger.warning(f"使用docx解析WPS失败，尝试使用zip解析: {e}")
            return DocumentReader._read_wps_via_zip(file_path)

    @staticmethod
    def _read_wps_via_zip(file_path: str) -> Dict[str, Any]:
        """通过ZIP方式读取WPS文档"""
        paragraphs = []
        tables = []
        
        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                # 读取document.xml
                if 'word/document.xml' in zf.namelist():
                    doc_xml = zf.read('word/document.xml').decode('utf-8')
                    # 简单提取文本（去除XML标签）
                    import re
                    text_content = re.sub(r'<[^>]+>', '', doc_xml)
                    paragraphs = [{'index': i, 'text': line.strip(), 'style': 'Normal'} 
                                  for i, line in enumerate(text_content.split('\n')) if line.strip()]
        except Exception as e:
            logger.error(f"读取WPS文档失败: {e}")
            raise

        full_text = '\n'.join([p['text'] for p in paragraphs])
        
        # 获取文件信息
        file_stat = os.stat(file_path)
        metadata = DocumentMetadata(
            title=Path(file_path).stem,
            author='',
            created_date=datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            modified_date=datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
            file_type='wps',
            file_size=file_stat.st_size,
            paragraph_count=len(paragraphs),
            table_count=len(tables)
        )

        logger.info(f"WPS读取完成: {len(paragraphs)}个段落")

        return {
            'success': True,
            'document_content': full_text,
            'document_metadata': {
                'title': metadata.title,
                'author': metadata.author,
                'created_date': metadata.created_date,
                'modified_date': metadata.modified_date,
                'file_type': metadata.file_type,
                'file_size': metadata.file_size,
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'page_count': metadata.page_count,
                'company': metadata.company
            },
            'structured_content': {
                'paragraphs': paragraphs,
                'tables': tables
            },
            'file_path': file_path
        }

    @staticmethod
    def _extract_docx_metadata(doc, file_path: str) -> DocumentMetadata:
        """提取docx文档元数据"""
        file_stat = os.stat(file_path)
        
        # 从core_properties获取元数据
        try:
            core_props = doc.core_properties
            title = core_props.title or Path(file_path).stem
            author = core_props.author or ''
            created = core_props.created.isoformat() if core_props.created else ''
            modified = core_props.modified.isoformat() if core_props.modified else ''
        except Exception:
            title = Path(file_path).stem
            author = ''
            created = datetime.fromtimestamp(file_stat.st_ctime).isoformat()
            modified = datetime.fromtimestamp(file_stat.st_mtime).isoformat()

        return DocumentMetadata(
            title=title,
            author=author,
            created_date=created,
            modified_date=modified,
            file_type='docx',
            file_size=file_stat.st_size,
            paragraph_count=len(doc.paragraphs),
            table_count=len(doc.tables)
        )


# ============================================================
# 系统信息读取器（从Excel读取系统基础信息）
# ============================================================

class SystemInfoReader:
    """系统信息读取器 - 从Excel文件读取系统基础信息（正式名称、开发部门、运维部门等）"""

    # 期望的Excel列标题映射（支持模糊匹配）
    EXPECTED_HEADER_MAP = {
        "系统中文名称": "system_name_cn",
        "系统英文简称": "system_name_en",
        "模块中文名称": "module_name_cn",
        "模块英文简称": "module_name_en",
        "研发牵头部门": "dev_department",
        "运维牵头部门": "ops_department"
    }

    @staticmethod
    def read_system_info(file_path: str) -> Dict[str, Any]:
        """
        从Excel文件中读取系统基础信息

        Args:
            file_path: Excel文件路径（.xlsx格式）

        Returns:
            Dict包含systems（按系统分组的完整信息列表）和raw_rows（原始行数据）
        """
        logger.info(f"开始读取系统信息表: {file_path}")

        if not file_path or not file_path.strip():
            return {
                "success": False,
                "error": "file_path不能为空",
                "error_type": "ValidationError"
            }

        file_path = file_path.strip()

        if not os.path.exists(file_path):
            return {
                "success": False,
                "error": f"文件不存在: {file_path}",
                "error_type": "FileNotFoundError"
            }

        ext = Path(file_path).suffix.lower()
        if ext not in ['.xlsx', '.xls']:
            return {
                "success": False,
                "error": f"不支持的文件格式: {ext}，仅支持 .xlsx、.xls",
                "error_type": "UnsupportedFormatError"
            }

        if not OPENPYXL_AVAILABLE:
            return {
                "success": False,
                "error": "openpyxl库未安装，请运行: pip install openpyxl",
                "error_type": "ImportError"
            }

        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            ws = wb.active
            if ws is None:
                return {
                    "success": False,
                    "error": "Excel文件中没有工作表",
                    "error_type": "EmptyWorkbookError"
                }

            # 读取所有行
            rows = list(ws.iter_rows(values_only=True))
            wb.close()

            if not rows:
                return {
                    "success": False,
                    "error": "Excel文件为空",
                    "error_type": "EmptyFileError"
                }

            # 解析表头
            header_row = rows[0]
            header_mapping = SystemInfoReader._map_headers(header_row)

            logger.info(f"表头映射: {header_mapping}")

            if not header_mapping:
                return {
                    "success": False,
                    "error": f"无法识别Excel表头，期望的列包括: {', '.join(SystemInfoReader.EXPECTED_HEADER_MAP.keys())}",
                    "error_type": "HeaderParseError"
                }

            # 解析数据行
            raw_rows = []
            systems_map = {}  # key: 系统英文简称, value: 系统信息

            for row_idx, row in enumerate(rows[1:], start=2):
                if all(cell is None or (isinstance(cell, str) and cell.strip() == '') for cell in row):
                    continue  # 跳过空行

                row_dict = {}
                for col_idx, (field_name, field_key) in header_mapping.items():
                    value = row[col_idx] if col_idx < len(row) else None
                    if value is not None:
                        value = str(value).strip()
                    row_dict[field_key] = value or ""

                raw_rows.append(row_dict)

                # 按系统英文简称聚合
                sys_en = row_dict.get("system_name_en", "").strip()
                sys_cn = row_dict.get("system_name_cn", "").strip()
                dev_dept = row_dict.get("dev_department", "").strip()
                ops_dept = row_dict.get("ops_department", "").strip()

                if sys_en and sys_cn:
                    if sys_en not in systems_map:
                        systems_map[sys_en] = {
                            "system_name_cn": sys_cn,
                            "system_name_en": sys_en,
                            "dev_department": dev_dept,
                            "ops_department": ops_dept,
                            "modules": []
                        }
                    # 如果存在模块信息，记录模块
                    module_cn = row_dict.get("module_name_cn", "").strip()
                    module_en = row_dict.get("module_name_en", "").strip()
                    if module_cn:
                        systems_map[sys_en]["modules"].append({
                            "module_name_cn": module_cn,
                            "module_name_en": module_en
                        })

            # 构建按系统分组的概要信息
            systems_list = []
            for sys_en, sys_info in systems_map.items():
                systems_list.append(sys_info)

            # 构建系统名称索引（便于LLM检索）
            system_name_index = {}
            for sys_info in systems_list:
                # 用中文名和英文名作为key
                system_name_index[sys_info["system_name_cn"]] = sys_info
                system_name_index[sys_info["system_name_en"]] = sys_info

            result = {
                "success": True,
                "file_path": file_path,
                "total_systems": len(systems_list),
                "total_records": len(raw_rows),
                "systems": systems_list,
                # 记录数超过100时省略raw_rows（避免超大JSON）
                "raw_rows": raw_rows if len(raw_rows) <= 100 else f"省略（共{len(raw_rows)}条，请通过systems查看聚合数据）",
                "system_name_index": {
                    "keys": list(system_name_index.keys()),
                    "description": "系统名称索引，可用于根据文档中出现的系统名称快速查找对应系统信息"
                }
            }

            logger.info(f"系统信息读取完成: {len(systems_list)}个系统, {len(raw_rows)}条记录")
            return result

        except Exception as e:
            logger.exception(f"读取系统信息表失败: {e}")
            return {
                "success": False,
                "error": f"读取系统信息表失败: {str(e)}",
                "error_type": type(e).__name__
            }

    @staticmethod
    def _map_headers(header_row: tuple) -> Dict[int, str]:
        """
        将Excel表头映射到内部字段名

        Args:
            header_row: Excel的第一行（表头）

        Returns:
            Dict: {列索引: (原始表头, 内部字段名)}
        """
        import re

        mapping = {}

        # 构建模糊匹配映射：中文表头（去除非中文字符）→ 字段名
        normalized_expected = {}
        # 补充英文缩写的别名映射：英文名 → 字段名
        alias_map = {
            "系统中文名称": "system_name_cn",
            "系统英文简称": "system_name_en",
            "模块中文名称": "module_name_cn",
            "模块英文简称": "module_name_en",
            "研发牵头部门": "dev_department",
            "运维牵头部门": "ops_department"
        }
        # 常见同义变体映射
        synonym_map = {
            "系统中文名": "系统中文名称",
            "系统英文名": "系统英文简称",
            "模块中文名": "模块中文名称",
            "模块英文名": "模块英文简称",
            "研发部门": "研发牵头部门",
            "运维部门": "运维牵头部门",
            "研发牵头": "研发牵头部门",
            "运维牵头": "运维牵头部门"
        }

        for cn_name, field_key in SystemInfoReader.EXPECTED_HEADER_MAP.items():
            cn_clean = re.sub(r'[^\u4e00-\u9fff]', '', cn_name)
            normalized_expected[cn_clean] = field_key

        for col_idx, header_value in enumerate(header_row):
            if header_value is None:
                continue
            header_str = str(header_value).strip()
            if not header_str:
                continue

            # 1. 精确匹配（原始字符串）
            if header_str in SystemInfoReader.EXPECTED_HEADER_MAP:
                mapping[col_idx] = (header_str, SystemInfoReader.EXPECTED_HEADER_MAP[header_str])
                continue

            # 2. 同义变体匹配
            if header_str in synonym_map:
                target = synonym_map[header_str]
                if target in SystemInfoReader.EXPECTED_HEADER_MAP:
                    mapping[col_idx] = (header_str, SystemInfoReader.EXPECTED_HEADER_MAP[target])
                    continue

            # 3. 模糊匹配（去除非中文字符后匹配）
            header_clean = re.sub(r'[^\u4e00-\u9fff]', '', header_str)

            # 跳过空字符串（纯英文/数字表头无法通过中文模糊匹配）
            if not header_clean:
                continue

            if header_clean in normalized_expected:
                mapping[col_idx] = (header_str, normalized_expected[header_clean])
                continue

            # 4. 部分匹配（包含关系）——要求最小匹配长度≥3个中文字符
            #    避免"名称"误匹配"系统中文名称"、"缩写"误匹配"英文简称"等
            for cn_clean, field_key in normalized_expected.items():
                min_match_len = 3
                if len(cn_clean) < min_match_len or len(header_clean) < min_match_len:
                    continue
                if cn_clean in header_clean or header_clean in cn_clean:
                    mapping[col_idx] = (header_str, field_key)
                    break

        return mapping


# ============================================================
# 文档修订器
# ============================================================

class DocumentReviser:
    """文档修订器 - 支持修订模式"""

    @staticmethod
    def clean_json_string(json_str: str) -> str:
        """
        清理JSON字符串，去除Markdown代码块标记
        
        用于处理LLM输出的带Markdown标记的JSON字符串
        例如：```json\n{...}\n``` → {...}
        
        Args:
            json_str: 可能包含Markdown标记的JSON字符串
            
        Returns:
            纯净的JSON字符串
        """
        if not json_str:
            return json_str
        
        # 去除开头的Markdown标记
        json_str = re.sub(r'^\s*```json\s*', '', json_str.strip(), flags=re.IGNORECASE)
        json_str = re.sub(r'^\s*```\s*', '', json_str.strip())
        
        # 去除结尾的Markdown标记
        json_str = re.sub(r'\s*```\s*$', '', json_str)
        
        return json_str.strip()

    @staticmethod
    def revise_document(file_path: str, suggestions_json: str, output_path: str = "", use_track_changes: str = "comment") -> Dict[str, Any]:
        """
        根据检查建议修订文档
        
        Args:
            file_path: 原文档路径
            suggestions_json: 检查建议JSON字符串
            output_path: 输出路径（可选）
            use_track_changes: 是否使用修订模式
            
        Returns:
            修订结果字典
        """
        logger.info(f"开始修订文档: {file_path}")
        
        if not file_path or not suggestions_json:
            return {
                'success': False,
                'error': '缺少必需参数',
                'error_type': 'ValidationError'
            }
        
        if not os.path.exists(file_path):
            return {
                'success': False,
                'error': f'文件不存在: {file_path}',
                'error_type': 'FileNotFoundError'
            }
        
        # 解析建议JSON（自动清理Markdown标记）
        try:
            # 清理可能的Markdown代码块标记
            clean_json = DocumentReviser.clean_json_string(suggestions_json)
            suggestions_data = json.loads(clean_json)
            if isinstance(suggestions_data, dict):
                suggestions = suggestions_data.get('suggestions', [])
            elif isinstance(suggestions_data, list):
                suggestions = suggestions_data
            else:
                suggestions = []
        except json.JSONDecodeError as e:
            return {
                'success': False,
                'error': f'JSON解析错误: {str(e)}',
                'error_type': 'JSONParseError'
            }
        
        if not suggestions:
            return {
                'success': False,
                'error': '没有可应用的修订建议',
                'error_type': 'NoSuggestionsError'
            }
        
        # 确定输出路径
        if not output_path:
            file_name = Path(file_path).stem
            file_ext = Path(file_path).suffix
            output_path = f"{file_name}_revised{file_ext}"
        
        # 执行修订
        # 注意：use_track_changes 可以是 'true', 'comment', 'false' 三种字符串
        
        try:
            ext = Path(file_path).suffix.lower()
            if ext == '.docx':
                result = DocumentReviser._revise_docx(file_path, suggestions, output_path, use_track_changes)
            elif ext in ['.wps', '.wpsx']:
                result = DocumentReviser._revise_wps(file_path, suggestions, output_path, use_track_changes)
            else:
                return {
                    'success': False,
                    'error': f'不支持的文档格式: {ext}',
                    'error_type': 'UnsupportedFormatError'
                }
            
            # 生成修改记录明细文档
            if result.get('success') and result.get('applied_revisions', 0) > 0:
                log_path = DocumentReviser._write_revision_log(file_path, result.get('revisions_detail', []), suggestions)
                if log_path:
                    result['revision_log_path'] = log_path
                    logger.info(f"修改记录文档已生成: {log_path}")
            
            return result
            
        except Exception as e:
            logger.exception("修订文档失败")
            return {
                'success': False,
                'error': f'修订失败: {str(e)}',
                'error_type': type(e).__name__
            }

    @staticmethod
    def _create_track_changes_element(element_type: str, text: str, author: str = "DocumentReviser", date: str = None) -> OxmlElement:
        """
        创建Word修订模式的XML元素

        Args:
            element_type: 'del' 或 'ins'
            text: 文本内容
            author: 修订作者
            date: 修订日期（ISO格式）

        Returns:
            OxmlElement: 修订元素
        """
        if date is None:
            date = datetime.now().astimezone().strftime("%Y-%m-%dT%H:%M:%SZ")

        # 创建修订元素
        if element_type == 'del':
            element = OxmlElement('w:del')
            element.set(qn('w:id'), str(hash(text) % 1000000))
            element.set(qn('w:author'), author)
            element.set(qn('w:date'), date)
        else:  # ins
            element = OxmlElement('w:ins')
            element.set(qn('w:id'), str(hash(text) % 1000000 + 1))
            element.set(qn('w:author'), author)
            element.set(qn('w:date'), date)

        # 创建 run 元素包含文本
        run = OxmlElement('w:r')

        # 创建文本元素
        text_elem = OxmlElement('w:t')
        text_elem.text = text

        run.append(text_elem)
        element.append(run)

        return element

    @staticmethod
    def _apply_track_changes_to_paragraph(para, original: str, new_text: str, revision_id: int) -> bool:
        """
        在段落中应用真正的Word修订模式（使用w:del和w:ins元素）

        Args:
            para: 段落对象
            original: 原文本
            new_text: 新文本
            revision_id: 修订ID

        Returns:
            bool: 是否成功应用
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # 获取段落的XML元素
        p = para._p

        # 查找包含原文的run
        found = False
        date_str = datetime.now().astimezone().strftime("%Y-%m-%dT%H:%M:%SZ")

        # 遍历所有run
        for run in para.runs:
            if original in run.text:
                # 找到包含原文的run
                r = run._r

                # 创建删除元素 (w:del)
                del_elem = OxmlElement('w:del')
                del_elem.set(qn('w:id'), str(revision_id))
                del_elem.set(qn('w:author'), 'DocumentReviser')
                del_elem.set(qn('w:date'), date_str)

                # 在删除元素中创建run包含原文
                del_run = OxmlElement('w:r')
                del_text = OxmlElement('w:t')
                del_text.text = original
                del_run.append(del_text)
                del_elem.append(del_run)

                # 创建插入元素 (w:ins)
                ins_elem = OxmlElement('w:ins')
                ins_elem.set(qn('w:id'), str(revision_id + 1))
                ins_elem.set(qn('w:author'), 'DocumentReviser')
                ins_elem.set(qn('w:date'), date_str)

                # 在插入元素中创建run包含新文本
                ins_run = OxmlElement('w:r')
                ins_text = OxmlElement('w:t')
                ins_text.text = new_text
                ins_run.append(ins_text)
                ins_elem.append(ins_run)

                # 替换原文run
                # 先添加删除和插入元素到段落
                p.append(del_elem)
                p.append(ins_elem)

                # 删除原文run
                p.remove(r)

                found = True
                break

        return found

    @staticmethod
    def _revise_docx(file_path: str, suggestions: List[Dict], output_path: str, use_track_changes: str) -> Dict[str, Any]:
        """修订Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装")

        doc = DocxDocument(file_path)

        applied_count = 0
        revisions_detail = []
        revision_id = 1000  # 起始修订ID

        # 按顺序应用建议
        for suggestion in suggestions:
            original = suggestion.get('original_text', '')
            new_text = suggestion.get('suggestion', '')

            if not original:
                continue

            # 在文档中查找并替换
            found = False
            for para in doc.paragraphs:
                if original in para.text:
                    if use_track_changes == 'true':
                        # 使用真正的Word修订模式
                        success = DocumentReviser._apply_track_changes_to_paragraph(
                            para, original, new_text, revision_id
                        )
                        if success:
                            found = True
                            applied_count += 1
                            revision_id += 2  # 每次修订使用两个ID（del和ins）
                    elif use_track_changes == 'comment':
                        # 在原文后插入红色括号说明（不修改原文）
                        success = DocumentReviser._add_red_comment_to_paragraph(
                            para, original, new_text, suggestion.get('reason', ''), 
                            suggestion.get('rule_id', ''), revision_id
                        )
                        if success:
                            found = True
                            applied_count += 1
                            revision_id += 1
                    else:
                        # 直接替换模式
                        para.text = para.text.replace(original, new_text)
                        found = True
                        applied_count += 1
                    break

            revisions_detail.append({
                'id': suggestion.get('id', ''),
                'rule_id': suggestion.get('rule_id', ''),
                'original': original,
                'new': new_text,
                'status': 'success' if found else 'not_found'
            })

        # 保存文档
        doc.save(output_path)

        return {
            'success': True,
            'output_path': output_path,
            'use_track_changes': use_track_changes,
            'total_suggestions': len(suggestions),
            'applied_revisions': applied_count,
            'revisions_detail': revisions_detail
        }

    @staticmethod
    def _add_red_comment_to_paragraph(para, original: str, new_text: str, reason: str, 
                                       rule_id: str, comment_id: int) -> bool:
        """
        在段落中原文后插入红色括号说明（不修改原文）
        通过XML操作在原文run后插入红色说明run，保留原文所有格式
        
        格式：（【建议】xxx 【原因】xxx）
        样式：红色字体，保留原文run的字体大小和字体名称
        
        Args:
            para: 段落对象
            original: 原文（保留不修改）
            new_text: 修改建议
            reason: 修改原因
            rule_id: 规则ID
            comment_id: 批注ID（用于标识）
            
        Returns:
            bool: 是否成功添加
        """
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            import copy
            
            # 查找原文位置
            if original not in para.text:
                return False
            
            # 构建说明文本
            comment_text = f"（【建议】{new_text} 【原因】{reason}）"
            
            # 遍历runs，找到包含原文的run
            for run in para.runs:
                if original in run.text:
                    # 创建新的r元素用于红色说明
                    new_r = OxmlElement('w:r')
                    
                    # 复制原run的字体属性（rPr），仅修改颜色为红色
                    # 这样红色说明的字体大小、字体名称等与原run一致
                    src_rpr = run._r.find(qn('w:rPr'))
                    if src_rpr is not None:
                        new_rpr = copy.deepcopy(src_rpr)
                        # 设置或覆盖颜色为红色
                        color = new_rpr.find(qn('w:color'))
                        if color is not None:
                            color.set(qn('w:val'), 'FF0000')
                        else:
                            color_elem = OxmlElement('w:color')
                            color_elem.set(qn('w:val'), 'FF0000')
                            new_rpr.append(color_elem)
                        new_r.append(new_rpr)
                    else:
                        # 原run无字体属性，创建仅含红色的rPr
                        new_rpr = OxmlElement('w:rPr')
                        color_elem = OxmlElement('w:color')
                        color_elem.set(qn('w:val'), 'FF0000')
                        new_rpr.append(color_elem)
                        new_r.append(new_rpr)
                    
                    # 创建文本元素
                    new_t = OxmlElement('w:t')
                    new_t.text = comment_text
                    new_r.append(new_t)
                    
                    # 在原run后面插入新r元素（保留所有原有run不变）
                    run._r.addnext(new_r)
                    
                    return True
            
            return False
            
        except Exception as e:
            logger.error(f"添加红色说明失败: {e}")
            import traceback
            traceback.print_exc()
            return False

    @staticmethod
    def _revise_wps(file_path: str, suggestions: List[Dict], output_path: str, use_track_changes: str) -> Dict[str, Any]:
        """修订WPS文档（复用docx逻辑）"""
        return DocumentReviser._revise_docx(file_path, suggestions, output_path, use_track_changes)

    @staticmethod
    def _get_revision_log_path(file_path: str) -> str:
        """
        获取修改记录文档路径
        与原文档同目录，文件名为：原文件名_revision_log.docx
        """
        file_dir = Path(file_path).parent
        file_stem = Path(file_path).stem
        return str(file_dir / f"{file_stem}_revision_log.docx")

    @staticmethod
    def _get_existing_revision_count(log_path: str) -> int:
        """
        读取已有修改记录文档，获取已存在的修改次数
        
        Args:
            log_path: 修改记录文档路径
            
        Returns:
            已有修改次数（行数）
        """
        if not os.path.exists(log_path):
            return 0
        
        try:
            doc = DocxDocument(log_path)
            count = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                # 统计"第N次修改"的标题出现的次数
                if text.startswith("第") and "次修改" in text:
                    count += 1
            return count
        except Exception:
            return 0

    @staticmethod
    def _write_revision_log(revised_file_path: str, revisions_detail: List[Dict], 
                             suggestions: List[Dict]) -> str:
        """
        生成或追加修改记录明细文档
        
        创建一个本地DOCX文档记录对文件的修改明细，内容包括：
        - 第几次修改
        - 修改时间
        - 修改概述（修改了哪些内容类型、严重程度分布等）
        - 修改内容明细（每条修改的原文、修改后内容、原因）
        
        Args:
            revised_file_path: 修订后的文档路径
            revisions_detail: 修订详情列表（来自修订结果）
            suggestions: 原始建议列表
            
        Returns:
            修改记录文档路径
        """
        if not DOCX_AVAILABLE:
            logger.warning("python-docx未安装，无法生成修改记录文档")
            return ""
        
        now = datetime.now()
        log_path = DocumentReviser._get_revision_log_path(revised_file_path)
        
        # 检查是否已有修改记录文档，决定是创建还是追加
        existing_count = DocumentReviser._get_existing_revision_count(log_path)
        revision_number = existing_count + 1
        
        # 生成修改概述
        by_severity = {}
        by_type = {}
        for sug in suggestions:
            sev = sug.get('severity', 'Unknown')
            by_severity[sev] = by_severity.get(sev, 0) + 1
            tp = sug.get('type', 'Unknown')
            by_type[tp] = by_type.get(tp, 0) + 1
        
        severity_summary = "、".join([f"{k}: {v}条" for k, v in sorted(by_severity.items())])
        type_summary = "、".join([f"{k}: {v}条" for k, v in sorted(by_type.items())])
        total = len(revisions_detail)
        success_count = sum(1 for r in revisions_detail if r.get('status') == 'success')
        
        summary_text = f"共修订{total}条内容，成功应用{success_count}条；严重程度分布：{severity_summary}；类型分布：{type_summary}"
        
        # 打开或创建文档
        if os.path.exists(log_path):
            doc = DocxDocument(log_path)
        else:
            doc = DocxDocument()
            # 添加文档标题
            para = doc.add_paragraph()
            run = para.add_run(f"文件修改记录明细")
            run.bold = True
            run.font.size = Pt(16)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加被修改文件信息
            doc.add_paragraph(f"被修改文件：{Path(revised_file_path).name}")
            doc.add_paragraph("")
        
        # 分隔线（追加时区分新旧记录）
        if existing_count > 0:
            p = doc.add_paragraph()
            run = p.add_run("─" * 50)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 第几次修改标题
        p = doc.add_paragraph()
        run = p.add_run(f"第{revision_number}次修改")
        run.bold = True
        run.font.size = Pt(14)
        
        # 修改时间
        doc.add_paragraph(f"修改时间：{now.strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 修改概述
        p = doc.add_paragraph()
        run = p.add_run("修改概述：")
        run.bold = True
        doc.add_paragraph(summary_text)
        doc.add_paragraph("")
        
        # 修改内容明细
        p = doc.add_paragraph()
        run = p.add_run("修改内容明细：")
        run.bold = True
        doc.add_paragraph("")
        
        for i, (rev, sug) in enumerate(zip(revisions_detail, suggestions)):
            seq = i + 1
            status_icon = "✅" if rev.get('status') == 'success' else "❌"
            
            # 序号和状态
            p = doc.add_paragraph()
            run = p.add_run(f"{seq}. [{status_icon}] {rev.get('id', '')}")
            run.bold = True
            
            # 规则信息
            rule_name = sug.get('rule_name', '')
            if rule_name:
                doc.add_paragraph(f"   规则：{rule_name}（{sug.get('rule_id', '')}）")
            
            doc.add_paragraph(f"   类型：{sug.get('type', '')}　严重程度：{sug.get('severity', '')}")
            doc.add_paragraph(f"   状态：{'已应用' if rev.get('status') == 'success' else '未找到原文'}")
            doc.add_paragraph(f"   原文：{rev.get('original', '')}")
            
            # 修改后内容
            p = doc.add_paragraph()
            run = p.add_run(f"   修改后：{rev.get('new', '')}")
            run.font.color.rgb = RGBColor(0, 100, 0)  # 深绿色
            
            # 修改原因
            reason = sug.get('reason', '')
            if reason:
                p = doc.add_paragraph()
                run = p.add_run(f"   原因：{reason}")
                run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
            
            doc.add_paragraph("")  # 空行分隔
        
        # 保存
        doc.save(log_path)
        logger.info(f"修改记录文档已保存: {log_path}, 第{revision_number}次修改, 共{total}条")
        
        return log_path


# ============================================================
# MCP服务器初始化 - 导入共享mcp实例，注册工具
# ============================================================

if FASTMCP_AVAILABLE:

    # ============================================================
    # 第1步 - 文档读取工具
    # ============================================================

    @mcp.tool()
    def read_document(file_path: str) -> str:
        """
        读取文档内容，返回文档文本和元数据（第1步）
        
        Args:
            file_path: 文档路径，支持 .docx、.wps、.wpsx 格式
            
        Returns:
            JSON字符串，包含document_content和document_metadata
        """
        start_time = datetime.now()
        call_id = f"read_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始读取文档: {file_path}")

        try:
            if not file_path or not file_path.strip():
                return json.dumps({
                    "success": False,
                    "error": "file_path不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id
                }, ensure_ascii=False)

            reader = DocumentReader()
            result = reader.read_document(file_path)
            result['call_id'] = call_id
            result['timestamp'] = datetime.now().isoformat()

            logger.info(f"[{call_id}] 读取完成: {result['document_metadata']['paragraph_count']}个段落")
            return json.dumps(result, ensure_ascii=False, indent=2)

        except Exception as e:
            logger.exception(f"[{call_id}] 读取异常")
            return json.dumps({
                "success": False,
                "error": f"读取失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)

    @mcp.tool()
    def get_document_info(file_path: str) -> str:
        """
        获取文档基本信息（第1步辅助工具）
        
        Args:
            file_path: 文档路径
            
        Returns:
            JSON字符串，包含文档基本信息
        """
        start_time = datetime.now()
        call_id = f"info_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 获取文档信息: {file_path}")

        try:
            if not file_path or not file_path.strip():
                return json.dumps({
                    "success": False,
                    "error": "file_path不能为空",
                    "call_id": call_id
                }, ensure_ascii=False)

            file_path = file_path.strip()

            if not os.path.exists(file_path):
                return json.dumps({
                    "success": False,
                    "error": f"文件不存在: {file_path}",
                    "call_id": call_id
                }, ensure_ascii=False)

            file_stat = os.stat(file_path)
            file_type = Path(file_path).suffix.lower()

            result = {
                "success": True,
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "file_type": file_type,
                "file_size": file_stat.st_size,
                "file_size_readable": f"{file_stat.st_size / 1024:.2f} KB",
                "created_time": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }

            return json.dumps(result, ensure_ascii=False, indent=2)

        except Exception as e:
            logger.exception(f"[{call_id}] 获取信息异常")
            return json.dumps({
                "success": False,
                "error": f"获取信息失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id
            }, ensure_ascii=False)

    # ============================================================
    # 新增第1.5步 - 系统信息读取工具
    # ============================================================

    @mcp.tool()
    def read_system_info(file_path: str) -> str:
        """
        从Excel系统信息表中读取系统基础信息（第1.5步 - 新增节点）

        从预先维护的系统信息Excel文件中读取各系统的正式中文名称、英文简称、
        研发牵头部门、运维牵头部门等信息，返回结构化JSON供后续LLM节点使用。

        Args:
            file_path: 系统信息Excel文件路径（.xlsx格式）

        Returns:
            JSON字符串，包含systems（按系统分组的信息列表）和raw_rows（原始行数据）

        输出结构示例:
        {
            "success": true,
            "total_systems": 6,
            "total_records": 20,
            "systems": [
                {
                    "system_name_cn": "客户关系管理系统",
                    "system_name_en": "CRM",
                    "dev_department": "产品研发部",
                    "ops_department": "基础运维部",
                    "modules": [...]
                }
            ],
            "raw_rows": [...],
            "system_name_index": {
                "keys": ["客户关系管理系统", "CRM", ...]
            }
        }
        """
        start_time = datetime.now()
        call_id = f"sysinfo_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始读取系统信息表: {file_path}")

        try:
            if not file_path or not file_path.strip():
                return json.dumps({
                    "success": False,
                    "error": "file_path不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id
                }, ensure_ascii=False)

            reader = SystemInfoReader()
            result = reader.read_system_info(file_path)
            result['call_id'] = call_id
            result['timestamp'] = datetime.now().isoformat()

            logger.info(f"[{call_id}] 系统信息读取完成: {result.get('total_systems', 0)}个系统, {result.get('total_records', 0)}条记录")
            # 注意：不传 indent=2 以减小大文件场景下的JSON体积
            return json.dumps(result, ensure_ascii=False)

        except Exception as e:
            logger.exception(f"[{call_id}] 读取系统信息异常")
            return json.dumps({
                "success": False,
                "error": f"读取失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)

    @mcp.tool()
    def match_system_from_document(document_text: str, system_info_json: str) -> str:
        """
        从文档内容中识别涉及的系统，并匹配系统信息表中的对应记录（第1.5步 - 辅助工具）

        该工具遍历文档文本中出现的系统名称（包括中英文名称），与系统信息表中的记录进行匹配，
        返回文档中所涉及系统的详细信息。配合LLM分析节点使用，可辅助LLM更快完成系统匹配。

        Args:
            document_text: 文档内容文本（从第1步read_document输出中获得）
            system_info_json: 系统信息表JSON字符串（从read_system_info输出中获得）

        Returns:
            JSON字符串，包含matched_systems（匹配到的系统列表）
        """
        start_time = datetime.now()
        call_id = f"matchsys_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始从文档中匹配系统信息")

        try:
            if not document_text or not document_text.strip():
                return json.dumps({
                    "success": False,
                    "error": "document_text不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id
                }, ensure_ascii=False)

            if not system_info_json or not system_info_json.strip():
                return json.dumps({
                    "success": False,
                    "error": "system_info_json不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id
                }, ensure_ascii=False)

            # 解析系统信息
            try:
                sys_info = json.loads(system_info_json)
            except json.JSONDecodeError as e:
                return json.dumps({
                    "success": False,
                    "error": f"system_info_json解析失败: {str(e)}",
                    "error_type": "JSONParseError",
                    "call_id": call_id
                }, ensure_ascii=False)

            if not sys_info.get("success"):
                return json.dumps({
                    "success": False,
                    "error": f"系统信息读取失败: {sys_info.get('error', '未知错误')}",
                    "error_type": "SystemInfoError",
                    "call_id": call_id
                }, ensure_ascii=False)

            systems = sys_info.get("systems", [])

            # ----------------------------------------------------------------
            # 收集所有系统名称关键词（中文名、英文名），排除通用模块名
            # ----------------------------------------------------------------
            # 如果模块名在超过3个系统中出现，视为通用名，不作为匹配关键词
            # 这避免"基础管理"、"数据管理"等通用模块名导致全部系统误匹配
            #
            # 分两轮构建：
            #   第1轮：统计每个模块名出现的系统数
            #   第2轮：只将"稀有模块名"（出现系统数<=阈值）作为关键词
            MODULE_DUPLICATE_THRESHOLD = 3

            # 第1轮：统计模块名频率
            module_freq = {}
            for sys_data in systems:
                seen_mods = set()
                for module in sys_data.get("modules", []):
                    module_cn = module.get("module_name_cn", "").strip()
                    if module_cn and module_cn not in seen_mods:
                        seen_mods.add(module_cn)
                        module_freq[module_cn] = module_freq.get(module_cn, 0) + 1

            # 第2轮：构建关键词 → 系统映射
            system_keywords = {}
            for sys_data in systems:
                cn_name = sys_data.get("system_name_cn", "")
                en_name = sys_data.get("system_name_en", "")
                keywords = set()

                # 中文名全称（高优先级）
                if cn_name:
                    keywords.add(cn_name)
                    # 生成可能的简称（如"客户关系管理系统"→"客户关系管理"）
                    if cn_name.endswith("系统"):
                        keywords.add(cn_name[:-2])
                    if cn_name.endswith("平台"):
                        keywords.add(cn_name[:-2])
                    if cn_name.endswith("中台"):
                        keywords.add(cn_name[:-2])

                # 英文名称（高优先级）
                if en_name:
                    keywords.add(en_name)

                # 模块名称——只使用稀有模块名（出现系统数<=阈值）
                for module in sys_data.get("modules", []):
                    module_cn = module.get("module_name_cn", "").strip()
                    module_en = module.get("module_name_en", "").strip()
                    module_freq_count = module_freq.get(module_cn, 999)
                    if module_cn and module_freq_count <= MODULE_DUPLICATE_THRESHOLD:
                        keywords.add(module_cn)
                    if module_en and module_freq_count <= MODULE_DUPLICATE_THRESHOLD:
                        keywords.add(module_en)

                # 存储关键词到系统的映射
                for kw in keywords:
                    if kw:
                        if kw not in system_keywords:
                            system_keywords[kw] = []
                        system_keywords[kw].append(sys_data)

            # ----------------------------------------------------------------
            # 在文档中搜索关键词（优先匹配完整系统名，再匹配模块名）
            # ----------------------------------------------------------------
            matched_systems_set = {}  # sys_key -> sys_data (去重用dict，O(1))
            matched_keywords_found = set()

            # 先收集所有匹配到的关键词
            for keyword, related_systems in system_keywords.items():
                if keyword and keyword in document_text:
                    matched_keywords_found.add(keyword)
                    for sys_data in related_systems:
                        sys_key = sys_data.get("system_name_en", "") or sys_data.get("system_name_cn", "")
                        if sys_key and sys_key not in matched_systems_set:
                            matched_systems_set[sys_key] = sys_data

            matched_systems = list(matched_systems_set.values())

            result = {
                "success": True,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat(),
                "matched_systems": matched_systems,
                "matched_count": len(matched_systems),
                "matched_keywords": list(matched_keywords_found),
                "total_systems_in_table": len(systems),
                "note": "提供了初步的文本关键词匹配结果，建议结合LLM进行语义级别的系统识别以提高准确率"
            }

            logger.info(f"[{call_id}] 匹配完成: 找到{len(matched_systems)}个关联系统, {len(matched_keywords_found)}个匹配关键词")
            return json.dumps(result, ensure_ascii=False, indent=2)

        except Exception as e:
            logger.exception(f"[{call_id}] 匹配系统信息异常")
            return json.dumps({
                "success": False,
                "error": f"匹配失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)

    # ============================================================
    # 第6步 - 文档修订工具
    # ============================================================

    @mcp.tool()
    def revise_document(file_path: str, suggestions_json: str, output_path: str = "", use_track_changes: str = "comment") -> str:
        """
        根据检查建议修订文档（第6步）

        支持三种修订模式：
        1. Word原生修订模式（Track Changes）：使用w:del和w:ins元素标记修改
        2. 红色括号说明模式：在原文后插入红色括号说明，不修改原文
        3. 直接替换模式：直接替换原文

        Args:
            file_path: 原文档路径
            suggestions_json: 检查建议JSON字符串
            output_path: 输出路径（可选）
            use_track_changes: 修订模式
                - 'true'（默认）：启用Word修订追踪（Track Changes）
                - 'comment'：在原文后插入红色括号说明
                - 'false'：直接替换原文

        Returns:
            JSON字符串，包含修订结果
        """
        start_time = datetime.now()
        call_id = f"revise_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始修订文档")
        logger.info(f"[{call_id}] file_path: {file_path}")
        logger.info(f"[{call_id}] use_track_changes: {use_track_changes}")

        try:
            reviser = DocumentReviser()
            result = reviser.revise_document(file_path, suggestions_json, output_path, use_track_changes)
            result['call_id'] = call_id
            result['timestamp'] = datetime.now().isoformat()

            logger.info(f"[{call_id}] 修订完成: {result.get('applied_revisions', 0)}条修订应用")
            return json.dumps(result, ensure_ascii=False, indent=2)

        except Exception as e:
            logger.exception(f"[{call_id}] 修订异常")
            return json.dumps({
                "success": False,
                "error": f"修订失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)

    @mcp.tool()
    def validate_suggestions(suggestions_json: str) -> str:
        """
        验证检查建议JSON格式
        
        Args:
            suggestions_json: 检查建议JSON字符串
            
        Returns:
            JSON字符串，包含验证结果
        """
        start_time = datetime.now()
        call_id = f"validate_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 验证建议JSON")

        try:
            data = json.loads(suggestions_json)
            
            if isinstance(data, dict):
                suggestions = data.get("suggestions", [])
            elif isinstance(data, list):
                suggestions = data
            else:
                return json.dumps({
                    "success": False,
                    "valid": False,
                    "error": "输入必须是JSON数组或包含suggestions的对象",
                    "call_id": call_id
                }, ensure_ascii=False)

            return json.dumps({
                "success": True,
                "valid": True,
                "count": len(suggestions),
                "message": f"验证通过，共{len(suggestions)}条建议",
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }, ensure_ascii=False)

        except json.JSONDecodeError as e:
            return json.dumps({
                "success": False,
                "valid": False,
                "error": f"JSON解析错误: {str(e)}",
                "call_id": call_id
            }, ensure_ascii=False)

    @mcp.tool()
    def get_document_tools_info() -> str:
        """获取文档修订工具信息"""
        tools_info = {
            "name": "UnifiedDocumentServer",
            "version": "1.2.0",
            "description": "统一文档MCP服务器，支持文档读取、系统信息获取和文档修订（含Word原生修订模式Track Changes）",
            "tools": [
                {
                    "name": "read_document",
                    "description": "读取文档内容（第1步）",
                    "category": "文档读取"
                },
                {
                    "name": "get_document_info",
                    "description": "获取文档基本信息（第1步辅助）",
                    "category": "文档读取"
                },
                {
                    "name": "read_system_info",
                    "description": "从Excel系统信息表读取系统基础信息（第1.5步 - 新增）",
                    "category": "系统信息"
                },
                {
                    "name": "match_system_from_document",
                    "description": "从文档文本中匹配关联的系统信息（第1.5步 - 辅助）",
                    "category": "系统信息"
                },
                {
                    "name": "revise_document",
                    "description": "根据建议修订文档（第6步）",
                    "category": "文档修订"
                },
                {
                    "name": "validate_suggestions",
                    "description": "验证建议JSON格式（第6步辅助）",
                    "category": "文档修订"
                }
            ]
        }
        return json.dumps(tools_info, ensure_ascii=False, indent=2)


# ============================================================
# 主入口
# ============================================================

def main():
    """主入口函数"""
    parser = argparse.ArgumentParser(description='统一文档MCP服务器')
    parser.add_argument(
        '--transport',
        choices=['stdio', 'sse'],
        default='sse',
        help='传输方式: stdio 或 sse (默认)'
    )
    parser.add_argument(
        '--host',
        default='40.129.21.85',
        help='SSE模式下的主机地址 (默认: 40.129.21.85)'
    )
    parser.add_argument(
        '--port',
        type=int,
        default=18080,
        help='SSE模式下的端口号 (默认: 18080)'
    )

    args = parser.parse_args()

    print("=" * 70)
    print("统一文档MCP服务器（第1步 + 第1.5步 + 第6步）v1.2.0")
    print("新增功能：系统信息读取 + 文档系统匹配")
    print("联动服务：时间轴图MCP服务器（timeline_mcp_server.py）")
    print("=" * 70)

    if not FASTMCP_AVAILABLE:
        print("\n错误: fastmcp未安装")
        print("请运行: pip install fastmcp")
        sys.exit(1)

    if not DOCX_AVAILABLE:
        print("\n错误: python-docx未安装")
        print("请运行: pip install python-docx")
        sys.exit(1)

    if not OPENPYXL_AVAILABLE:
        print("\n警告: openpyxl未安装，系统信息读取功能不可用")
        print("如需使用系统信息读取功能，请运行: pip install openpyxl")

    print(f"\n传输方式: {args.transport.upper()}")

    if args.transport == 'sse':
        print(f"服务地址: http://{args.host}:{args.port}")
        print(f"SSE端点: http://{args.host}:{args.port}/sse")
        print(f"消息端点: http://{args.host}:{args.port}/messages")

    print("\n可用工具:")
    print("  【第1步 - 文档读取】")
    print("    1. read_document          - 读取文档内容")
    print("    2. get_document_info      - 获取文档基本信息")
    print("  【第1.5步 - 系统信息获取（新增）】")
    print("    3. read_system_info       - 从Excel系统信息表读取系统基础信息")
    print("    4. match_system_from_document - 从文档文本中匹配关联的系统信息")
    print("  【第6步 - 文档修订】")
    print("    5. revise_document        - 根据建议修订文档")
    print("    6. validate_suggestions   - 验证建议JSON")
    print("  【时间轴图服务（同一端口）】")
    print("    7. generate_timeline       - 根据事件数据生成时间轴PNG图")
    print("    8. validate_timeline_data  - 验证时间轴数据格式")
    print("    9. get_timeline_template   - 获取时间轴数据模板")
    print("\n启动服务...")
    print("-" * 70)

    # 不再需要子进程启动timeline服务，所有工具已在同一个共享mcp实例中注册

    # 启动FastMCP服务
    if args.transport == 'sse':
        mcp.run(transport='sse', host=args.host, port=args.port)
    else:
        mcp.run()


def _start_timeline_server(transport: str):
    """
    启动时间轴图MCP服务（旧方案，不再使用）
    
    新方案通过 main.py 统一启动，所有工具在同一个共享 mcp 实例中注册。
    此函数保留仅用于向后兼容（直接运行 unified_mcp_server.py 时仍然独立启动timeline）。
    """
    # 当作为独立服务运行时，仍然需要启动timeline的子进程
    if not hasattr(mcp, '_SHARED_MODE') or not mcp._SHARED_MODE:
        timeline_file = Path(__file__).parent / "timeline_mcp_server.py"
        if not timeline_file.exists():
            logger.info("timeline_mcp_server.py不存在，跳过启动")
            return
        
        try:
            import subprocess
            if transport == 'sse':
                cmd = [sys.executable, str(timeline_file), '--transport', 'sse', '--host', '40.129.21.85', '--port', '8002']
                process = subprocess.Popen(
                    cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT,
                    creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
                )
                logger.info(f"时间轴MCP服务已启动: http://40.129.21.85:8002 (PID: {process.pid})")
                print(f"\n  【时间轴图服务 - http://40.129.21.85:8002】")
                print(f"    6. generate_timeline       - 根据事件数据生成时间轴PNG图")
                print(f"    7. validate_timeline_data  - 验证时间轴数据格式")
                print(f"    8. get_timeline_template   - 获取时间轴数据模板")
        except Exception as e:
            logger.warning(f"启动时间轴MCP服务失败: {e}")


if __name__ == "__main__":
    main()
