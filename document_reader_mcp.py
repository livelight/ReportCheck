"""
文档读取MCP节点 - 第1步
使用 FastMCP 框架实现
功能：读取文档内容，输出为JSON格式，作为第2步（硬性规则检查）和第3步（软性规则检查）的输入
支持格式：Word文档(.docx)、WPS文档(.wps和.wpsx)
"""

import json
import os
import sys
import logging
import zipfile
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass
from datetime import datetime

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger('DocumentReaderMCPNode')

# FastMCP导入
try:
    from fastmcp import FastMCP
    FASTMCP_AVAILABLE = True
except ImportError:
    FASTMCP_AVAILABLE = False
    logger.warning("fastmcp未安装，请运行: pip install fastmcp")

# 文档处理库
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx未安装，请运行: pip install python-docx")


@dataclass
class DocumentMetadata:
    """文档元数据模型"""
    title: str
    author: str
    created_date: str
    modified_date: str
    file_type: str
    file_size: int
    paragraph_count: int
    table_count: int
    page_count: Optional[int] = None
    company: Optional[str] = None


@dataclass
class DocumentSegment:
    """文档分段模型"""
    segment_id: int
    start_index: int
    end_index: int
    content: str
    char_count: int
    estimated_tokens: int
    start_paragraph_idx: int
    end_paragraph_idx: int
    is_table_segment: bool = False
    table_info: Optional[Dict] = None


class DocumentReader:
    """文档读取器 - 支持docx/wps/wpsx"""

    @staticmethod
    def read_document(file_path: str) -> Dict[str, Any]:
        """
        读取文档，返回文本内容和元数据
        
        Args:
            file_path: 文档路径
            
        Returns:
            Dict包含document_content和document_metadata
        """
        logger.info(f"开始读取文档: {file_path}")

        if not file_path or not file_path.strip():
            raise ValueError("文件路径不能为空")

        file_path = file_path.strip()

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = Path(file_path).suffix.lower()

        if ext == '.docx':
            return DocumentReader._read_docx(file_path)
        elif ext in ['.wps', '.wpsx']:
            return DocumentReader._read_wps(file_path)
        else:
            raise ValueError(f"不支持的文档格式: {ext}，仅支持 .docx, .wps, .wpsx")

    @staticmethod
    def _read_docx(file_path: str) -> Dict[str, Any]:
        """读取Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装，请运行: pip install python-docx")

        logger.info("读取Word文档 (.docx)")

        doc = DocxDocument(file_path)
        
        # 提取段落文本
        paragraphs = []
        full_text_lines = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({
                    'index': i,
                    'text': text,
                    'style': para.style.name if para.style else 'Normal'
                })
                full_text_lines.append(text)

        # 提取表格内容
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.rows[0].cells) if table.rows else 0,
                'data': table_data
            })

        # 提取元数据
        metadata = DocumentReader._extract_docx_metadata(doc, file_path)

        # 构建完整文本
        full_text = '\n'.join(full_text_lines)
        
        # 添加表格文本表示
        if tables:
            full_text += '\n\n【表格内容】\n'
            for table in tables:
                full_text += f'\n表{table["index"] + 1}:\n'
                for row_data in table['data']:
                    full_text += '| ' + ' | '.join(row_data) + ' |\n'

        logger.info(f"读取完成: {len(paragraphs)}个段落, {len(tables)}个表格")

        return {
            'success': True,
            'document_content': full_text,
            'document_metadata': {
                'title': metadata.title,
                'author': metadata.author,
                'created_date': metadata.created_date,
                'modified_date': metadata.modified_date,
                'file_type': metadata.file_type,
                'file_size': metadata.file_size,
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'page_count': metadata.page_count,
                'company': metadata.company
            },
            'structured_content': {
                'paragraphs': paragraphs,
                'tables': tables
            },
            'file_path': file_path
        }

    @staticmethod
    def _read_wps(file_path: str) -> Dict[str, Any]:
        """读取WPS文档（使用OOXML格式解析）"""
        logger.info("读取WPS文档 (.wps/.wpsx)")

        # WPS文档实际上是OOXML格式，可以像docx一样解析
        try:
            doc = DocxDocument(file_path)
            return DocumentReader._read_docx(file_path)
        except Exception as e:
            logger.warning(f"使用docx解析WPS失败，尝试使用zip解析: {e}")
            return DocumentReader._read_wps_via_zip(file_path)

    @staticmethod
    def _read_wps_via_zip(file_path: str) -> Dict[str, Any]:
        """通过ZIP方式读取WPS文档"""
        paragraphs = []
        tables = []
        
        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                # 读取document.xml
                if 'word/document.xml' in zf.namelist():
                    doc_xml = zf.read('word/document.xml').decode('utf-8')
                    # 简单提取文本（去除XML标签）
                    import re
                    text_content = re.sub(r'<[^>]+>', '', doc_xml)
                    paragraphs = [{'index': i, 'text': line.strip(), 'style': 'Normal'} 
                                  for i, line in enumerate(text_content.split('\n')) if line.strip()]
        except Exception as e:
            logger.error(f"读取WPS文档失败: {e}")
            raise

        full_text = '\n'.join([p['text'] for p in paragraphs])
        
        # 获取文件信息
        file_stat = os.stat(file_path)
        metadata = DocumentMetadata(
            title=Path(file_path).stem,
            author='',
            created_date=datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            modified_date=datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
            file_type='wps',
            file_size=file_stat.st_size,
            paragraph_count=len(paragraphs),
            table_count=len(tables)
        )

        logger.info(f"WPS读取完成: {len(paragraphs)}个段落")

        return {
            'success': True,
            'document_content': full_text,
            'document_metadata': {
                'title': metadata.title,
                'author': metadata.author,
                'created_date': metadata.created_date,
                'modified_date': metadata.modified_date,
                'file_type': metadata.file_type,
                'file_size': metadata.file_size,
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'page_count': metadata.page_count,
                'company': metadata.company
            },
            'structured_content': {
                'paragraphs': paragraphs,
                'tables': tables
            },
            'file_path': file_path
        }

    @staticmethod
    def _extract_docx_metadata(doc, file_path: str) -> DocumentMetadata:
        """提取docx文档元数据"""
        file_stat = os.stat(file_path)
        
        # 从core_properties获取元数据
        try:
            core_props = doc.core_properties
            title = core_props.title or Path(file_path).stem
            author = core_props.author or ''
            created = core_props.created.isoformat() if core_props.created else ''
            modified = core_props.modified.isoformat() if core_props.modified else ''
        except Exception:
            title = Path(file_path).stem
            author = ''
            created = datetime.fromtimestamp(file_stat.st_ctime).isoformat()
            modified = datetime.fromtimestamp(file_stat.st_mtime).isoformat()

        return DocumentMetadata(
            title=title,
            author=author,
            created_date=created,
            modified_date=modified,
            file_type='docx',
            file_size=file_stat.st_size,
            paragraph_count=len(doc.paragraphs),
            table_count=len(doc.tables)
        )


class DocumentSegmenter:
    """文档分段器 - 将长文档分割成多个段落块"""
    
    # 默认配置：每个分段约4000 tokens（留出余量避免溢出）
    DEFAULT_MAX_TOKENS_PER_SEGMENT = 4000
    # 中文字符与token的估算比例（1个汉字约0.5-0.75个token，英文约0.25-0.5个token）
    # 保守估算：1 token ≈ 1个中文字符 或 4个英文字符
    CHARS_PER_TOKEN_CN = 1.0
    CHARS_PER_TOKEN_EN = 4.0
    
    def __init__(self, max_tokens_per_segment: int = DEFAULT_MAX_TOKENS_PER_SEGMENT):
        self.max_tokens_per_segment = max_tokens_per_segment
    
    def estimate_tokens(self, text: str) -> int:
        """估算文本的token数量（简化版）"""
        if not text:
            return 0
        
        cn_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
        en_chars = len(text) - cn_chars
        
        # 保守估算
        estimated = int(cn_chars / self.CHARS_PER_TOKEN_CN + en_chars / self.CHARS_PER_TOKEN_EN)
        return max(1, estimated)
    
    def segment_document(self, paragraphs: List[Dict[str, Any]], tables: List[Dict[str, Any]]) -> List[DocumentSegment]:
        """
        将文档内容分段
        
        策略：
        1. 保持段落完整性（不切割单个段落）
        2. 表格作为独立分段
        3. 每个分段token数不超过阈值
        4. 优先在章节边界处分割
        
        Args:
            paragraphs: 段落列表
            tables: 表格列表
            
        Returns:
            分段列表
        """
        segments = []
        current_segment_content = []
        current_tokens = 0
        current_char_count = 0
        segment_id = 0
        start_para_idx = 0
        global_char_idx = 0
        
        # 识别章节标题模式
        section_patterns = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、',
                           '（一）', '（二）', '（三）', '（四）', '（五）',
                           '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.',
                           '1、', '2、', '3、', '4、', '5、']
        
        def is_section_boundary(text: str) -> bool:
            """检查是否是章节边界"""
            for pattern in section_patterns:
                if text.strip().startswith(pattern):
                    return True
            return False
        
        def flush_segment(end_para_idx: int, is_final: bool = False):
            """将当前累积的内容保存为一个分段"""
            nonlocal segment_id, current_segment_content, current_tokens, current_char_count
            nonlocal start_para_idx, global_char_idx
            
            if not current_segment_content:
                return
            
            content = '\n'.join(current_segment_content)
            segment = DocumentSegment(
                segment_id=segment_id,
                start_index=global_char_idx - current_char_count,
                end_index=global_char_idx,
                content=content,
                char_count=current_char_count,
                estimated_tokens=current_tokens,
                start_paragraph_idx=start_para_idx,
                end_paragraph_idx=end_para_idx if is_final else end_para_idx - 1
            )
            segments.append(segment)
            segment_id += 1
            current_segment_content = []
            current_tokens = 0
            current_char_count = 0
            start_para_idx = end_para_idx
        
        # 处理段落
        for i, para in enumerate(paragraphs):
            text = para.get('text', '')
            if not text:
                continue
            
            para_tokens = self.estimate_tokens(text)
            para_chars = len(text)
            
            # 检查是否是章节边界且当前分段已有内容
            if is_section_boundary(text) and current_segment_content:
                flush_segment(i)
            
            # 检查加入当前段落后是否超过token限制
            if current_tokens + para_tokens > self.max_tokens_per_segment and current_segment_content:
                flush_segment(i)
            
            current_segment_content.append(text)
            current_tokens += para_tokens
            current_char_count += para_chars + 1  # +1 for newline
            global_char_idx += para_chars + 1
        
        # 刷新最后一个段落分段
        if current_segment_content:
            flush_segment(len(paragraphs), is_final=True)
        
        # 处理表格 - 每个表格作为独立分段
        for table in tables:
            table_text = self._table_to_text(table)
            table_tokens = self.estimate_tokens(table_text)
            
            segment = DocumentSegment(
                segment_id=segment_id,
                start_index=global_char_idx,
                end_index=global_char_idx + len(table_text),
                content=table_text,
                char_count=len(table_text),
                estimated_tokens=table_tokens,
                start_paragraph_idx=-1,  # 表格没有段落索引
                end_paragraph_idx=-1,
                is_table_segment=True,
                table_info={
                    'table_index': table.get('index', 0),
                    'rows': table.get('rows', 0),
                    'cols': table.get('cols', 0)
                }
            )
            segments.append(segment)
            segment_id += 1
            global_char_idx += len(table_text)
        
        return segments
    
    def _table_to_text(self, table: Dict[str, Any]) -> str:
        """将表格转换为文本表示"""
        lines = [f"【表{table.get('index', 0) + 1}】"]
        for row_data in table.get('data', []):
            lines.append('| ' + ' | '.join(row_data) + ' |')
        return '\n'.join(lines)
    
    def get_segment_info(self, segments: List[DocumentSegment]) -> Dict[str, Any]:
        """获取分段统计信息"""
        if not segments:
            return {
                'total_segments': 0,
                'total_tokens': 0,
                'avg_tokens_per_segment': 0,
                'max_tokens': 0,
                'min_tokens': 0
            }
        
        token_counts = [s.estimated_tokens for s in segments]
        return {
            'total_segments': len(segments),
            'total_tokens': sum(token_counts),
            'avg_tokens_per_segment': sum(token_counts) / len(token_counts),
            'max_tokens': max(token_counts),
            'min_tokens': min(token_counts),
            'table_segments': sum(1 for s in segments if s.is_table_segment),
            'text_segments': sum(1 for s in segments if not s.is_table_segment)
        }


class DocumentReaderMCPNode:
    """文档读取MCP节点"""
    
    name = "DocumentReaderNode"
    version = "1.1.0"
    description = "读取文档内容，支持文档分段，输出为JSON格式供后续检查步骤使用"

    def __init__(self, max_tokens_per_segment: int = 4000):
        self.reader = DocumentReader()
        self.segmenter = DocumentSegmenter(max_tokens_per_segment)

    def read_document(self, file_path: str) -> Dict[str, Any]:
        """
        读取文档并返回内容
        
        Args:
            file_path: 文档路径
            
        Returns:
            包含document_content和document_metadata的字典
        """
        return self.reader.read_document(file_path)
    
    def read_document_segmented(self, file_path: str, max_tokens_per_segment: Optional[int] = None) -> Dict[str, Any]:
        """
        读取文档并分段返回
        
        Args:
            file_path: 文档路径
            max_tokens_per_segment: 每段最大token数（可选，覆盖默认值）
            
        Returns:
            包含document_content、segments和segment_info的字典
        """
        # 读取文档
        result = self.reader.read_document(file_path)
        if not result.get('success'):
            return result
        
        # 使用指定的token限制或默认限制
        if max_tokens_per_segment:
            segmenter = DocumentSegmenter(max_tokens_per_segment)
        else:
            segmenter = self.segmenter
        
        # 获取结构化内容
        structured = result.get('structured_content', {})
        paragraphs = structured.get('paragraphs', [])
        tables = structured.get('tables', [])
        
        # 分段
        segments = segmenter.segment_document(paragraphs, tables)
        segment_info = segmenter.get_segment_info(segments)
        
        # 转换分段为可序列化的字典
        segments_data = []
        for seg in segments:
            segments_data.append({
                'segment_id': seg.segment_id,
                'start_index': seg.start_index,
                'end_index': seg.end_index,
                'content': seg.content,
                'char_count': seg.char_count,
                'estimated_tokens': seg.estimated_tokens,
                'start_paragraph_idx': seg.start_paragraph_idx,
                'end_paragraph_idx': seg.end_paragraph_idx,
                'is_table_segment': seg.is_table_segment,
                'table_info': seg.table_info
            })
        
        # 构建结果
        result['segments'] = segments_data
        result['segment_info'] = segment_info
        result['segmentation_config'] = {
            'max_tokens_per_segment': max_tokens_per_segment or self.segmenter.max_tokens_per_segment
        }
        
        return result

    def get_tool_definition(self) -> Dict[str, Any]:
        """获取工具定义（用于MCP）"""
        return {
            "name": self.name,
            "description": self.description,
            "version": self.version,
            "tools": [
                {
                    "name": "read_document",
                    "description": "读取文档内容，返回文档文本和元数据",
                    "input_schema": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "文档路径，支持.docx, .wps, .wpsx格式"
                            }
                        },
                        "required": ["file_path"]
                    }
                },
                {
                    "name": "get_document_info",
                    "description": "获取文档基本信息（不读取完整内容）",
                    "input_schema": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "文档路径"
                            }
                        },
                        "required": ["file_path"]
                    }
                },
                {
                    "name": "read_document_segmented",
                    "description": "读取文档并分段返回，支持长文档处理",
                    "input_schema": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "文档路径，支持.docx, .wps, .wpsx格式"
                            },
                            "max_tokens_per_segment": {
                                "type": "integer",
                                "description": "每段最大token数（默认4000），用于控制分段粒度",
                                "default": 4000
                            }
                        },
                        "required": ["file_path"]
                    }
                },
                {
                    "name": "get_tools_info",
                    "description": "获取工具信息和定义",
                    "input_schema": {
                        "type": "object",
                        "properties": {}
                    }
                }
            ]
        }


# ============================================================
# MCP工具定义
# ============================================================

if FASTMCP_AVAILABLE:
    mcp = FastMCP("DocumentReaderNode")

    @mcp.tool()
    def read_document(file_path: str) -> str:
        """
        读取文档内容，返回文档文本和元数据
        
        按流程完成文档读取、内容提取、元数据提取，输出为JSON格式。
        支持docx、wps、wpsx格式，输出包含document_content和document_metadata字段。
        
        Args:
            file_path: 文档路径，支持 .docx、.wps、.wpsx 格式
            
        Returns:
            str: JSON字符串，包含以下字段：
                - success (bool): 操作是否成功
                - document_content (str): 文档完整文本内容
                - document_metadata (dict): 文档元数据
                    - title (str): 文档标题
                    - author (str): 作者
                    - created_date (str): 创建日期
                    - modified_date (str): 修改日期
                    - file_type (str): 文件类型
                    - file_size (int): 文件大小（字节）
                    - paragraph_count (int): 段落数量
                    - table_count (int): 表格数量
                - structured_content (dict): 结构化内容
                    - paragraphs (list): 段落列表
                    - tables (list): 表格列表
                - file_path (str): 文件路径
                - error (str): 错误描述（失败时）
                - error_type (str): 错误类型（失败时）
        """
        start_time = datetime.now()
        call_id = f"read_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始读取文档")
        logger.info(f"[{call_id}] 输入 - file_path: {file_path}")

        try:
            # 1. 验证参数
            if not file_path or not file_path.strip():
                error_result = {
                    "success": False,
                    "error": "file_path不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            file_path = file_path.strip()

            # 2. 验证文件存在
            if not os.path.exists(file_path):
                error_result = {
                    "success": False,
                    "error": f"文件不存在: {file_path}",
                    "error_type": "FileNotFoundError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            # 3. 读取文档
            reader = DocumentReader()
            result = reader.read_document(file_path)
            result['call_id'] = call_id
            result['timestamp'] = datetime.now().isoformat()

            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            result['duration_seconds'] = duration

            logger.info(f"[{call_id}] 读取完成! 耗时: {duration:.2f}秒")
            logger.info(f"[{call_id}] 段落数: {result['document_metadata']['paragraph_count']}")
            logger.info(f"[{call_id}] 表格数: {result['document_metadata']['table_count']}")

            return json.dumps(result, ensure_ascii=False, indent=2)

        except Exception as e:
            error_result = {
                "success": False,
                "error": f"读取失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }
            logger.exception(f"[{call_id}] 读取异常")
            return json.dumps(error_result, ensure_ascii=False)

    @mcp.tool()
    def read_document_segmented(file_path: str, max_tokens_per_segment: int = 4000) -> str:
        """
        读取文档并按智能分段返回，适用于长文档处理

        将文档分割成多个段落块，每个块包含约4000 tokens（可配置），
        保持段落完整性，优先在章节边界处分割。返回分段信息和统计。

        Args:
            file_path: 文档路径，支持 .docx、.wps、.wpsx 格式
            max_tokens_per_segment: 每段最大token数（默认4000），建议范围2000-8000

        Returns:
            str: JSON字符串，包含以下字段：
                - success (bool): 操作是否成功
                - document_content (str): 文档完整文本内容
                - document_metadata (dict): 文档元数据
                - structured_content (dict): 结构化内容（段落、表格）
                - segments (list): 分段列表，每个分段包含：
                    - segment_id (int): 分段ID
                    - content (str): 分段内容
                    - estimated_tokens (int): 估算token数
                    - start_index/end_index (int): 在全文中的字符位置
                    - is_table_segment (bool): 是否是表格分段
                - segment_info (dict): 分段统计信息
                    - total_segments (int): 总分段数
                    - total_tokens (int): 总token数
                    - avg_tokens_per_segment (float): 平均每段token数
                - segmentation_config (dict): 分段配置
                - file_path (str): 文件路径
        """
        start_time = datetime.now()
        call_id = f"seg_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 开始分段读取文档")
        logger.info(f"[{call_id}] 输入 - file_path: {file_path}, max_tokens: {max_tokens_per_segment}")

        try:
            # 1. 验证参数
            if not file_path or not file_path.strip():
                error_result = {
                    "success": False,
                    "error": "file_path不能为空",
                    "error_type": "ValidationError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            file_path = file_path.strip()

            # 2. 验证文件存在
            if not os.path.exists(file_path):
                error_result = {
                    "success": False,
                    "error": f"文件不存在: {file_path}",
                    "error_type": "FileNotFoundError",
                    "call_id": call_id,
                    "timestamp": start_time.isoformat()
                }
                logger.error(f"[{call_id}] {error_result['error']}")
                return json.dumps(error_result, ensure_ascii=False)

            # 3. 分段读取文档
            node = DocumentReaderMCPNode(max_tokens_per_segment=max_tokens_per_segment)
            result = node.read_document_segmented(file_path, max_tokens_per_segment)
            result['call_id'] = call_id
            result['timestamp'] = datetime.now().isoformat()

            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            result['duration_seconds'] = duration

            segment_info = result.get('segment_info', {})
            logger.info(f"[{call_id}] 分段读取完成! 耗时: {duration:.2f}秒")
            logger.info(f"[{call_id}] 总分段数: {segment_info.get('total_segments', 0)}")
            logger.info(f"[{call_id}] 总估算token数: {segment_info.get('total_tokens', 0)}")
            logger.info(f"[{call_id}] 平均每段token数: {segment_info.get('avg_tokens_per_segment', 0):.1f}")

            return json.dumps(result, ensure_ascii=False, indent=2)

        except Exception as e:
            error_result = {
                "success": False,
                "error": f"分段读取失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }
            logger.exception(f"[{call_id}] 分段读取异常")
            return json.dumps(error_result, ensure_ascii=False)

    @mcp.tool()
    def get_document_info(file_path: str) -> str:
        """
        获取文档基本信息（不读取完整内容）
        
        用于快速检查文档是否存在及基本属性。
        
        Args:
            file_path: 文档路径
            
        Returns:
            str: JSON字符串，包含文档基本信息
        """
        start_time = datetime.now()
        call_id = f"info_{start_time.strftime('%Y%m%d_%H%M%S')}"

        logger.info(f"[{call_id}] 获取文档信息: {file_path}")

        try:
            if not file_path or not file_path.strip():
                return json.dumps({
                    "success": False,
                    "error": "file_path不能为空",
                    "call_id": call_id
                }, ensure_ascii=False)

            file_path = file_path.strip()

            if not os.path.exists(file_path):
                return json.dumps({
                    "success": False,
                    "error": f"文件不存在: {file_path}",
                    "call_id": call_id
                }, ensure_ascii=False)

            file_stat = os.stat(file_path)
            file_type = Path(file_path).suffix.lower()

            result = {
                "success": True,
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "file_type": file_type,
                "file_size": file_stat.st_size,
                "file_size_readable": f"{file_stat.st_size / 1024:.2f} KB",
                "created_time": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }

            logger.info(f"[{call_id}] 文档信息获取成功")
            return json.dumps(result, ensure_ascii=False, indent=2)

        except Exception as e:
            error_result = {
                "success": False,
                "error": f"获取信息失败: {str(e)}",
                "error_type": type(e).__name__,
                "call_id": call_id,
                "timestamp": datetime.now().isoformat()
            }
            logger.exception(f"[{call_id}] 获取信息异常")
            return json.dumps(error_result, ensure_ascii=False)

    @mcp.tool()
    def get_tools_info() -> str:
        """获取工具信息"""
        node = DocumentReaderMCPNode()
        return json.dumps(node.get_tool_definition(), ensure_ascii=False, indent=2)


# ============================================================
# 主入口
# ============================================================

def main():
    """主入口函数 - 支持stdio和SSE两种传输方式"""
    import argparse

    parser = argparse.ArgumentParser(description='文档读取MCP节点服务')
    parser.add_argument(
        '--transport',
        choices=['stdio', 'sse'],
        default='stdio',
        help='传输方式: stdio (默认) 或 sse'
    )
    parser.add_argument(
        '--host',
        default='0.0.0.0',
        help='SSE模式下的主机地址 (默认: 0.0.0.0)'
    )
    parser.add_argument(
        '--port',
        type=int,
        default=8001,
        help='SSE模式下的端口号 (默认: 8001)'
    )

    args = parser.parse_args()

    print("=" * 60)
    print("文档读取MCP节点服务（第1步）v1.1.0")
    print("新增功能：智能文档分段，支持长文档处理")
    print("=" * 60)

    if not FASTMCP_AVAILABLE:
        print("\n错误: fastmcp未安装")
        print("请运行: pip install fastmcp")
        sys.exit(1)

    if not DOCX_AVAILABLE:
        print("\n错误: python-docx未安装")
        print("请运行: pip install python-docx")
        sys.exit(1)

    print(f"\n传输方式: {args.transport.upper()}")

    if args.transport == 'sse':
        print(f"服务地址: http://{args.host}:{args.port}")
        print(f"SSE端点: http://{args.host}:{args.port}/sse")
        print(f"消息端点: http://{args.host}:{args.port}/messages")

    print("\n可用工具:")
    print("  1. read_document          - 读取文档内容")
    print("  2. read_document_segmented - 分段读取文档（支持长文档）")
    print("  3. get_document_info      - 获取文档基本信息")
    print("  4. get_tools_info         - 工具信息")
    print("\n启动服务...")
    print("-" * 60)

    # 启动FastMCP服务
    if args.transport == 'sse':
        mcp.run(transport='sse', host=args.host, port=args.port)
    else:
        mcp.run()


if __name__ == "__main__":
    main()
