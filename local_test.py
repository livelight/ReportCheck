"""
创建示例文档并进行修订测试
"""

import json
import os
from docx import Document
from revision_mcp_node import RevisionMCPNode, DocumentParser


def create_sample_docx():
    """创建示例Word文档"""
    doc = Document()

    # 标题
    doc.add_heading('企业文档管理系统项目报告', level=1)

    # 章节1 - 项目概述
    doc.add_heading('一、项目概述', level=2)
    doc.add_paragraph('本项目旨在开发一套企业级文档管理系统，以提高企业内部文档管理效率。')

    # 章节2 - 项目目标
    doc.add_heading('二、项目目标', level=2)
    doc.add_paragraph('1. 实现文档的集中存储和分类管理')
    doc.add_paragraph('2. 支持多用户协作编辑')
    doc.add_paragraph('3. 提供文档搜索和版本控制功能')
    doc.add_paragraph('4. 确保文档安全性和权限管理')

    # 章节3 - 当前进展
    doc.add_heading('三、当前进展', level=2)
    doc.add_paragraph('目前已完成核心功能的开发和测试工作，进展顺利。')

    # 章节4 - 风险分析
    doc.add_heading('四、风险分析', level=2)
    doc.add_paragraph('项目进展顺利，暂无风险。')

    # 章节5 - 下一步计划
    doc.add_heading('五、下一步计划', level=2)
    doc.add_paragraph('继续优化系统性能，并进行用户验收测试。')

    # 章节6 - 总结
    doc.add_heading('六、总结', level=2)
    doc.add_paragraph('整体项目进度符合预期。')

    # 保存
    path = 'sample_document.docx'
    doc.save(path)
    return path


def create_sample_wpsx():
    """创建示例WPS文档（模拟WPSX格式）"""
    import zipfile
    from xml.etree import ElementTree as ET

    path = 'sample_wpsx.wpsx'

    # WPSX本质上是一个ZIP压缩包，包含XML内容
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zf:
        # 创建content.xml
        content_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:rPr><w:b/></w:rPr><w:t>企业文档管理系统项目报告</w:t></w:r></w:p>
    <w:p><w:r><w:rPr><w:b/><w:sz w:val="28"/></w:rPr><w:t>一、项目概述</w:t></w:r></w:p>
    <w:p><w:r><w:t>本项目旨在开发一套企业级文档管理系统，以提高企业内部文档管理效率。</w:t></w:r></w:p>
    <w:p><w:r><w:rPr><w:b/><w:sz w:val="28"/></w:rPr><w:t>二、项目目标</w:t></w:r></w:p>
    <w:p><w:r><w:t>1. 实现文档的集中存储和分类管理</w:t></w:r></w:p>
    <w:p><w:r><w:t>2. 支持多用户协作编辑</w:t></w:r></w:p>
    <w:p><w:r><w:t>3. 提供文档搜索和版本控制功能</w:t></w:r></w:p>
    <w:p><w:r><w:t>4. 确保文档安全性和权限管理</w:t></w:r></w:p>
    <w:p><w:r><w:rPr><w:b/><w:sz w:val="28"/></w:rPr><w:t>三、当前进展</w:t></w:r></w:p>
    <w:p><w:r><w:t>目前已完成核心功能的开发和测试工作，进展顺利。</w:t></w:r></w:p>
    <w:p><w:r><w:rPr><w:b/><w:sz w:val="28"/></w:rPr><w:t>四、风险分析</w:t></w:r></w:p>
    <w:p><w:r><w:t>项目进展顺利，暂无风险。</w:t></w:r></w:p>
    <w:p><w:r><w:rPr><w:b/><w:sz w:val="28"/></w:rPr><w:t>五、下一步计划</w:t></w:r></w:p>
    <w:p><w:r><w:t>继续优化系统性能，并进行用户验收测试。</w:t></w:r></w:p>
    <w:p><w:r><w:rPr><w:b/><w:sz w:val="28"/></w:rPr><w:t>六、总结</w:t></w:r></w:p>
    <w:p><w:r><w:t>整体项目进度符合预期。</w:t></w:r></w:p>
  </w:body>
</w:document>'''

        zf.writestr('content.xml', content_xml)

        # 创建[Content_Types].xml
        content_types = '''<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="xml" ContentType="application/xml"/>
</Types>'''
        zf.writestr('[Content_Types].xml', content_types)

    return path


def run_test():
    """运行完整测试"""
    print("=" * 70)
    print("文档修订MCP节点 - 本地文档测试")
    print("=" * 70)

    node = RevisionMCPNode()

    # 创建示例文档
    print("\n[步骤1] 创建示例文档...")

    docx_path = create_sample_docx()
    print(f"  创建Word文档: {docx_path}")
    print(f"  文件大小: {os.path.getsize(docx_path)} bytes")

    wpsx_path = create_sample_wpsx()
    print(f"  创建WPS文档: {wpsx_path}")
    print(f"  文件大小: {os.path.getsize(wpsx_path)} bytes")

    # 定义检查建议
    print("\n[步骤2] 定义检查建议...")
    suggestions = [
        {
            "id": "ref_001",
            "rule_id": "FORMAT_001",
            "rule_name": "标题层级不规范",
            "type": "format",
            "severity": "Medium",
            "section": "sec_001",
            "original_text": "企业文档管理系统项目报告",
            "suggestion": "企业文档管理系统项目报告（修订版）",
            "reason": "标题应明确标注版本信息"
        },
        {
            "id": "ref_002",
            "rule_id": "CONTENT_015",
            "rule_name": "项目目标描述不完整",
            "type": "content",
            "severity": "High",
            "section": "sec_002",
            "original_text": "1. 实现文档的集中存储和分类管理",
            "suggestion": "1. 实现文档的集中存储和分类管理\n2. 支持多用户协作编辑\n3. 提供文档搜索和版本控制功能\n4. 确保文档安全性和权限管理\n5. 支持移动端访问和离线编辑",
            "reason": "缺少移动端支持的目标描述"
        },
        {
            "id": "ref_003",
            "rule_id": "CONTENT_023",
            "rule_name": "风险分析缺失",
            "type": "content",
            "severity": "High",
            "section": "sec_004",
            "original_text": "项目进展顺利，暂无风险。",
            "suggestion": "当前主要风险及应对措施：\n1. 人员流动性风险：加强知识文档建设，确保关键人员备份\n2. 技术风险：预留技术预研时间，及时评估新技术方案\n3. 进度风险：采用敏捷开发模式，分阶段交付",
            "reason": "风险分析过于简单，缺乏具体风险点和应对措施"
        },
        {
            "id": "ref_004",
            "rule_id": "LANG_002",
            "rule_name": "语句不够精炼",
            "type": "language",
            "severity": "Low",
            "section": "sec_003",
            "original_text": "目前已完成核心功能的开发和测试工作，进展顺利。",
            "suggestion": "目前已完成核心功能的开发和内部测试，整体质量符合预期。",
            "reason": "'进展顺利'表述不够专业"
        },
        {
            "id": "ref_005",
            "rule_id": "LOGIC_001",
            "rule_name": "逻辑顺序需调整",
            "type": "logic",
            "severity": "Medium",
            "section": "sec_005",
            "original_text": "继续优化系统性能，并进行用户验收测试。",
            "suggestion": "1. 完成用户验收测试（UAT）\n2. 根据测试反馈进行系统优化\n3. 准备生产环境部署\n4. 正式上线并跟踪运行状态",
            "reason": "应先测试后优化，最后部署上线"
        }
    ]

    suggestions_json = json.dumps(suggestions, ensure_ascii=False, indent=2)
    print(f"  共 {len(suggestions)} 条检查建议")
    print("\n  建议摘要：")
    for s in suggestions:
        print(f"    [{s['severity']:8}] {s['rule_name']} ({s['type']})")

    # ============ 测试 DOCX ============
    print("\n" + "-" * 70)
    print("[步骤3] 测试 Word文档 (.docx) 修订...")
    print("-" * 70)

    result_docx = node.revise_document(
        file_path=docx_path,
        suggestions_json=suggestions_json
    )

    print(f"\n修订结果:")
    print(f"  成功: {result_docx.get('success')}")
    print(f"  输出文件: {result_docx.get('output_path')}")
    print(f"  总建议数: {result_docx.get('total_suggestions')}")
    print(f"  已应用: {result_docx.get('applied_revisions')}")
    print(f"  已跳过: {result_docx.get('skipped_revisions')}")

    if result_docx.get('success'):
        print(f"\n输出文件存在: {os.path.exists(result_docx.get('output_path', ''))}")
        print(f"文件大小: {os.path.getsize(result_docx.get('output_path', ''))} bytes")

    print("\n  修订详情:")
    for rev in result_docx.get('revisions_detail', []):
        print(f"    [{rev['severity']:8}] {rev['rule_name']}")
        print(f"      原文: {rev['original'][:40]}...")
        print(f"      新文: {rev['new'][:40]}...")
        print(f"      状态: {rev['status']}")

    # 读取修订后的docx内容
    print("\n  修订后文档内容预览:")
    revised_doc = Document(result_docx.get('output_path'))
    for i, para in enumerate(revised_doc.paragraphs[:12]):
        if para.text.strip():
            print(f"    {para.text[:60]}")

    # ============ 测试 WPSX ============
    print("\n" + "-" * 70)
    print("[步骤4] 测试 WPS文档 (.wpsx) 修订...")
    print("-" * 70)

    # 解析WPSX内容
    print("\nWPSX文档内容:")
    doc_data = DocumentParser.parse_document(wpsx_path)
    for para in doc_data.get('paragraphs', [])[:8]:
        print(f"  {para['text'][:50]}")

    # 执行修订
    result_wpsx = node.revise_document(
        file_path=wpsx_path,
        suggestions_json=suggestions_json
    )

    print(f"\n修订结果:")
    print(f"  成功: {result_wpsx.get('success')}")
    print(f"  总建议数: {result_wpsx.get('total_suggestions')}")
    print(f"  已应用: {result_wpsx.get('applied_revisions')}")
    print(f"  提示: {result_wpsx.get('note', 'N/A')}")

    # ============ 统计汇总 ============
    print("\n" + "=" * 70)
    print("[测试总结]")
    print("=" * 70)

    print("\n修订统计 (by_severity):")
    stats = result_docx.get('statistics', {})
    for severity, count in stats.get('by_severity', {}).items():
        bar = '█' * count
        print(f"  {severity:10}: {count} {bar}")

    print("\n修订统计 (by_type):")
    for type_name, count in stats.get('by_type', {}).items():
        bar = '█' * count
        print(f"  {type_name:10}: {count} {bar}")

    print("\n输出文件:")
    print(f"  DOCX: {os.path.abspath(result_docx.get('output_path', ''))}")

    print("\n" + "=" * 70)
    print("测试完成!")
    print("=" * 70)


if __name__ == "__main__":
    run_test()
