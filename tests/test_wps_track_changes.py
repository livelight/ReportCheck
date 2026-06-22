"""
测试 WPS/WPSX 文档修订模式功能
"""

import json
import os
from docx import Document
from revision_mcp_node import (
    FASTMCP_AVAILABLE,
    DOCX_AVAILABLE,
    CheckSuggestion,
    DocumentParser,
    TrackChangesReviser
)


def create_sample_wps():
    """创建示例WPS文档 (.wps) - 使用OOXML格式"""
    from docx import Document
    from docx.shared import Pt

    path = 'test_sample.wps'

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    doc.add_heading('企业文档管理系统项目报告', level=0)
    doc.add_heading('一、项目概述', level=1)
    doc.add_paragraph('本项目旨在开发一套企业级文档管理系统。')
    doc.add_heading('二、项目目标', level=1)
    doc.add_paragraph('1. 实现文档的集中存储和分类管理')
    doc.add_paragraph('2. 支持多用户协作编辑')
    doc.add_heading('三、当前进展', level=1)
    doc.add_paragraph('目前已完成核心功能的开发和测试工作，进展顺利。')
    doc.add_heading('四、风险分析', level=1)
    doc.add_paragraph('项目进展顺利，暂无风险。')
    doc.add_heading('五、下一步计划', level=1)
    doc.add_paragraph('继续优化系统性能，并进行用户验收测试。')

    # 保存为docx格式，然后重命名为.wps
    temp_docx = path.replace('.wps', '_temp.docx')
    doc.save(temp_docx)

    import zipfile
    with zipfile.ZipFile(temp_docx, 'r') as zf_src:
        with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zf_dst:
            for item in zf_src.namelist():
                zf_dst.writestr(item, zf_src.read(item))

    try:
        os.remove(temp_docx)
    except:
        pass

    return path


def create_sample_wpsx():
    """创建示例WPSX文档 (.wpsx)"""
    from docx import Document
    from docx.shared import Pt

    path = 'test_sample.wpsx'

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    doc.add_heading('企业文档管理系统项目报告', level=0)
    doc.add_heading('一、项目概述', level=1)
    doc.add_paragraph('本项目旨在开发一套企业级文档管理系统。')
    doc.add_heading('二、项目目标', level=1)
    doc.add_paragraph('1. 实现文档的集中存储和分类管理')
    doc.add_paragraph('2. 支持多用户协作编辑')
    doc.add_heading('三、当前进展', level=1)
    doc.add_paragraph('目前已完成核心功能的开发和测试工作，进展顺利。')
    doc.add_heading('四、风险分析', level=1)
    doc.add_paragraph('项目进展顺利，暂无风险。')
    doc.add_heading('五、下一步计划', level=1)
    doc.add_paragraph('继续优化系统性能，并进行用户验收测试。')

    # 保存为docx格式，然后重命名为.wpsx
    temp_docx = path.replace('.wpsx', '_temp.docx')
    doc.save(temp_docx)

    import zipfile
    with zipfile.ZipFile(temp_docx, 'r') as zf_src:
        with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zf_dst:
            for item in zf_src.namelist():
                zf_dst.writestr(item, zf_src.read(item))

    try:
        os.remove(temp_docx)
    except:
        pass

    return path


def test_wps_revision_with_track_changes():
    """测试WPS文档修订模式"""
    print("=" * 70)
    print("WPS/WPSX 文档修订模式测试")
    print("=" * 70)

    # 创建测试文档
    print("\n[步骤1] 创建测试文档")
    wps_path = create_sample_wps()
    wpsx_path = create_sample_wpsx()

    print(f"  WPS文件: {wps_path} ({os.path.getsize(wps_path)} bytes)")
    print(f"  WPSX文件: {wpsx_path} ({os.path.getsize(wpsx_path)} bytes)")

    # 定义检查建议
    suggestions = [
        CheckSuggestion(
            id="ref_001",
            rule_id="CONTENT_023",
            rule_name="风险分析缺失",
            type="content",
            severity="High",
            section="sec_004",
            original_text="项目进展顺利，暂无风险。",
            suggestion="当前主要风险及应对措施：\n1. 人员流动性风险\n2. 技术架构风险",
            reason="风险分析过于简单"
        ),
        CheckSuggestion(
            id="ref_002",
            rule_id="LANG_001",
            rule_name="表述不规范",
            type="language",
            severity="Medium",
            section="sec_003",
            original_text="目前已完成核心功能的开发和测试工作，进展顺利。",
            suggestion="目前已完成核心功能的开发和内部测试，整体质量符合预期。",
            reason="表述不够专业"
        )
    ]

    print(f"\n检查建议:")
    for s in suggestions:
        print(f"  [{s.severity}] {s.rule_name}")

    # 测试WPS修订
    print("\n" + "-" * 70)
    print("[步骤2] 测试WPS文档修订（使用修订模式）")
    print("-" * 70)

    try:
        # 解析文档
        doc_data = DocumentParser.parse_document(wps_path)
        print(f"[OK] 文档解析成功，{len(doc_data['paragraphs'])}个段落")

        # 使用TrackChangesReviser
        reviser = TrackChangesReviser(doc_data)
        result = reviser.apply_suggestions(suggestions)

        print(f"[OK] 修订应用成功")
        print(f"     总建议: {result['total_suggestions']}")
        print(f"     已应用: {result['revision_count']}")

        # 保存文档
        output_path = 'test_sample_revised_tc.wps'
        doc_data['document'].save(output_path)

        print(f"[OK] 文档已保存: {output_path}")
        print(f"     文件大小: {os.path.getsize(output_path)} bytes")

        # 验证修订标记
        print("\n验证修订标记:")
        revised_doc = Document(output_path)
        for para in revised_doc.paragraphs:
            if '风险' in para.text or '进展顺利' in para.text:
                print(f"  段落: {para.text[:60]}...")
                # 检查XML中是否有修订标记
                xml_str = str(para._element)
                has_del = 'w:del' in xml_str
                has_ins = 'w:ins' in xml_str
                print(f"    - 删除标记(w:del): {'[OK]' if has_del else '[无]'}")
                print(f"    - 插入标记(w:ins): {'[OK]' if has_ins else '[无]'}")

    except Exception as e:
        import traceback
        print(f"[FAIL] 修订失败: {e}")
        traceback.print_exc()

    # 测试WPSX修订
    print("\n" + "-" * 70)
    print("[步骤3] 测试WPSX文档修订（使用修订模式）")
    print("-" * 70)

    try:
        # 解析文档
        doc_data = DocumentParser.parse_document(wpsx_path)
        print(f"[OK] 文档解析成功，{len(doc_data['paragraphs'])}个段落")

        # 使用TrackChangesReviser
        reviser = TrackChangesReviser(doc_data)
        result = reviser.apply_suggestions(suggestions)

        print(f"[OK] 修订应用成功")
        print(f"     总建议: {result['total_suggestions']}")
        print(f"     已应用: {result['revision_count']}")

        # 保存文档
        output_path = 'test_sample_revised_tc.wpsx'
        doc_data['document'].save(output_path)

        print(f"[OK] 文档已保存: {output_path}")
        print(f"     文件大小: {os.path.getsize(output_path)} bytes")

        # 验证修订标记
        print("\n验证修订标记:")
        revised_doc = Document(output_path)
        for para in revised_doc.paragraphs:
            if '风险' in para.text or '进展顺利' in para.text:
                print(f"  段落: {para.text[:60]}...")
                # 检查XML中是否有修订标记
                xml_str = str(para._element)
                has_del = 'w:del' in xml_str
                has_ins = 'w:ins' in xml_str
                print(f"    - 删除标记(w:del): {'[OK]' if has_del else '[无]'}")
                print(f"    - 插入标记(w:ins): {'[OK]' if has_ins else '[无]'}")

    except Exception as e:
        import traceback
        print(f"[FAIL] 修订失败: {e}")
        traceback.print_exc()

    print("\n" + "=" * 70)
    print("测试完成")
    print("=" * 70)
    print("\n说明:")
    print("  - 修订后的文档在WPS/Word中会显示修订标记")
    print("  - 删除的文本显示为删除线（红色）")
    print("  - 插入的文本显示为下划线（蓝色）")
    print("  - 可在审阅选项卡中接受/拒绝修订")


if __name__ == "__main__":
    test_wps_revision_with_track_changes()
