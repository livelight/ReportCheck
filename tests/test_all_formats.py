"""
多格式文档功能测试 - 测试DOCX、WPS、WPSX文件
"""
import json
import os
import sys
from docx import Document

def create_test_files():
    """创建测试文档（DOCX、WPS、WPSX）"""
    print("="*70)
    print("创建测试文档")
    print("="*70)
    
    # 创建DOCX
    docx_path = "test_all_formats.docx"
    doc = Document()
    doc.add_heading('格式兼容性测试文档', level=1)
    doc.add_paragraph('这是一段测试文本，用于验证文档读取功能。')
    doc.add_heading('一、测试章节', level=2)
    doc.add_paragraph('测试段落一：系统崩溃导致服务中断。')
    doc.add_paragraph('测试段落二：需要修改的内容在这里。')
    doc.add_heading('二、表格测试', level=2)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = '项目'
    table.cell(0, 1).text = '状态'
    table.cell(1, 0).text = '测试'
    table.cell(1, 1).text = '通过'
    doc.save(docx_path)
    print(f"[OK] 创建DOCX: {docx_path}")
    
    # WPS和WPSX实际上是OOXML格式，可以复制DOCX后改扩展名
    wps_path = "test_all_formats.wps"
    wpsx_path = "test_all_formats.wpsx"
    
    import shutil
    shutil.copy(docx_path, wps_path)
    shutil.copy(docx_path, wpsx_path)
    
    print(f"[OK] 创建WPS: {wps_path}")
    print(f"[OK] 创建WPSX: {wpsx_path}")
    
    return docx_path, wps_path, wpsx_path

def test_read_all_formats(docx_path, wps_path, wpsx_path):
    """测试读取所有格式"""
    print("\n" + "="*70)
    print("测试文档读取功能")
    print("="*70)
    
    from unified_mcp_server import DocumentReader
    reader = DocumentReader()
    
    results = {}
    for file_path in [docx_path, wps_path, wpsx_path]:
        ext = os.path.splitext(file_path)[1].upper()
        print(f"\n[{ext}] 测试文件: {file_path}")
        try:
            result = reader.read_document(file_path)
            results[ext] = result
            print(f"  读取成功: {result['success']}")
            print(f"  段落数: {result['document_metadata']['paragraph_count']}")
            print(f"  表格数: {result['document_metadata']['table_count']}")
            print(f"  内容长度: {len(result['document_content'])} 字符")
        except Exception as e:
            print(f"  [FAIL] 读取失败: {e}")
            results[ext] = None
    
    return results

def test_revision_all_formats(docx_path, wps_path, wpsx_path):
    """测试修订所有格式"""
    print("\n" + "="*70)
    print("测试文档修订功能（所有格式）")
    print("="*70)
    
    from unified_mcp_server import DocumentReviser
    reviser = DocumentReviser()
    
    suggestions = [
        {
            'id': 'test_001',
            'rule_id': 'TEST_001',
            'original_text': '系统崩溃',
            'suggestion': '系统故障'
        },
        {
            'id': 'test_002',
            'rule_id': 'TEST_002',
            'original_text': '需要修改的内容',
            'suggestion': '已修改的内容'
        }
    ]
    
    results = {}
    for file_path in [docx_path, wps_path, wpsx_path]:
        ext = os.path.splitext(file_path)[1].upper()
        output_path = f"test_all_formats_revised{ext.lower()}"
        
        print(f"\n[{ext}] 修订测试")
        print(f"  输入: {file_path}")
        print(f"  输出: {output_path}")
        
        try:
            result = reviser.revise_document(
                file_path,
                json.dumps(suggestions),
                output_path,
                'true'  # 启用修订模式
            )
            results[ext] = result
            print(f"  修订成功: {result['success']}")
            print(f"  应用修订数: {result['applied_revisions']}/{result['total_suggestions']}")
            
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                print(f"  输出文件大小: {file_size} 字节")
                
                # 验证文件可被读取
                from unified_mcp_server import DocumentReader
                try:
                    verify_reader = DocumentReader()
                    verify_result = verify_reader.read_document(output_path)
                    print(f"  验证读取: {verify_result['success']}")
                except Exception as e:
                    print(f"  [WARN] 验证读取失败: {e}")
            
        except Exception as e:
            print(f"  [FAIL] 修订失败: {e}")
            results[ext] = None
    
    return results

def test_segmentation_all_formats(docx_path, wps_path, wpsx_path):
    """测试分段读取所有格式"""
    print("\n" + "="*70)
    print("测试文档分段功能（所有格式）")
    print("="*70)
    
    from document_reader_mcp import DocumentReaderMCPNode
    node = DocumentReaderMCPNode(max_tokens_per_segment=100)
    
    results = {}
    for file_path in [docx_path, wps_path, wpsx_path]:
        ext = os.path.splitext(file_path)[1].upper()
        print(f"\n[{ext}] 分段测试: {file_path}")
        
        try:
            result = node.read_document_segmented(file_path, max_tokens_per_segment=100)
            results[ext] = result
            print(f"  分段成功: {result['success']}")
            
            if 'segment_info' in result:
                info = result['segment_info']
                print(f"  总分段数: {info['total_segments']}")
                print(f"  总token数: {info['total_tokens']}")
                print(f"  文本分段: {info.get('text_segments', 'N/A')}")
                print(f"  表格分段: {info.get('table_segments', 'N/A')}")
            
            if 'segments' in result:
                print(f"  分段详情:")
                for seg in result['segments'][:3]:  # 只显示前3个分段
                    preview = seg['content'][:50].replace('\n', ' ')
                    print(f"    - 段{seg['segment_id']}: {preview}... ({seg['estimated_tokens']} tokens)")
                if len(result['segments']) > 3:
                    print(f"    ... 还有 {len(result['segments'])-3} 个分段")
                    
        except Exception as e:
            print(f"  [FAIL] 分段失败: {e}")
            import traceback
            traceback.print_exc()
            results[ext] = None
    
    return results

def cleanup_test_files():
    """清理测试文件"""
    print("\n" + "="*70)
    print("清理测试文件")
    print("="*70)
    
    test_files = [
        "test_all_formats.docx",
        "test_all_formats.wps",
        "test_all_formats.wpsx",
        "test_all_formats_revised.docx",
        "test_all_formats_revised.wps",
        "test_all_formats_revised.wpsx"
    ]
    
    for f in test_files:
        if os.path.exists(f):
            os.remove(f)
            print(f"  [OK] 删除: {f}")
        else:
            print(f"  [SKIP] 不存在: {f}")

def main():
    print("="*70)
    print("多格式文档功能测试 (DOCX / WPS / WPSX)")
    print("="*70)
    
    try:
        # 创建测试文件
        docx_path, wps_path, wpsx_path = create_test_files()
        
        # 测试读取
        read_results = test_read_all_formats(docx_path, wps_path, wpsx_path)
        
        # 测试修订
        revise_results = test_revision_all_formats(docx_path, wps_path, wpsx_path)
        
        # 测试分段
        seg_results = test_segmentation_all_formats(docx_path, wps_path, wpsx_path)
        
        # 汇总结果
        print("\n" + "="*70)
        print("测试汇总")
        print("="*70)
        
        all_pass = True
        for ext in ['.DOCX', '.WPS', '.WPSX']:
            print(f"\n{ext} 格式:")
            read_ok = read_results.get(ext) is not None and read_results[ext].get('success')
            revise_ok = revise_results.get(ext) is not None and revise_results[ext].get('success')
            seg_ok = seg_results.get(ext) is not None and seg_results[ext].get('success')
            
            print(f"  读取功能: {'[PASS]' if read_ok else '[FAIL]'}")
            print(f"  修订功能: {'[PASS]' if revise_ok else '[FAIL]'}")
            print(f"  分段功能: {'[PASS]' if seg_ok else '[FAIL]'}")
            
            if not (read_ok and revise_ok and seg_ok):
                all_pass = False
        
        # 清理
        cleanup_test_files()
        
        print("\n" + "="*70)
        if all_pass:
            print("[SUCCESS] 所有格式测试通过!")
            return 0
        else:
            print("[FAILED] 部分格式测试失败")
            return 1
            
    except Exception as e:
        print(f"\n[ERROR] 测试过程异常: {e}")
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    sys.exit(main())
