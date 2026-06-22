"""
测试 WPS/WPSX 文档修订功能
创建符合WPS规范的完整文档结构
"""

import json
import os
import zipfile
from xml.etree import ElementTree as ET
from docx import Document
from revision_mcp_node import (
    FASTMCP_AVAILABLE,
    DOCX_AVAILABLE,
    CheckSuggestion,
    DocumentParser,
    TrackChangesReviser
)


def create_sample_wps():
    """创建符合WPS规范的WPS文档 (.wps) - 使用docx库生成OOXML格式"""
    path = 'test_sample.wps'

    # 使用python-docx创建文档
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 添加内容
    doc.add_heading('企业文档管理系统项目报告', level=0)
    doc.add_heading('一、项目概述', level=1)
    doc.add_paragraph('本项目旨在开发一套企业级文档管理系统。')
    doc.add_heading('二、项目目标', level=1)
    doc.add_paragraph('1. 实现文档的集中存储和分类管理')
    doc.add_paragraph('2. 支持多用户协作编辑')
    doc.add_heading('三、当前进展', level=1)
    doc.add_paragraph('目前已完成核心功能的开发和测试工作，进展顺利。')
    doc.add_heading('四、风险分析', level=1)
    doc.add_paragraph('项目进展顺利，暂无风险。')
    doc.add_heading('五、下一步计划', level=1)
    doc.add_paragraph('继续优化系统性能，并进行用户验收测试。')

    # 保存为docx格式，然后重命名为.wps
    temp_docx = path.replace('.wps', '_temp.docx')
    doc.save(temp_docx)

    # 将docx内容复制到.wps文件
    with zipfile.ZipFile(temp_docx, 'r') as zf_src:
        with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zf_dst:
            for item in zf_src.namelist():
                zf_dst.writestr(item, zf_src.read(item))

    # 删除临时文件
    try:
        os.remove(temp_docx)
    except:
        pass

    return path


def create_sample_wpsx():
    """创建符合WPS规范的WPSX文档 (.wpsx) - 使用类似docx的结构"""
    path = 'test_sample.wpsx'

    # WPSX使用word/document.xml结构
    document_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
  <w:body>
    <w:p>
      <w:pPr><w:jc w:val="center"/></w:pPr>
      <w:r>
        <w:rPr><w:b/><w:sz w:val="32"/></w:rPr>
        <w:t>企业文档管理系统项目报告</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr><w:outlineLvl w:val="0"/></w:pPr>
      <w:r>
        <w:rPr><w:b/><w:sz w:val="28"/></w:rPr>
        <w:t>一、项目概述</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:rPr><w:sz w:val="24"/></w:rPr>
        <w:t>本项目旨在开发一套企业级文档管理系统。</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr><w:outlineLvl w:val="0"/></w:pPr>
      <w:r>
        <w:rPr><w:b/><w:sz w:val="28"/></w:rPr>
        <w:t>二、项目目标</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:rPr><w:sz w:val="24"/></w:rPr>
        <w:t>1. 实现文档的集中存储和分类管理</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:rPr><w:sz w:val="24"/></w:rPr>
        <w:t>2. 支持多用户协作编辑</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr><w:outlineLvl w:val="0"/></w:pPr>
      <w:r>
        <w:rPr><w:b/><w:sz w:val="28"/></w:rPr>
        <w:t>三、当前进展</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:rPr><w:sz w:val="24"/></w:rPr>
        <w:t>目前已完成核心功能的开发和测试工作，进展顺利。</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr><w:outlineLvl w:val="0"/></w:pPr>
      <w:r>
        <w:rPr><w:b/><w:sz w:val="28"/></w:rPr>
        <w:t>四、风险分析</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:rPr><w:sz w:val="24"/></w:rPr>
        <w:t>项目进展顺利，暂无风险。</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:pPr><w:outlineLvl w:val="0"/></w:pPr>
      <w:r>
        <w:rPr><w:b/><w:sz w:val="28"/></w:rPr>
        <w:t>五、下一步计划</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:rPr><w:sz w:val="24"/></w:rPr>
        <w:t>继续优化系统性能，并进行用户验收测试。</w:t>
      </w:r>
    </w:p>
    <w:sectPr>
      <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="1440" w:right="1800" w:bottom="1440" w:left="1800"/>
    </w:sectPr>
  </w:body>
</w:document>'''

    # 样式定义
    styles_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:docDefaults>
    <w:rPrDefault>
      <w:rPr>
        <w:rFonts w:ascii="宋体" w:eastAsia="宋体" w:hAnsi="宋体"/>
        <w:sz w:val="24"/>
        <w:szCs w:val="24"/>
      </w:rPr>
    </w:rPrDefault>
  </w:docDefaults>
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:rPr>
      <w:rFonts w:ascii="宋体" w:eastAsia="宋体"/>
      <w:sz w:val="24"/>
    </w:rPr>
  </w:style>
</w:styles>'''

    # 内容类型定义
    content_types_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>'''

    # 根关系定义
    rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''

    # word目录下的关系定义
    word_rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>'''

    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('word/document.xml', document_xml)
        zf.writestr('word/styles.xml', styles_xml)
        zf.writestr('word/_rels/document.xml.rels', word_rels_xml)
        zf.writestr('[Content_Types].xml', content_types_xml)
        zf.writestr('_rels/.rels', rels_xml)

    return path


def test_wps_parsing():
    """测试WPS文档解析"""
    print("=" * 70)
    print("WPS/WPSX 文档解析测试")
    print("=" * 70)

    # 创建测试文档
    print("\n[步骤1] 创建测试文档")
    wps_path = create_sample_wps()
    wpsx_path = create_sample_wpsx()

    print(f"  WPS文件: {wps_path} ({os.path.getsize(wps_path)} bytes)")
    print(f"  WPSX文件: {wpsx_path} ({os.path.getsize(wpsx_path)} bytes)")

    # 测试解析
    print("\n[步骤2] 解析WPS文档")
    try:
        wps_data = DocumentParser.parse_document(wps_path)
        print(f"  [OK] 解析成功")
        print(f"       类型: {wps_data['type']}")
        print(f"       段落数: {len(wps_data['paragraphs'])}")
        print("\n  内容预览:")
        for i, para in enumerate(wps_data['paragraphs'][:6]):
            text = para['text'][:50]
            print(f"    {i+1}. {text}")
    except Exception as e:
        print(f"  [FAIL] 解析失败: {e}")

    print("\n[步骤3] 解析WPSX文档")
    try:
        wpsx_data = DocumentParser.parse_document(wpsx_path)
        print(f"  [OK] 解析成功")
        print(f"       类型: {wpsx_data['type']}")
        print(f"       段落数: {len(wpsx_data['paragraphs'])}")
        print("\n  内容预览:")
        for i, para in enumerate(wpsx_data['paragraphs'][:6]):
            text = para['text'][:50]
            print(f"    {i+1}. {text}")
    except Exception as e:
        print(f"  [FAIL] 解析失败: {e}")


def test_wps_revision():
    """测试WPS文档修订"""
    print("\n" + "=" * 70)
    print("WPS/WPSX 文档修订测试")
    print("=" * 70)

    # 创建测试文档
    wps_path = 'test_sample.wps'
    wpsx_path = 'test_sample.wpsx'

    if not os.path.exists(wps_path):
        create_sample_wps()
    if not os.path.exists(wpsx_path):
        create_sample_wpsx()

    # 定义检查建议
    suggestions = [
        CheckSuggestion(
            id="ref_001",
            rule_id="CONTENT_023",
            rule_name="风险分析缺失",
            type="content",
            severity="High",
            section="sec_004",
            original_text="项目进展顺利，暂无风险。",
            suggestion="当前主要风险及应对措施：\n1. 人员流动性风险\n2. 技术架构风险",
            reason="风险分析过于简单"
        ),
        CheckSuggestion(
            id="ref_002",
            rule_id="LANG_001",
            rule_name="表述不规范",
            type="language",
            severity="Medium",
            section="sec_003",
            original_text="目前已完成核心功能的开发和测试工作，进展顺利。",
            suggestion="目前已完成核心功能的开发和内部测试，整体质量符合预期。",
            reason="表述不够专业"
        )
    ]

    print(f"\n检查建议:")
    for s in suggestions:
        print(f"  [{s.severity}] {s.rule_name}")

    # 测试WPS修订
    print("\n[步骤1] 测试WPS文档修订")
    try:
        from revision_mcp_node import DocumentParser, TrackChangesReviser

        # 解析文档
        doc_data = DocumentParser.parse_document(wps_path)
        print(f"  [OK] 文档解析成功，{len(doc_data['paragraphs'])}个段落")

        # 应用修订
        reviser = TrackChangesReviser(doc_data)
        result = reviser.apply_suggestions(suggestions)

        print(f"  [OK] 修订应用成功")
        print(f"       总建议: {result['total_suggestions']}")
        print(f"       已应用: {result['revision_count']}")

        # 保存文档
        output_path = 'test_sample_revised.wps'
        import zipfile
        with zipfile.ZipFile(wps_path, 'r') as src_zf:
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as dst_zf:
                for name in src_zf.namelist():
                    if name == 'content.xml':
                        # 需要重新生成content.xml
                        content = _generate_wps_content(doc_data)
                        dst_zf.writestr(name, content)
                    else:
                        dst_zf.writestr(name, src_zf.read(name))

        print(f"  [OK] 文档已保存: {output_path}")
        print(f"       文件大小: {os.path.getsize(output_path)} bytes")

    except Exception as e:
        import traceback
        print(f"  [FAIL] 修订失败: {e}")
        traceback.print_exc()

    # 测试WPSX修订
    print("\n[步骤2] 测试WPSX文档修订")
    try:
        # 解析文档
        doc_data = DocumentParser.parse_document(wpsx_path)
        print(f"  [OK] 文档解析成功，{len(doc_data['paragraphs'])}个段落")

        # 应用修订
        reviser = TrackChangesReviser(doc_data)
        result = reviser.apply_suggestions(suggestions)

        print(f"  [OK] 修订应用成功")
        print(f"       总建议: {result['total_suggestions']}")
        print(f"       已应用: {result['revision_count']}")

        # 保存文档
        output_path = 'test_sample_revised.wpsx'
        import zipfile
        with zipfile.ZipFile(wpsx_path, 'r') as src_zf:
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as dst_zf:
                for name in src_zf.namelist():
                    if name == 'word/document.xml':
                        # 需要重新生成document.xml
                        content = _generate_wpsx_content(doc_data)
                        dst_zf.writestr(name, content)
                    else:
                        dst_zf.writestr(name, src_zf.read(name))

        print(f"  [OK] 文档已保存: {output_path}")
        print(f"       文件大小: {os.path.getsize(output_path)} bytes")

    except Exception as e:
        import traceback
        print(f"  [FAIL] 修订失败: {e}")
        traceback.print_exc()

    print("\n" + "=" * 70)
    print("测试完成")
    print("=" * 70)


def _generate_wps_content(doc_data):
    """生成WPS内容XML"""
    paragraphs = doc_data.get('paragraphs', [])
    def escape_xml(text):
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    xml_parts = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>']
    xml_parts.append('<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">')
    xml_parts.append('<w:body>')
    for para in paragraphs:
        text = escape_xml(para['text'])
        xml_parts.append(f'<w:p><w:r><w:t>{text}</w:t></w:r></w:p>')
    xml_parts.append('</w:body></w:document>')
    return '\n'.join(xml_parts)


def _generate_wpsx_content(doc_data):
    """生成WPSX内容XML"""
    paragraphs = doc_data.get('paragraphs', [])
    def escape_xml(text):
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    xml_parts = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>']
    xml_parts.append('<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">')
    xml_parts.append('<w:body>')
    for para in paragraphs:
        text = escape_xml(para['text'])
        xml_parts.append(f'<w:p><w:r><w:t>{text}</w:t></w:r></w:p>')
    xml_parts.append('</w:body></w:document>')
    return '\n'.join(xml_parts)


if __name__ == "__main__":
    test_wps_parsing()
    test_wps_revision()
