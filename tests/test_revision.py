"""
测试脚本 - 文档修订MCP节点测试
"""

import json
import os
import sys
import tempfile
from docx import Document

# 添加项目根目录到模块搜索路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from revision_mcp_node import (
    RevisionMCPNode,
    CheckSuggestion as OldCheckSuggestion,
    DocumentParser,
    TrackChangesReviser,
    FASTMCP_AVAILABLE,
    DOCX_AVAILABLE
)


def test_revision_mcp_node():
    """测试RevisionMCPNode - 使用revise_document接口"""
    import json as json_mod
    from docx import Document as DocxDocument
    
    # 创建测试文档
    doc = DocxDocument()
    doc.add_heading('测试文档', level=1)
    doc.add_paragraph('这是需要修改的原文内容。')
    doc.add_paragraph('项目进展顺利，暂无风险。')
    test_path = 'test_revision_temp.docx'
    doc.save(test_path)

    # 直接使用RevisionMCPNode的revise_document方法
    node = RevisionMCPNode()
    suggestions = [
        {
            'id': 'test_001',
            'rule_id': 'CONTENT_001',
            'rule_name': '测试修订',
            'type': 'content',
            'severity': 'High',
            'section': 'sec_001',
            'original_text': '这是需要修改的原文内容。',
            'suggestion': '这是修改后的新内容。',
            'reason': '测试原因'
        }
    ]
    result = node.revise_document(
        file_path=test_path,
        suggestions_json=json_mod.dumps(suggestions, ensure_ascii=False),
        output_path='test_revision_output.docx',
        use_track_changes="false"
    )

    # 输出结果
    print("\n[修订结果]")
    print(json.dumps(result, indent=2, ensure_ascii=False))

    # 验证输出
    if result.get('success'):
        output_path = result.get('output_path', '')
        if os.path.exists(output_path):
            print(f"\n[验证通过] 修订后文档已生成: {output_path}")
            print(f"文件大小: {os.path.getsize(output_path)} bytes")

            # 读取修订后的文档内容
            from docx import Document
            revised_doc = Document(output_path)
            print("\n[修订后文档内容预览]")
            for i, para in enumerate(revised_doc.paragraphs[:10]):
                if para.text.strip():
                    print(f"  {para.text[:80]}")
        else:
            print(f"\n[警告] 输出文件不存在: {output_path}")
    else:
        print(f"\n[错误] 修订失败: {result.get('error')}")

    # 清理
    print("\n" + "=" * 60)
    print("测试完成")
    print("=" * 60)

    return result


def test_json_parsing():
    """测试JSON解析"""
    print("\n[测试JSON解析]")

    test_cases = [
        # 正常JSON
        '[{"id":"001","rule_id":"R001","rule_name":"测试规则","type":"format","severity":"High","section":"sec1","original_text":"原文","suggestion":"建议","reason":"原因"}]',
        # 空数组
        '[]',
        # 无效JSON
        '{invalid json}'
    ]

    node = RevisionMCPNode()

    for i, test_json in enumerate(test_cases):
        print(f"\n测试用例 {i+1}:")
        print(f"  输入: {test_json[:50]}...")
        try:
            suggestions = json.loads(test_json)
            print(f"  结果: 解析成功，共{len(suggestions)}条建议")
        except json.JSONDecodeError as e:
            print(f"  结果: JSON解析错误 - {e}")


def test_suggestion_model():
    """测试建议数据模型"""
    print("\n[测试建议数据模型]")

    test_data = {
        "id": "ref_001",
        "rule_id": "CONTENT_023",
        "rule_name": "缺少风险分析",
        "type": "content",
        "severity": "High",
        "section": "sec_005",
        "original_text": "项目进展顺利，暂无风险。",
        "suggestion": "项目整体进展顺利，目前主要风险包括：1）人员流动性风险；2）技术架构调整风险。",
        "reason": "风险分析过于简单"
    }

    suggestion = OldCheckSuggestion.from_dict(test_data)

    print(f"  ID: {suggestion.id}")
    print(f"  规则名称: {suggestion.rule_name}")
    print(f"  类型: {suggestion.type}")
    print(f"  严重程度: {suggestion.severity}")
    print(f"  原始文本: {suggestion.original_text[:30]}...")
    print(f"  修订建议: {suggestion.suggestion[:30]}...")

    print("\n[测试通过]")


if __name__ == "__main__":
    # 运行所有测试
    test_suggestion_model()
    test_json_parsing()
    result = test_revision_mcp_node()
