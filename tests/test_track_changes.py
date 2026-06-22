"""
测试修订模式（Track Changes）功能
"""

import json
import os
from docx import Document
from revision_mcp_node import (
    FASTMCP_AVAILABLE,
    DOCX_AVAILABLE,
    CheckSuggestion,
    DocumentParser,
    TrackChangesReviser
)


def create_sample_docx():
    """创建示例Word文档"""
    doc = Document()

    doc.add_heading('企业文档管理系统项目报告', level=1)
    doc.add_heading('一、项目概述', level=2)
    doc.add_paragraph('本项目旨在开发一套企业级文档管理系统。')
    doc.add_heading('二、项目目标', level=2)
    doc.add_paragraph('1. 实现文档的集中存储和分类管理')
    doc.add_paragraph('2. 支持多用户协作编辑')
    doc.add_heading('三、当前进展', level=2)
    doc.add_paragraph('目前已完成核心功能的开发和测试工作，进展顺利。')
    doc.add_heading('四、风险分析', level=2)
    doc.add_paragraph('项目进展顺利，暂无风险。')
    doc.add_heading('五、下一步计划', level=2)
    doc.add_paragraph('继续优化系统性能，并进行用户验收测试。')

    path = 'test_sample.docx'
    doc.save(path)
    return path


def test_track_changes():
    """测试修订模式"""
    print("=" * 70)
    print("修订模式（Track Changes）测试")
    print("=" * 70)

    print(f"\n依赖检查:")
    print(f"  fastmcp: {'[OK]' if FASTMCP_AVAILABLE else '[FAIL]'}")
    print(f"  python-docx: {'[OK]' if DOCX_AVAILABLE else '[FAIL]'}")

    if not FASTMCP_AVAILABLE or not DOCX_AVAILABLE:
        print("\n[ERROR] 缺少必要依赖，无法测试")
        return

    # 创建测试文档
    print("\n" + "-" * 70)
    print("[步骤1] 创建测试文档")
    print("-" * 70)

    docx_path = create_sample_docx()
    print(f"文档路径: {docx_path}")
    print(f"文件大小: {os.path.getsize(docx_path)} bytes")

    # 显示原文档内容
    print("\n原文档内容:")
    original_doc = Document(docx_path)
    for i, para in enumerate(original_doc.paragraphs[:8]):
        if para.text.strip():
            print(f"  {para.text[:60]}")

    # 定义检查建议
    print("\n" + "-" * 70)
    print("[步骤2] 定义检查建议")
    print("-" * 70)

    suggestions = [
        CheckSuggestion(
            id="ref_001",
            rule_id="CONTENT_023",
            rule_name="风险分析缺失",
            type="content",
            severity="High",
            section="sec_004",
            original_text="项目进展顺利，暂无风险。",
            suggestion="当前主要风险及应对措施：\n1. 人员流动性风险：加强知识文档建设\n2. 技术架构风险：预留技术预研时间",
            reason="风险分析过于简单"
        ),
        CheckSuggestion(
            id="ref_002",
            rule_id="LANG_001",
            rule_name="表述不规范",
            type="language",
            severity="Medium",
            section="sec_003",
            original_text="目前已完成核心功能的开发和测试工作，进展顺利。",
            suggestion="目前已完成核心功能的开发和内部测试，整体质量符合预期。",
            reason="'进展顺利'表述不够专业"
        )
    ]

    print(f"建议数量: {len(suggestions)}")
    for s in suggestions:
        print(f"  [{s.severity}] {s.rule_name}")
        print(f"    原文: {s.original_text[:40]}...")
        print(f"    建议: {s.suggestion[:40]}...")

    # 测试修订模式
    print("\n" + "-" * 70)
    print("[步骤3] 应用修订（使用修订模式）")
    print("-" * 70)

    try:
        # 解析文档
        doc_data = DocumentParser.parse_document(docx_path)
        print("[OK] 文档解析成功")

        # 使用TrackChangesReviser
        reviser = TrackChangesReviser(doc_data)
        result = reviser.apply_suggestions(suggestions)

        print(f"[OK] 修订应用成功")
        print(f"     总建议: {result['total_suggestions']}")
        print(f"     成功: {result['revision_count']}")

        # 保存文档
        output_path = 'test_sample_revised_track_changes.docx'
        doc_data['document'].save(output_path)
        print(f"[OK] 文档已保存: {output_path}")

        # 验证输出
        if os.path.exists(output_path):
            print(f"[OK] 输出文件存在, 大小: {os.path.getsize(output_path)} bytes")

        # 读取修订后内容（检查XML）
        print("\n修订详情:")
        for rev in result['revisions']:
            print(f"  [{rev['severity']:7}] {rev['rule_name']} - {rev['status']}")
            print(f"    原文: {rev['original']}")
            print(f"    新文: {rev['new']}")

        print("\n" + "=" * 70)
        print("[说明] 修订后的文档使用Word修订追踪功能标记修改")
        print("       - 删除的文本显示为删除线（红色）")
        print("       - 插入的文本显示为下划线（蓝色）")
        print("       - 可在Word中接受/拒绝修订")
        print("=" * 70)

        return True

    except Exception as e:
        import traceback
        print(f"[FAIL] 修订失败: {e}")
        traceback.print_exc()
        return False


if __name__ == "__main__":
    test_track_changes()
