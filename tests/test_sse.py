"""
SSE服务器测试脚本

使用方法:
1. 启动SSE服务器: python mcp_sse_server.py
2. 运行测试: python test_sse.py
"""

import requests
import json
import sys

BASE_URL = "http://localhost:8000"


def test_health():
    """测试健康检查"""
    print("=" * 60)
    print("测试1: 健康检查")
    print("-" * 60)

    try:
        response = requests.get(f"{BASE_URL}/health", timeout=5)
        print(f"状态码: {response.status_code}")
        print(f"响应: {json.dumps(response.json(), indent=2, ensure_ascii=False)}")
        return response.status_code == 200
    except Exception as e:
        print(f"错误: {e}")
        return False


def test_tools():
    """测试工具定义"""
    print("\n" + "=" * 60)
    print("测试2: 工具定义")
    print("-" * 60)

    try:
        response = requests.get(f"{BASE_URL}/tools", timeout=5)
        print(f"状态码: {response.status_code}")
        data = response.json()
        print(f"工具名称: {data.get('name')}")
        print(f"工具数量: {len(data.get('tools', []))}")
        return response.status_code == 200
    except Exception as e:
        print(f"错误: {e}")
        return False


def test_sync_revise():
    """测试同步修订接口"""
    print("\n" + "=" * 60)
    print("测试3: 同步修订接口")
    print("-" * 60)

    # 先创建一个测试文档
    from docx import Document
    doc = Document()
    doc.add_heading("测试文档", level=1)
    doc.add_paragraph("这是一段测试文本。")
    doc.add_paragraph("项目进展顺利，暂无风险。")
    doc.save("test_sse_doc.docx")
    print("已创建测试文档: test_sse_doc.docx")

    suggestions = [
        {
            "id": "ref_001",
            "rule_id": "CONTENT_001",
            "rule_name": "缺少风险分析",
            "type": "content",
            "severity": "High",
            "section": "风险分析",
            "original_text": "项目进展顺利，暂无风险。",
            "suggestion": "项目整体进展顺利，目前主要风险包括：人员流动性风险。",
            "reason": "风险分析过于简单"
        }
    ]

    try:
        response = requests.post(
            f"{BASE_URL}/revise",
            json={
                "file_path": "test_sse_doc.docx",
                "suggestions_json": json.dumps(suggestions),
                "output_path": "test_sse_doc_revised.docx"
            },
            timeout=30
        )
        print(f"状态码: {response.status_code}")
        data = response.json()
        print(f"成功: {data.get('success')}")
        print(f"输出路径: {data.get('output_path')}")
        print(f"应用修订数: {data.get('applied_revisions')}")
        return data.get('success', False)
    except Exception as e:
        print(f"错误: {e}")
        return False


def test_sse_endpoint():
    """测试SSE端点"""
    print("\n" + "=" * 60)
    print("测试4: SSE端点连接")
    print("-" * 60)

    try:
        import sseclient
        import urllib3

        http = urllib3.PoolManager()
        response = http.request(
            'GET',
            f"{BASE_URL}/sse",
            preload_content=False,
            headers={'Accept': 'text/event-stream'}
        )

        client = sseclient.SSEClient(response)
        print("已连接到SSE端点")

        # 读取几个事件
        events = []
        for event in client.events():
            events.append(event)
            print(f"  事件类型: {event.event}")
            if len(events) >= 2:
                break

        response.release_conn()
        return len(events) > 0

    except ImportError:
        print("需要安装 sseclient-py: pip install sseclient-py")
        print("跳过SSE测试")
        return True
    except Exception as e:
        print(f"错误: {e}")
        return False


def cleanup():
    """清理测试文件"""
    import os
    files = ["test_sse_doc.docx", "test_sse_doc_revised.docx"]
    for f in files:
        if os.path.exists(f):
            os.remove(f)
            print(f"已清理: {f}")


def main():
    print("=" * 60)
    print("MCP SSE服务器测试")
    print("=" * 60)
    print(f"测试地址: {BASE_URL}")
    print("=" * 60)

    # 检查服务器是否运行
    try:
        requests.get(f"{BASE_URL}/health", timeout=2)
    except:
        print("\n错误: 无法连接到SSE服务器")
        print("请先启动服务器: python mcp_sse_server.py")
        sys.exit(1)

    # 运行测试
    results = []

    results.append(("健康检查", test_health()))
    results.append(("工具定义", test_tools()))
    results.append(("同步修订", test_sync_revise()))
    results.append(("SSE端点", test_sse_endpoint()))

    # 清理
    print("\n" + "=" * 60)
    print("清理测试文件")
    print("-" * 60)
    cleanup()

    # 汇总
    print("\n" + "=" * 60)
    print("测试结果汇总")
    print("=" * 60)
    for name, passed in results:
        status = "通过" if passed else "失败"
        print(f"  {name}: {status}")

    all_passed = all(r[1] for r in results)
    print("-" * 60)
    print(f"总计: {len(results)}项, 通过: {sum(r[1] for r in results)}项")
    print("=" * 60)

    return 0 if all_passed else 1


if __name__ == "__main__":
    sys.exit(main())
